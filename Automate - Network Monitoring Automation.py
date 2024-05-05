import time
import tkinter
from tkinter import messagebox
import customtkinter
from tkinter import filedialog
import paramiko
import pyautogui
from customtkinter import *
from tkinter import StringVar
import os
from selenium import webdriver
from selenium.webdriver import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.chrome.service import Service
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from datetime import date
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import win32com.client.gencache
from pathlib import Path
import glob
from selenium.webdriver.chrome.options import Options
from PIL import Image
import datetime
import win32com.client as win32
from tkinter import simpledialog
from selenium.webdriver.support.ui import Select

customtkinter.set_appearance_mode("System")
customtkinter.set_default_color_theme("green")


class MainFrame(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        # configure window
        self.checklist = None
        self.usernameform = None
        self.mainbutton = None
        self.titlelabel = None
        self.mainpageframe = None
        self.checklistform = None
        self.driver = None
        self.codeIPAddNodeNCR = None
        self.configframe = None
        self.nceframe = None
        self.codepasswordpdh = None
        self.codepdhaddress1 = None
        self.codeusernamepdh = None
        self.pdhframe = None
        self.codemin2 = None
        self.codemin1 = None
        self.codepasswordmin = None
        self.codeusernamemin = None
        self.prtgminframe = None
        self.ncmframe = None
        self.exfoframe = None
        self.cnframe = None
        self.cnframe = None
        self.idmframe = None
        self.meraframe = None
        self.combo1 = None
        self.combomeraki = None
        self.bistemplate = None
        self.ncrminframe = None
        self.conframe = None
        self.contentframe = None
        self.ncrframe = None
        self.button_template = None
        self.button_meraki_idm = None
        self.button_meraki_main = None
        self.button_nms = None
        self.button_ncr_min = None
        self.button_ncr = None
        self.bisframe = None
        self.nmspage = None
        self.Selected_Value2 = None
        self.Selected_Value = None
        self.DepartmentInput = None
        self.OtherInfoInput = None
        self.UserName = None
        self.CircuitIDInput = None
        self.ProjectCodeInput = None
        self.combobox = None
        self.Department = None
        self.OtherInfo = None
        self.PreparedBy = None
        self.ProjectCode = None
        self.circuitidmin = None
        self.apnamemin = None
        self.passwordmin = None
        self.usernamemin = None
        self.mainframe = None
        self.userpass = None
        self.usermin = None
        self.userAP = None
        self.usercircuitid = None
        self.usercircuitid1 = None
        self.username = None
        self.userAP1 = None
        self.userpass1 = None
        self.username1 = None
        self.maestro = None
        self.CircuitID = None
        self.APNAME = None
        self.Password = None
        self.ID = None
        self.bot = None
        self.LRD = None
        self.NodeNAME = None
        self.SDWANID = None
        self.NodeNAME = None
        self.PASSWORD = None
        self.conn = None
        self.commands = None
        self.PortNumber = None
        self.IPAddNode = None
        self.noodname = None
        self.options = None
        self.windows = None
        self.s = None
        self.circuitid11 = None
        self.apname11 = None
        self.password11 = None
        self.username11 = None
        self.paramiko = None
        self.codePortNumber = None
        self.codeIPAddNode = None
        self.codeSlotNo = None
        self.codeSlotNo = None
        self.config = None
        self.codesdwanid1 = None
        self.codepassword1 = None
        self.codeusername1 = None
        self.merak = None
        self.codesdwanid = None
        self.codepassword = None
        self.codeusername = None
        self.naan = None
        self.nodename = None
        self.mera = None
        self.nceelements = None
        self.ncepassword = None
        self.nceid = None
        self.nce = None
        self.nodename1 = None
        self.exo = None
        self.lrd1 = None
        self.maes = None
        self.nms = None
        self.nms1 = None
        self.screen = None
        self.DC = None
        self.screen1 = None
        self.circuitid1 = None
        self.password1 = None
        self.username1 = None
        self.mainframe1 = None
        self.content_frame4 = None
        self.filename = None
        self.button1 = None
        self.button5 = None
        self.button4 = None
        self.button3 = None
        self.button2 = None
        self.date = None
        self.dual = None
        self.managed = None
        self.branded = None
        self.content_frame2 = None
        self.level = None
        self.server = None
        self.racky = None
        self.content_frame3 = None
        self.mode = None
        self.number = None
        self.VendorName11 = None
        self.ServiceTypeBW11 = None
        self.UserName11 = None
        self.CircuitID11 = None
        self.CustomerName11 = None
        self.content_frame = None
        self.codeusernamenceip = None
        self.codepasswordnceip = None
        self.codenodenamenceip = None
        self.codeportnonceip = None
        self.MINAO1 = None
        self.MINAO2 = None
        self.usernameMIN11 = None
        self.passwordMIN11 = None
        self.apnameMIN11 = None
        self.AOMIN11 = None
        self.MINAO22 = None
        self.codeusernameConfigNCE = None
        self.codepasswordConfigNCE = None
        self.codeIPAddNodeConfigNCE = None
        self.codePortNumberConfigNCE = None
        self.codePortNoVLANIDConfigNCE = None
        self.s = None
        self.options = None
        self.title("AUTOMATE")
        self.geometry(f"{1200}x{680}")
        self.resizable(False, False)

        self.options = Options()

        # this parameter tells Chrome that
        # it should be run without UI (Headless)
        self.options.add_argument("--disable-infobars")
        self.options.add_argument("--start-maximized")
        self.options.add_argument("--disable-extensions")
        self.options.add_argument('--window-size=1920,1080')
        self.options.add_argument("--ignore-certificate-errors")
        self.options.add_argument('--ignore-ssl-errors=yes')
        prefs = {"download.default_directory": "C:\\Automate\\Checklist Output\\Download Files"}
        self.options.add_experimental_option("prefs", prefs)
        self.options.add_argument("--headless=new")

        self.s = Service('images\\chromedriver.exe')

        # configure grid layout (4x4)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure((2, 3), weight=0)
        self.grid_rowconfigure((0, 1, 2), weight=1)

        # create sidebar frame with widgets
        self.sidebar_frame = customtkinter.CTkFrame(self, width=140, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=6, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(6, weight=1)
        self.logo_label = customtkinter.CTkLabel(self.sidebar_frame, text="Menu",
                                                 font=customtkinter.CTkFont(size=20, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))
        self.sidebar_button_1 = customtkinter.CTkButton(self.sidebar_frame, text="BIS AUTOMATION",
                                                        command=self.Automate)
        self.sidebar_button_1.grid(row=2, column=0, padx=20, pady=10)
        self.sidebar_button_2 = customtkinter.CTkButton(self.sidebar_frame, text="HOD", command=self.HODTry)
        self.sidebar_button_2.grid(row=3, column=0, padx=20, pady=10)
        self.sidebar_button_3 = customtkinter.CTkButton(self.sidebar_frame, text="LOGIN", command=self.Mainpage)
        self.sidebar_button_3.grid(row=1, column=0, padx=20, pady=10)
        self.sidebar_button_4 = customtkinter.CTkButton(self.sidebar_frame, text="INITIATE AP", command=self.GUIAP)
        self.sidebar_button_4.grid(row=4, column=0, padx=20, pady=10)
        #self.sidebar_button_5 = customtkinter.CTkButton(self.sidebar_frame, text="BIS ACCEPTANCE", command=self.GUI_BIS_Acceptance)
        #self.sidebar_button_5.grid(row=5, column=0, padx=20, pady=10)

        self.appearance_mode_label = customtkinter.CTkLabel(self.sidebar_frame, text="Appearance Mode:", anchor="w")
        self.appearance_mode_label.grid(row=7, column=0, padx=20, pady=(10, 0))
        self.appearance_mode_optionmenu = customtkinter.CTkOptionMenu(self.sidebar_frame,
                                                                      values=["System", "Dark", "Light"],
                                                                      command=self.change_appearance_mode_event)
        self.appearance_mode_optionmenu.grid(row=8, column=0, padx=20, pady=(10, 10))
        self.scaling_label = customtkinter.CTkLabel(self.sidebar_frame, text="UI Scaling:", anchor="w")
        self.scaling_label.grid(row=9, column=0, padx=20, pady=(10, 0))
        self.scaling_optionmenu = customtkinter.CTkOptionMenu(self.sidebar_frame,
                                                              values=["80%", "90%", "100%", "110%", "120%"],
                                                              command=self.change_scaling_event)
        self.scaling_optionmenu.grid(row=10, column=0, padx=20, pady=(10, 20))

        # set default values
        self.content_frame = customtkinter.CTkFrame(self, width=960, corner_radius=0)
        self.content_frame.grid(row=0, column=1, rowspan=4, sticky="nsew")
        self.content_frame.grid_rowconfigure(0, weight=1)
        self.content_frame.grid_columnconfigure(0, weight=1)
        self.pageframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                border_color="Gray", border_width=20, corner_radius=20)
        self.pageframe.place(x=40, y=40)
        self.title_label = customtkinter.CTkLabel(self.content_frame, text="Welcome to Automation",
                                                  font=customtkinter.CTkFont(size=50, weight="bold"))
        self.title_label.place(x=240, y=180)
        self.enter_button = customtkinter.CTkButton(master=self.content_frame, width=220, height=120, border_width=0,
                                                    corner_radius=8, text="LOGIN", command=self.Mainpage)
        self.enter_button.place(x=390, y=350)

    def change_appearance_mode_event(self, new_appearance_mode: str):
        customtkinter.set_appearance_mode(new_appearance_mode)

    def change_scaling_event(self, new_scaling: str):
        new_scaling_float = int(new_scaling.replace("%", "")) / 100
        customtkinter.set_widget_scaling(new_scaling_float)

    def GUI_BIS_Acceptance(self):

        # set default values
        self.mainframeBISAccpt = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
        border_color="Gray", border_width=20, corner_radius=20)
        self.mainframeBISAccpt.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.mainframeBISAccpt, image=my_image).place(x=60, y=120)

        self.usernameBISAccpt = StringVar()
        self.passwordBISAccpt = StringVar()
        self.circuitid1BISAccpt = StringVar()
        self.APnameBISAccpt = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head1 = CTkLabel(self.mainframeBISAccpt, width=120, text="PLEASE ENTER THE DETAILS BELOW*",
        font=("Tahoma", 15, "bold"))
        head1.place(x=300, y=40)
        CTkLabel(self.mainframeBISAccpt, width=120, text="This system functions to insert all the required attachments in the Checklists",
                 font=("Tahoma", 12)).place(x=250, y=60)
        # ---------------------------------------Username-----------------------------------------------------------
        CTkLabel(self.mainframeBISAccpt, width=120, text="USERNAME*", fg_color="transparent",
        font=("Tahoma", 15, "bold")).place(x=545, y=160)
        user113 = customtkinter.CTkEntry(master=self.mainframeBISAccpt,
        placeholder_text="USERNAME",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.usernameBISAccpt)
        user113.place(x=550, y=190)
        # ---------------------------------------Password-----------------------------------------------------------
        self.show_password = tkinter.BooleanVar(value=False)
        CTkLabel(self.mainframeBISAccpt, width=120, text="PASSWORD*", fg_color="transparent",
        font=("Tahoma", 15, "bold")).place(x=545, y=240)
        self.password_entry = customtkinter.CTkEntry(master=self.mainframeBISAccpt,
        placeholder_text="PASSWORD",show="*",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.passwordBISAccpt)
        self.password_entry.place(x=550, y=270)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.mainframeBISAccpt, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=855, y=267)
        # ---------------------------------------Ap NAME-----------------------------------------------------------
        user12 = CTkLabel(self.mainframeBISAccpt, width=120, text="AP NAME*", fg_color="transparent",
        font=("Tahoma", 15, "bold"))
        user12.place(x=535, y=310)
        user11 = customtkinter.CTkEntry(master=self.mainframeBISAccpt,
        placeholder_text="AP NAME",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.APnameBISAccpt)
        user11.place(x=550, y=340)
        # --------------------------------------circuitid-------------------------------------------------------------
        CTkLabel(self.mainframeBISAccpt, width=120, text="CIRCUIT ID*", fg_color="transparent",
                 font=("Tahoma", 15, "bold")).place(x=545, y=380)
        circuitid13 = customtkinter.CTkEntry(master=self.mainframeBISAccpt,
        placeholder_text="CIRCUIT ID",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.circuitid1BISAccpt)
        circuitid13.place(x=550, y=410)
        # --------------------------------------BUTTON-------------------------------------------------------------
        CTkButton(self.mainframeBISAccpt, width=100, height=30, text='Submit',border_width=0,
        command=self.BIS_Acceptance).place(x=650, y=530)
   
    #function for BIS Acceptance in NCR
    def BIS_Acceptance(self):

        self.ID_BISAccpt = self.usernameBISAccpt.get()
        self.Password_BISAccpt = self.passwordBISAccpt.get()
        self.circuitID_BISAccpt = self.circuitid1BISAccpt.get()
        self.APNAME_BISAccpt = self.APnameBISAccpt.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        # get IP Address
        url = "http://10.204.96.65:7200/"

        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_BISAccpt).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_BISAccpt).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)

        #open Search BIS Process in NCR
        time.sleep(2)
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_BISAccpt,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()

        # This is actual coding start in the AP. Please uncomment this step when testing with real AP
        #reviewSOCchecklist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Review SOC Service Checklist']")))
        #a.move_to_element(reviewSOCchecklist).click().perform()
        

        #to check if there is SOC Service Acceptance or not. please comment below try/except/else code SOC/SOC2/C/D once using testing real AP
        try:
            SOC = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='SOC Service Acceptance']")))
            a.move_to_element(SOC).click().perform()
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No SOC Service Acceptance found in the AP")           
        else:
            SOC2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='SOC Acceptance']")))
            a.move_to_element(SOC2).click().perform() 

        time.sleep(5)
        try:
            C = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='SOC Service Acceptance Checklist']"))) #click on the SOC Service Acceptance Checklist
            a.move_to_element(C).click().perform()     
        except:
            messagebox.showerror("Error", "No SOC Service Acceptance Checklist found in the AP")     
        else:
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='SOC Acceptance Checklist']"))) #click on the SOC Service Acceptance Checklist
            a.move_to_element(D).click().perform()
        #---------------------------------------------------------------------upload attachments-----------------------------------------------------------------------------------------#
        #------------------------------------------------------------------------first box (Summary Page)--------------------------------------------------------------------------------#
        SummaryPage1_attachment = "C:\\Automate\\Checklist Output\\NCR\\Summary Page1.png"
        SummaryPage2_attachment = "C:\\Automate\\Checklist Output\\NCR\\Summary Page2.png"
        time.sleep(5)
        if os.path.exists(SummaryPage1_attachment) and os.path.exists(SummaryPage2_attachment):
            E = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[12]"))) #click on the first attachment box
            a.move_to_element(E).click().perform()
            F = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(F).click().perform()
            time.sleep(3)
            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(G).click().perform()
            time.sleep(5)
            pyautogui.typewrite(SummaryPage1_attachment)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)

            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File for 2nd summary page attachment
            a.move_to_element(G).click().perform()
            time.sleep(5)
            pyautogui.typewrite(SummaryPage2_attachment)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: Summary Page1.png & Summary Page2.png!")


        #------------------------------------------------------------------------2nd box (Network Diagram)--------------------------------------------------------------------------------#
        path_dir_ND = "C:\\Automate\\Checklist Output\\NCR"
        latest_NDfile = max(glob.glob(os.path.join(path_dir_ND, "*NWO*")) + glob.glob(os.path.join(path_dir_ND, "*MWO*")), key=os.path.getctime)
        time.sleep(5)
        if os.path.exists(latest_NDfile):
            H = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[12]"))) #click on the 2nd attachment box
            a.move_to_element(H).click().perform()
            I = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(I).click().perform()
            time.sleep(3)
            J = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(J).click().perform()
            time.sleep(5)
            pyautogui.typewrite(latest_NDfile)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: ND")


        #------------------------------------------------------------------------3rd box (Path Element)--------------------------------------------------------------------------------#
        path_element = "C:\\Automate\\Checklist Output\\NCR\\Path Element.png"
        AP_Description = "C:\\Automate\\Checklist Output\\NCR\\AP Description.png"
        Parameters_Tab = "C:\\Automate\\Checklist Output\\NCR\\Parameters Tab.png"
        time.sleep(5)
        if os.path.exists(exfo_file1) and os.path.exists(AP_Description) and os.path.exists(Parameters_Tab):
            K = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[12]"))) #click on the 3rd attachment box
            a.move_to_element(K).click().perform()
            L = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(L).click().perform()
            time.sleep(3)
            M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(M).click().perform()
            time.sleep(5)
            pyautogui.typewrite(path_element)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #___________________________________________________AP_Description__________________________________________________________________#
            M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(M).click().perform()
            time.sleep(5)
            pyautogui.typewrite(AP_Description)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #___________________________________________________Parameters Tab__________________________________________________________________#
            M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(M).click().perform()
            time.sleep(5)
            pyautogui.typewrite(Parameters_Tab)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)        
        else:
            print("File not found: Path Element.png, AP Description.png, Parameters Tab.png")

            
        #------------------------------------------------------------------------4th box (Alarms)--------------------------------------------------------------------------------#
        #___________________________________________________BREEZE AIR Folder__________________________________________________________________#
        radios_file = "C:\\Automate\\Checklist Output\\BREEZE AIR\radios.csv"
        if os.path.exists(radios_file):
            N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(N).click().perform()
            O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(O).click().perform()
            time.sleep(3)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(radios_file)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: radios.csv!")
        #___________________________________________________CNMaestro Folder__________________________________________________________________#
        cnmaestro_file1 = "C:\\Automate\\Checklist Output\\CNMAESTRO\\Performance1.png"
        cnmaestro_file2 = "C:\\Automate\\Checklist Output\\CNMAESTRO\\Performance2.png"
        if os.path.exists(cnmaestro_file1) and os.path.exists(cnmaestro_file2):
            N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(N).click().perform()
            O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(O).click().perform()
            time.sleep(3)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(cnmaestro_file1)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(cnmaestro_file2)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: Performance1.png & Performance2.png!")
        #___________________________________________________EXFO Folder__________________________________________________________________#
        exfo_file1 = "C:\\Automate\\Checklist Output\\EXFO\\monitors.png"
        exfo_file2 = "C:\\Automate\\Checklist Output\\EXFO\\monitors (1).png"

        if os.path.exists(exfo_file1) and os.path.exists(cnmaestro_file2):
            N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(N).click().perform()
            O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(O).click().perform()
            time.sleep(3)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(cnmaestro_file1)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(exfo_file2)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: monitors.png & monitors (1).png!")

        directory = r'C:\\Automate\\Checklist Output\\EXFO' # Define the directory where the files are located
        file_list = os.listdir(directory) # Get a list of all files in the directory
        prefix = "Warning Alert_" # Define a prefix to look for
        files_with_timestamps = [] # Create a list to store file paths and their creation times
    
        for file_name in file_list: # Iterate through the files in the directory
            if file_name.startswith(prefix): # Check if the file starts with the specified prefix
                full_file_path = os.path.join(directory, file_name) # Construct the full file path
                creation_time = os.path.getctime(full_file_path) # Get the creation time of the file   
                files_with_timestamps.append((full_file_path, creation_time)) # Add the file path and creation time to the list
       
        files_with_timestamps.sort(key=lambda x: x[1], reverse=True)  # Sort the list of files by creation time in descending order

        for file_path, _ in files_with_timestamps: # Iterate through the sorted files and process them
            if os.path.exists(file_path): # Check if the file exists before attempting to use it
                N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
                a.move_to_element(N).click().perform()
                O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
                a.move_to_element(O).click().perform()
                time.sleep(3)
                P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
                a.move_to_element(P).click().perform()
                time.sleep(5)
                pyautogui.typewrite(file_path)        # Type the file path
                pyautogui.press('enter')         # Press Enter
                time.sleep(10)
                WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
                time.sleep(10)
                # Perform your desired action with the file (e.g., upload it)
                print(f"Processing file: {file_path}")
            else:
                print(f"File not found: {file_path}")
        #___________________________________________________MERAKI Folder__________________________________________________________________#
        DRO_livedata = "C:\\Automate\\Checklist Output\\MERAKI\\DRO Device Live Data.png"
        DRO_configuration = "C:\\Automate\\Checklist Output\\MERAKI\\DRO Device Configuration.png"
        #check on DRO screenshot and then upload if exixts
        if os.path.exists(DRO_livedata) and os.path.exists(DRO_configuration):
            Q = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(Q).click().perform()
            R = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(R).click().perform()
            time.sleep(3)
            T = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(T).click().perform()
            time.sleep(5)
            pyautogui.typewrite(DRO_livedata)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            U = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(U).click().perform()
            time.sleep(5)
            pyautogui.typewrite(DRO_configuration)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: DRO Device Live Data.png &  DRO Device Configuration.png")

        #DCG files in meraki folder (DCG Device, DCG Device Configuration, DCG Device Live Data, DCG Device Latency & Loss)
        directoryDCG = r'C:\\Automate\\Checklist Output\\MERAKI' # Define the directory where the files are located (DCG files)
        file_listDCG = os.listdir(directoryDCG) # Get a list of all files in the directory
        prefix = "DCG " # Define a prefix to look for
        filesDCG_with_timestamps = [] # Create a list to store file paths and their creation times
    
        for file_name in file_listDCG: # Iterate through the files in the directory
            if file_name.startswith(prefix): # Check if the file starts with the specified prefix
                full_file_pathDCG = os.path.join(directoryDCG, file_name) # Construct the full file path
                creation_timeDCG = os.path.getctime(full_file_pathDCG) # Get the creation time of the file   
                filesDCG_with_timestamps.append((full_file_pathDCG, creation_timeDCG)) # Add the file path and creation time to the list
       
        filesDCG_with_timestamps.sort(key=lambda x: x[1], reverse=True)  # Sort the list of files by creation time in descending order

        for file_path, _ in filesDCG_with_timestamps: # Iterate through the sorted files and process them
            if os.path.exists(file_path): # Check if the file exists before attempting to use it
                N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
                a.move_to_element(N).click().perform()
                O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
                a.move_to_element(O).click().perform()
                time.sleep(3)
                P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
                a.move_to_element(P).click().perform()
                time.sleep(5)
                pyautogui.typewrite(file_path)        # Type the file path
                pyautogui.press('enter')         # Press Enter
                time.sleep(10)
                WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
                time.sleep(10)
                # Perform your desired action with the file (e.g., upload it)
                print(f"Processing file: {file_path}")
            else:
                print(f"File not found: {file_path}")
                
        merakiDSW = "C:\\Automate\\Checklist Output\\MERAKI\\DSW Device Connectivity.png"
        merakiDAP = "C:\\Automate\\Checklist Output\\MERAKI\\DAP Device Connectivity.png"

        if os.path.exists(merakiDSW) and os.path.exists(merakiDAP):
            N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(N).click().perform()
            O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(O).click().perform()
            time.sleep(3)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(merakiDSW)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(merakiDAP)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: monitors.png & monitors (1).png!")
        #___________________________________________________NCE-FAN Folder__________________________________________________________________#
        ncefan_file = "C:\\Automate\\Checklist Output\\NCE-FAN\\Node Visibility.png"
        if os.path.exists(ncefan_file):
            N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(N).click().perform()
            O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(O).click().perform()
            time.sleep(3)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(ncefan_file)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: Node Visibility.png!")
        #___________________________________________________NCE-IP Folder__________________________________________________________________#
        nceip_file = "C:\\Automate\\Checklist Output\\NCE-IP\\Node Visibility.png"
        if os.path.exists(nceip_file):
            N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
            a.move_to_element(N).click().perform()
            O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
            a.move_to_element(O).click().perform()
            time.sleep(3)
            P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
            a.move_to_element(P).click().perform()
            time.sleep(5)
            pyautogui.typewrite(nceip_file)        # Type the file path
            pyautogui.press('enter')         # Press Enter
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
        else:
            print("File not found: Node Visibility.png!")
        #___________________________________________________________NCM folder____________________________________________________________________#
        #DCG files in NCM folder
        directoryNCM = r'C:\\Automate\\Checklist Output\\NCM' # Define the directory where the files are located (DCG files)
        file_listNCM = os.listdir(directoryNCM) # Get a list of all files in the directory
        prefixNCM = "NCM_" # Define a prefix to look for
        filesNCM_with_timestamps = [] # Create a list to store file paths and their creation times
    
        for file_name in file_listNCM: # Iterate through the files in the directory
            if file_name.startswith(prefixNCM): # Check if the file starts with the specified prefix
                full_file_pathNCM = os.path.join(directoryNCM, file_name) # Construct the full file path
                creation_timeNCM = os.path.getctime(full_file_pathNCM) # Get the creation time of the file   
                filesNCM_with_timestamps.append((full_file_pathNCM, creation_timeNCM)) # Add the file path and creation time to the list
       
        filesNCM_with_timestamps.sort(key=lambda x: x[1], reverse=True)  # Sort the list of files by creation time in descending order

        for file_path, _ in filesNCM_with_timestamps: # Iterate through the sorted files and process them
            if os.path.exists(file_path): # Check if the file exists before attempting to use it
                N = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]"))) #click on the 4th attachment box
                a.move_to_element(N).click().perform()
                O = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[12]/div[1]/div[2]/img[1]")))#click on the upload img
                a.move_to_element(O).click().perform()
                time.sleep(3)
                P = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))#click Add or Drop File
                a.move_to_element(P).click().perform()
                time.sleep(5)
                pyautogui.typewrite(file_path)        # Type the file path
                pyautogui.press('enter')         # Press Enter
                time.sleep(10)
                WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
                time.sleep(10)
                # Perform your desired action with the file (e.g., upload it)
                print(f"Processing file: {file_path}")
            else:
                print(f"File not found: {file_path}")


    #___________________________________________________________________INITIATE AP________________________________________________________________________________________________#
    def GUIAP(self):
        # set default values
        self.mainframeAP = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
        border_color="Gray", border_width=20, corner_radius=20)
        self.mainframeAP.place(x=40, y=40)

        self.APnameAP = StringVar()
        self.acceptanceobjectAP = StringVar()
        self.IPServiceCircuit = StringVar()
        self.attachment_Folder_Path = StringVar()
        self.BIS_Checklist_Template_Folder_Path = StringVar()
        self.extra_file_Path = StringVar()
        self.usernameAP = StringVar()
        self.passwordAP = StringVar()
        self.circuitid1AP = StringVar()
        self.CustNameinput = StringVar()
        self.VendorNameinput = StringVar()
        # ---------------------------------------label-----------------------------------------------------------
        head1 = CTkLabel(self.mainframeAP, width=120, text="PLEASE ENTER THE DETAILS BELOW*",
        font=("Tahoma", 15, "bold"))
        head1.place(x=300, y=40)
        CTkLabel(self.mainframeAP, width=120, text="This system functions to insert all the required elements and upload attachments in the Acceptance Plan",
                 font=("Tahoma", 12)).place(x=200, y=60)
        # ---------------------------------------Username-----------------------------------------------------------
        CTkLabel(self.mainframeAP, width=120, text="USERNAME*", fg_color="transparent",
        font=("Tahoma", 15, "bold")).place(x=85, y=135)
        user113 = customtkinter.CTkEntry(master=self.mainframeAP,
        placeholder_text="USERNAME",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.usernameAP)
        user113.place(x=90, y=165)
        # ---------------------------------------Password-----------------------------------------------------------
        self.show_password = tkinter.BooleanVar(value=False)
        CTkLabel(self.mainframeAP, width=120, text="PASSWORD*", fg_color="transparent",
        font=("Tahoma", 15, "bold")).place(x=85, y=215)
        self.password_entry = customtkinter.CTkEntry(master=self.mainframeAP,
        placeholder_text="PASSWORD",show="*",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.passwordAP)
        self.password_entry.place(x=90, y=245)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.mainframeAP, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=395, y=242)
        # ---------------------------------------Ap NAME-----------------------------------------------------------
        user12 = CTkLabel(self.mainframeAP, width=120, text="AP NAME*", fg_color="transparent",
        font=("Tahoma", 15, "bold"))
        user12.place(x=75, y=285)
        user11 = customtkinter.CTkEntry(master=self.mainframeAP,
        placeholder_text="AP NAME",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.APnameAP)
        user11.place(x=90, y=315)
        # --------------------------------------circuitid-------------------------------------------------------------
        CTkLabel(self.mainframeAP, width=120, text="CIRCUIT ID", fg_color="transparent",
                 font=("Tahoma", 15, "bold")).place(x=465, y=135)
        circuitid13 = customtkinter.CTkEntry(master=self.mainframeAP,
        placeholder_text="CIRCUIT ID",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.circuitid1AP)
        circuitid13.place(x=475, y=165)
        # --------------------------------------customer name-------------------------------------------------------------
        CTkLabel(self.mainframeAP, width=120, text="CUSTOMER NAME*", fg_color="transparent",
                 font=("Tahoma", 15, "bold")).place(x=480, y=215)
        circuitid13 = customtkinter.CTkEntry(master=self.mainframeAP,
        placeholder_text="CUSTOMER NAME",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.CustNameinput)
        circuitid13.place(x=475, y=245)
        #------------------------------------------dialog file path--------------------------------------------------#
        def open_file_dialog_attachment(): 
            attachment_file_path = filedialog.askopenfilename()
            # Process the selected file path         
            if attachment_file_path:
                attachment_file_path = attachment_file_path.replace('/', '\\') 
                attachment_entry.delete(0, 'end')  
                attachment_entry.insert(0, attachment_file_path)  
                     
        # Create a button to trigger the file dialog
        CTkButton(self.mainframeAP, width=50, height=20, text="Open", command=open_file_dialog_attachment).place(x=780, y=316)
        # --------------------------------------Upload Folder Path-------------------------------------------------------------
        CTkLabel(self.mainframeAP, width=120, text="UPLOAD FILE*", fg_color="transparent",
                 font=("Tahoma", 15, "bold")).place(x=475, y=285)
        CTkLabel(self.mainframeAP, width=120, text="(Please zip all required files)", fg_color="transparent",
                 font=("Tahoma", 12)).place(x=600, y=285)
        attachment_entry = customtkinter.CTkEntry(master=self.mainframeAP,
        placeholder_text="Attachment Folder Path",
        width=300,
        height=25,
        border_width=2,
        corner_radius=10, textvariable=self.attachment_Folder_Path)
        attachment_entry.place(x=475, y=315)  
        # --------------------------------------Acceptance Object-------------------------------------------------------------
        CTkLabel(self.mainframeAP, width=120, text="ACCEPTANCE OBJECT : ", 
                 fg_color="transparent",font=("Tahoma", 15, "bold")).place(x=480, y=355)
        CTkLabel(self.mainframeAP,text="Integration Acceptance Phase ",
                 fg_color="transparent",font=("Tahoma", 15, "italic")).place(x=660, y=355)
        CTkLabel(self.mainframeAP,text="eg. WOZROAIMS103 or AWCPEPRKM101 or ACDROSH50001",
                 fg_color="transparent",font=("Tahoma", 10)).place(x=480, y=380)

        #CTkLabel(self.mainframeAP,text="1.", fg_color="transparent",font=("Tahoma", 13)).place(x=470, y=187)
        customtkinter.CTkEntry(master=self.mainframeAP,placeholder_text="ACCEPTANCE OBJECT",width=300,  #first entry box
        height=25,border_width=2,corner_radius=10, textvariable=self.acceptanceobjectAP).place(x=475, y=405)
        #_______________________________DROPDOWN MENU_____________________________________________________#
        CTkLabel(self.mainframeAP,text="Please choose AP Template from dropdown below", fg_color="transparent",font=("Tahoma", 12)).place(x=90, y=355)
        optionAPTemplate = ["Template 1", "Template 2", "Template 3", "Template 4", "Template 5", "Template 6"]
        self.optionMenuTemplate = CTkOptionMenu(self.mainframeAP, values=optionAPTemplate)
        self.optionMenuTemplate.place(x=90, y=385)      
        # --------------------------------------Vendor Name-------------------------------------------------------------
        CTkLabel(self.mainframeAP, width=120, text="Please choose Vendor from dropdown below", 
                 fg_color="transparent",font=("Tahoma", 12)).place(x=90, y=425)
        optionVendorName = ["Vendor 1","Vendor 2","Vendor 3", "Vendor 4", "Vendor 5"]
        self.optionMenuVendor = CTkOptionMenu(self.mainframeAP, values=optionVendorName)
        self.optionMenuVendor.place(x=90, y=455)
        # --------------------------------------BUTTON-------------------------------------------------------------
        CTkButton(self.mainframeAP, width=100, height=30, text='Submit',border_width=0,
        command=self.submit).place(x=400, y=530)

    def submit(self):
        # Check the current value of the OptionMenu
        currentSelection = self.optionMenuTemplate.get()

        # Call the appropriate function based on the current selection
        if currentSelection == "CPE":
            self.APCPE()
        elif currentSelection == "Managed Services":
            self.APMANAGEDSERVICE()
        elif currentSelection == "Service Acceptance":
            self.APService()
        elif currentSelection == "CPE+Service":
            self.APCPEService()
        elif currentSelection == "SDWAN+Service":
            self.APSDWANService()
        elif currentSelection == "Managed Services+Service":
            self.APMSandSERVICE()

    def load_eye_images(self):
        self.eye_open_image = CTkImage(dark_image=Image.open("images\\show.png"))  # Replace with the path to your custom eye open image
        self.eye_closed_image = CTkImage(dark_image=Image.open("images\\hide.png")) 
        #self.eye_open_image = self.eye_open_image.subsample(3)
        #self.eye_closed_image = self.eye_closed_image.subsample(3)

    def toggle_password(self):
        self.show_password.set(not self.show_password.get())
        if self.show_password.get():
            self.password_entry.configure(show="")
            self.password_toggle_button.configure(image=self.eye_open_image)
        else:
            self.password_entry.configure(show="*")
            self.password_toggle_button.configure(image=self.eye_closed_image)

    def APMANAGEDSERVICE(self):
        self.ID_InAP = self.usernameAP.get()
        self.Password_InAP = self.passwordAP.get()
        self.AO_InAP = self.acceptanceobjectAP.get()
        self.circuitID = self.circuitid1AP.get()
        self.APNAME_InAP = self.APnameAP.get()
        self.attachmentfile = self.attachment_Folder_Path.get()
        self.CustName = self.CustNameinput.get()
        self.VendorName = self.optionMenuVendor.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        url = "http://"
        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_InAP).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_InAP).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)
        #------------------------------------------Initiate AP process---------------------------------------#
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_InAP,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()
        #------------------------------------------AP Description----------------------------------------------------------------------#
        Edit_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='pcEdit']")))
        a.move_to_element(Edit_but).click().perform()

        Desc_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='id_common_descr']")))
        current_text = Desc_but.get_attribute("value")
        Desc_but.clear()
        additional_text = f"\n{self.CustName}\n{self.circuitID}"
        updated_text = current_text + additional_text
        a.move_to_element(Desc_but)
        time.sleep(3)
        a.click().send_keys(updated_text).perform()

        # Locate the dropdown element
        try:
            dropdown_element = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[2]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/select[1]")))
        except:
            print("No element found")
        else:
            # Initialize the Select class
            dropdown = Select(dropdown_element)
            # Get the available options
            available_options = [option.text for option in dropdown.options]
            # Check if "Yes" is in the available options
            if "Yes" in available_options:
                # Select the "Yes" option
                dropdown.select_by_visible_text("Yes")

        Vendorbut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[6]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        a.move_to_element(Vendorbut).click().send_keys(self.VendorName).perform()
        Vendorchoice = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                                                 "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]")))
        a.move_to_element(Vendorchoice).click().perform()
        update_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='theform_update']")))
        a.move_to_element(update_but).click().perform()
        time.sleep(3)
        #_________________________________________________________Integration__________________________________________________________________#
        #checkif there is any Integration Acceptance Phase or not
        try:
            FOP_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(FOP_AObox).perform()
        except:
            print("No element found")
        else:
            A11 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(A11).click().perform()
            FOP_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FOP_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D1).click().perform()
        #-----------------------------------------------------check if there is any input on 2nd box AO-------------------------------------------------------#
        try:
            TSTM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(TSTM_AObox).perform()
        except:
            print("No element found")
        else:
            A22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(A22).perform()
            TSTM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(TSTM_AOlist).click(TSTM_AOlist).send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 3rd box AO-------------------------------------------------------#
        try:
            CMB_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(CMB_AObox).perform()
        except:
            print("No element found")
        else:
            A33 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(A33).perform()
            CMB_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(CMB_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()
        #-----------------------------------------------------check if there is any input on 4th box AO-------------------------------------------------------#
        try:
            SPM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(SPM_AObox).perform()
        except:
            print("No element found")
        else:
            A44 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(A44).perform()
            SPM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(SPM_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 5th box AO-------------------------------------------------------#
        try:
            FM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(FM_AObox).perform()
        except:
            print("No element found")
        else:
            A55 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(A55).perform()
            FM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FM_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #click on Save button
        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
        time.sleep(3)
        #-----------------------------------------upload attachments on Integration Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        z = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[17]")))
        a.move_to_element(z).click().perform()

        try:
            E = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E).click().perform()
        except:
            print("No element found")
        else:
            IA1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(IA1).click().perform()
            F = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F).click().perform()
            time.sleep(3)
            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #click on Show Items per Page
            K = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE).click().perform()
        except:
            print("No element found")
        else:
            IA2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(IA2).click().perform()
            FF = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF).click().perform()
            time.sleep(3)
            GG = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE1).click().perform()
        except:
            print("No element found")
        else:
            IA3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(IA3).click().perform()
            FF1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF1).click().perform()
            time.sleep(3)
            GG1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG1).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK1).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE2).click().perform()
        except:
            print("No element found")
        else:
            IA4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(IA4).click().perform()
            FF2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF2).click().perform()
            time.sleep(3)
            GG2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG2).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK2).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            EE3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(EE3).click().perform()
        except:
            print("No element found")
        else:
            IA5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(IA5).click().perform()
            FF3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF3).click().perform()
            time.sleep(3)
            GG3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG3).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)   
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK3).click().perform()
            time.sleep(3)

        #click on Save button
        L = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L).click().perform()

        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
 
        L1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L1).click().perform()
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your AP in NCR system')

    def APMSandSERVICE(self):
        self.ID_InAP = self.usernameAP.get()
        self.Password_InAP = self.passwordAP.get()
        self.AO_InAP = self.acceptanceobjectAP.get()
        self.circuitID = self.circuitid1AP.get()
        self.APNAME_InAP = self.APnameAP.get()
        self.attachmentfile = self.attachment_Folder_Path.get()
        self.CustName = self.CustNameinput.get()
        self.VendorName = self.optionMenuVendor.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        url = "http://"
        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_InAP).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_InAP).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)
       #--------------------------------------search cct ID in NCR--------------------------------------------#
        servicecircuittab = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(servicecircuittab).perform()
        servicecircuittab2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search Service Circuit")))
        a.move_to_element(servicecircuittab2).click().perform()
        time.sleep(2)

        o = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/div[1]/table[1]/tbody[1]/tr[2]/td[1]/table[1]/tbody[1]/tr[2]/td[1]/div[1]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/input[1]")))
        a.move_to_element(o).click().send_keys(self.circuitID).send_keys(Keys.ENTER).perform()
        #q = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='pagerPageShow']")))
        #a.move_to_element(q).perform()
        filtericon = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                        "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/thead/tr/th[3]/table/tbody/tr/td[2]")))        # filter to In Service only
        a.move_to_element(filtericon).click(filtericon).perform()
        s = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//input[@value='9124352540013835640']")))  # checkbox for In Service
        a.move_to_element(s).click().perform()
        t = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")))  # button apply
        a.move_to_element(t).click().perform()
        time.sleep(5)

        #stored IP Service circuit value in a variable
        try:
            u = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(u).perform()
            IPvalue = u.text.strip()
            print(IPvalue)
        except:
            print("No circuit ID inserted!")

        #IPServiceCircuit_value = IPvalue
        #------------------------------------------Initiate AP process---------------------------------------#
        time.sleep(2)
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_InAP,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()
        #------------------------------------------AP Description----------------------------------------------------------------------#
        Edit_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='pcEdit']")))
        a.move_to_element(Edit_but).click().perform()

        Desc_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='id_common_descr']")))
        current_text = Desc_but.get_attribute("value")
        Desc_but.clear()
        additional_text = f"\n{self.CustName}\n{self.circuitID}"
        updated_text = current_text + additional_text
        a.move_to_element(Desc_but)
        time.sleep(3)
        a.click().send_keys(updated_text).perform()

        try:
            Acceptance_Detail = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[3]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        except:
            print("No element found")
        else:        
            a.move_to_element(Acceptance_Detail).click().send_keys(IPvalue).perform()
            time.sleep(5)
            linkIP = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]/div[2]")))
            a.move_to_element(linkIP).click().click().perform()

        # Locate the dropdown element
        try:
            dropdown_element = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[2]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/select[1]")))
        except:
            print("No element found")
        else:
            # Initialize the Select class
            dropdown = Select(dropdown_element)
            # Get the available options
            available_options = [option.text for option in dropdown.options]
            # Check if "Yes" is in the available options
            if "Yes" in available_options:
                # Select the "Yes" option
                dropdown.select_by_visible_text("Yes")

        Vendorbut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[6]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        a.move_to_element(Vendorbut).click().send_keys(self.VendorName).perform()
        Vendorchoice = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                                                 "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]")))
        a.move_to_element(Vendorchoice).click().perform()
        update_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='theform_update']")))
        a.move_to_element(update_but).click().perform()
        time.sleep(3)
        #--------------------------------------------------Equipment Assignment-----------------------------------------------------------#
        equip = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[5]/a[1]")))
        a.move_to_element(equip).click(equip).perform()
        clickontick = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(clickontick).click().perform()
        Savebut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(Savebut).click().perform()
        #_________________________________________________________Integration__________________________________________________________________#
        #checkif there is any Integration Acceptance Phase or not
        try:
            FOP_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(FOP_AObox).perform()
        except:
            print("No element found")
        else:
            A11 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(A11).perform()
            FOP_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FOP_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D1).click().perform()
        #-----------------------------------------------------check if there is any input on 2nd box AO-------------------------------------------------------#
        try:
            TSTM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(TSTM_AObox).perform()
        except:
            print("No element found")
        else:
            A22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(A22).perform()
            TSTM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(TSTM_AOlist).click(TSTM_AOlist).send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 3rd box AO-------------------------------------------------------#
        try:
            CMB_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(CMB_AObox).perform()
        except:
            print("No element found")
        else:
            A33 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(A33).perform()
            CMB_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(CMB_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()
        #-----------------------------------------------------check if there is any input on 4th box AO-------------------------------------------------------#
        try:
            SPM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(SPM_AObox).perform()
        except:
            print("No element found")
        else:
            A44 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(A44).perform()
            SPM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(SPM_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 5th box AO-------------------------------------------------------#
        try:
            FM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(FM_AObox).perform()
        except:
            print("No element found")
        else:
            A55 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(A55).perform()
            FM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FM_AOlist).click().send_keys(self.AO_InAP).perform()
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #click on Save button
        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
        time.sleep(3)
        #-----------------------------------------upload attachments on Integration Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        z = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[17]")))
        a.move_to_element(z).click().perform()

        try:
            E = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E).click().perform()
        except:
            print("No element found")
        else:
            IA1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(IA1).click().perform()
            F = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F).click().perform()
            time.sleep(3)
            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #click on Show Items per Page
            K = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE).click().perform()
        except:
            print("No element found")
        else:
            IA2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(IA2).click().perform()
            FF = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF).click().perform()
            time.sleep(3)
            GG = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE1).click().perform()
        except:
            print("No element found")
        else:
            IA3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(IA3).click().perform()
            FF1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF1).click().perform()
            time.sleep(3)
            GG1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG1).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK1).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE2).click().perform()
        except:
            print("No element found")
        else:
            IA4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(IA4).click().perform()
            FF2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF2).click().perform()
            time.sleep(3)
            GG2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG2).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK2).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            EE3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(EE3).click().perform()
        except:
            print("No element found")
        else:
            IA5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(IA5).click().perform()
            FF3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF3).click().perform()
            time.sleep(3)
            GG3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG3).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)   
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK3).click().perform()
            time.sleep(3)

        #click on Save button
        L = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L).click().perform()

        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
 
        L1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L1).click().perform()
        #-----------------------------------------upload attachments on Service Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        zx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[16]")))
        a.move_to_element(zx).click().perform()
        try:
            E7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E7).click().perform()
        except:
            print("No element found")
        else:
            A1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(A1).click().perform()
            F7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F7).click().perform()
            time.sleep(3)
            G7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G7).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile, interval=0.01)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            K7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K7).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE7).click().perform()
        except:
            print("No element found")
        else:
            A2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(A2).click().perform()
            FF7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF7).click().perform()
            time.sleep(3)
            GG7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG7).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK7).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE12).click().perform()
        except:
            print("No element found")
        else:
            A3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(A3).click().perform()
            FF12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF12).click().perform()
            time.sleep(3)
            GG12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG12).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK12).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE22).click().perform()
        except:
            print("No element found")
        else:
            A4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(A4).click().perform()
            FF22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF22).click().perform()
            time.sleep(3)
            GG22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG22).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK22).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            box5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(box5).click().perform()
        except:
            print("No element found")
        else:
            A6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(A6).click().perform()
            box51 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box51).click().perform()
            time.sleep(3)
            box52 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box52).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(5)
            box53 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
            a.move_to_element(box53).click(box53).perform()
            time.sleep(3)
        #---------------------------------------------sixth box attachment------------------------------------------------------------#
        try:
            time.sleep(3)
            box6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(box6).click(box6).perform()       
        except:
            print("No element found")
        else:
            A7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(A7).click(A7).perform()    
            box61 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box61).click(box61).perform()
            time.sleep(3)
            box62 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box62).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            box63 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[16]")))
            a.move_to_element(box63).click(box63).perform()
            time.sleep(3)
        clickonsave = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(clickonsave).click().perform()
        time.sleep(2)
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your AP in NCR system')
        
    def APCPE(self):
        self.ID_InAP = self.usernameAP.get()
        self.Password_InAP = self.passwordAP.get()
        self.AO_InAP = self.acceptanceobjectAP.get()
        self.circuitID = self.circuitid1AP.get()
        self.APNAME_InAP = self.APnameAP.get()
        self.attachmentfile = self.attachment_Folder_Path.get()
        self.CustName = self.CustNameinput.get()
        self.VendorName = self.optionMenuVendor.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        url = "http://"
        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_InAP).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_InAP).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)
       #--------------------------------------search cct ID in NCR--------------------------------------------#
        servicecircuittab = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(servicecircuittab).perform()
        servicecircuittab2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search Service Circuit")))
        a.move_to_element(servicecircuittab2).click().perform()
        time.sleep(2)

        o = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/div[1]/table[1]/tbody[1]/tr[2]/td[1]/table[1]/tbody[1]/tr[2]/td[1]/div[1]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/input[1]")))
        a.move_to_element(o).click().send_keys(self.circuitID).send_keys(Keys.ENTER).perform()
        #q = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[2]")))
        #a.move_to_element(q).perform()
        filtericon = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[3]/table[1]/tbody[1]/tr[1]/td[2]/div[1]/a[1]/img[1]")))        # filter to In Service only
        a.move_to_element(filtericon).click(filtericon).perform()
        s = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//input[@value='9124352540013835640']")))  # checkbox for In Service
        a.move_to_element(s).click().perform()
        t = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")))  # button apply
        a.move_to_element(t).click().perform()
        time.sleep(5)

        #stored IP Service circuit value in a variable
        try:
            u = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(u).perform()
            IPvalue = u.text.strip()
            print(IPvalue)
        except:
            print("No circuit ID inserted!")

        #IPServiceCircuit_value = IPvalue
        #------------------------------------------Initiate AP process---------------------------------------#
        time.sleep(2)
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_InAP,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()
        #------------------------------------------AP Description----------------------------------------------------------------------#
        Edit_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='pcEdit']")))
        a.move_to_element(Edit_but).click().perform()

        Desc_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='id_common_descr']")))
        current_text = Desc_but.get_attribute("value")
        Desc_but.clear()
        additional_text = f"\n{self.CustName}\n{self.circuitID}"
        updated_text = current_text + additional_text
        a.move_to_element(Desc_but)
        time.sleep(3)
        a.click().send_keys(updated_text).perform()

        try:
            Acceptance_Detail = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[3]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        except:
            print("No element found")
        else:        
            a.move_to_element(Acceptance_Detail).click().send_keys(IPvalue).perform()
            time.sleep(5)
            linkIP = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]/div[2]")))
            a.move_to_element(linkIP).click().click().perform()

        # Locate the dropdown element
        try:
            dropdown_element = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[2]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/select[1]")))
        except:
            print("No element found")
        else:
            # Initialize the Select class
            dropdown = Select(dropdown_element)
            # Get the available options
            available_options = [option.text for option in dropdown.options]
            # Check if "Yes" is in the available options
            if "Yes" in available_options:
                # Select the "Yes" option
                dropdown.select_by_visible_text("Yes")

        Vendorbut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[6]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        a.move_to_element(Vendorbut).click().send_keys(self.VendorName).perform()
        Vendorchoice = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                                                 "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]")))
        a.move_to_element(Vendorchoice).click().perform()
        update_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='theform_update']")))
        a.move_to_element(update_but).click().perform()
        time.sleep(3)
        #--------------------------------------------------Equipment Assignment-----------------------------------------------------------#
        equip = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[5]/a[1]")))
        a.move_to_element(equip).click(equip).perform()
        clickontick = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(clickontick).click().perform()
        Savebut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(Savebut).click().perform()
        #-----------------------------------------Integration Acceptance Phase----------------------------------------------------------------------#
        #check to change AO Type first
        try:
            BOXAO1=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[4]")))
            a.move_to_element(BOXAO1).perform()
        except:
            print("No element found")
        else:
            BOXAO1=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[4]")))
            a.move_to_element(BOXAO1).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 2nd
        try:
            BOXAO2=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[4]")))
            a.move_to_element(BOXAO2).perform()
        except:
            print("No element found")
        else:
            BOXAO2=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[4]")))
            a.move_to_element(BOXAO2).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 3rd
        try:
            BOXAO3=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[4]")))
            a.move_to_element(BOXAO3).perform()
        except:
            print("No element found")
        else:
            BOXAO3=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[4]")))
            a.move_to_element(BOXAO3).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 4th
        try:
            BOXAO4=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[4]")))
            a.move_to_element(BOXAO4).perform()
        except:
            print("No element found")
        else:
            BOXAO4=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[4]")))
            a.move_to_element(BOXAO4).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 5th
        try:
            BOXAO5=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[4]")))
            a.move_to_element(BOXAO5).perform()
        except:
            print("No element found")
        else:
            BOXAO5=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[4]")))
            a.move_to_element(BOXAO5).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()

        #click on Save button
        Save_int = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(Save_int).click().perform()
        Save_int2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(Save_int2).click().perform()
        time.sleep(3)
        #_________________________________________________________Integration__________________________________________________________________#
        #checkif there is any Integration Acceptance Phase or not
        try:
            FOP_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(FOP_AObox).perform()
        except:
            print("No element found")
        else:
            A11 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(A11).perform()
            FOP_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FOP_AOlist).click().perform()
            textarea = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D1).click().perform()
        #-----------------------------------------------------check if there is any input on 2nd box AO-------------------------------------------------------#
        try:
            TSTM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(TSTM_AObox).perform()
        except:
            print("No element found")
        else:
            A22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(A22).perform()
            TSTM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(TSTM_AOlist).click(TSTM_AOlist).perform()
            textarea2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 3rd box AO-------------------------------------------------------#
        try:
            CMB_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(CMB_AObox).perform()
        except:
            print("No element found")
        else:
            A33 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(A33).perform()
            CMB_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(CMB_AOlist).click().perform()
            textarea3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea3).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 4th box AO-------------------------------------------------------#
        try:
            SPM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(SPM_AObox).perform()
        except:
            print("No element found")
        else:
            A44 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(A44).perform()
            SPM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(SPM_AOlist).click().perform()
            textarea4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea4).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 5th box AO-------------------------------------------------------#
        try:
            FM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(FM_AObox).perform()
        except:
            print("No element found")
        else:
            A55 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(A55).perform()
            FM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FM_AOlist).click().perform()
            textarea5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea5).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #click on Save button
        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
        time.sleep(3)
        #-----------------------------------------upload attachments on Integration Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        z = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[17]")))
        a.move_to_element(z).click().perform()

        try:
            E = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E).click().perform()
        except:
            print("No element found")
        else:
            IA1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(IA1).click().perform()
            F = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F).click().perform()
            time.sleep(3)
            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #click on Show Items per Page
            K = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE).click().perform()
        except:
            print("No element found")
        else:
            IA2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(IA2).click().perform()
            FF = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF).click().perform()
            time.sleep(3)
            GG = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE1).click().perform()
        except:
            print("No element found")
        else:
            IA3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(IA3).click().perform()
            FF1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF1).click().perform()
            time.sleep(3)
            GG1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG1).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK1).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE2).click().perform()
        except:
            print("No element found")
        else:
            IA4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(IA4).click().perform()
            FF2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF2).click().perform()
            time.sleep(3)
            GG2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG2).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK2).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            EE3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(EE3).click().perform()
        except:
            print("No element found")
        else:
            IA5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(IA5).click().perform()
            FF3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF3).click().perform()
            time.sleep(3)
            GG3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG3).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)   
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK3).click().perform()
            time.sleep(3)

        #click on Save button
        L = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L).click().perform()

        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
 
        L1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L1).click().perform()
        #-----------------------------------------upload attachments on Service Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        zx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[16]")))
        a.move_to_element(zx).click().perform()
        try:
            E7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E7).click().perform()
        except:
            print("No element found")
        else:
            A1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(A1).click().perform()
            F7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F7).click().perform()
            time.sleep(3)
            G7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G7).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile, interval=0.01)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            K7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K7).click().perform()
            time.sleep(3)

        clickonsave = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(clickonsave).click().perform()
        time.sleep(2)
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your AP in NCR system')        
    
    def APService(self):
        self.ID_InAP = self.usernameAP.get()
        self.Password_InAP = self.passwordAP.get()
        self.AO_InAP = self.acceptanceobjectAP.get()
        self.circuitID = self.circuitid1AP.get()
        self.APNAME_InAP = self.APnameAP.get()
        self.attachmentfile = self.attachment_Folder_Path.get()
        self.CustName = self.CustNameinput.get()
        self.VendorName = self.optionMenuVendor.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        # get IP Address
        #url = "http://"
        url = "http://"
        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_InAP).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_InAP).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)
       #--------------------------------------search cct ID in NCR--------------------------------------------#
        servicecircuittab = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(servicecircuittab).perform()
        servicecircuittab2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search Service Circuit")))
        a.move_to_element(servicecircuittab2).click().perform()
        time.sleep(2)

        o = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/div[1]/table[1]/tbody[1]/tr[2]/td[1]/table[1]/tbody[1]/tr[2]/td[1]/div[1]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/input[1]")))
        a.move_to_element(o).click().send_keys(self.circuitID).send_keys(Keys.ENTER).perform()
        #q = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='pagerPageShow']")))
        #a.move_to_element(q).perform()
        filtericon = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                        "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/thead/tr/th[3]/table/tbody/tr/td[2]")))        # filter to In Service only
        a.move_to_element(filtericon).click(filtericon).perform()
        s = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//input[@value='9124352540013835640']")))  # checkbox for In Service
        a.move_to_element(s).click().perform()
        t = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")))  # button apply
        a.move_to_element(t).click().perform()
        time.sleep(5)

        #stored IP Service circuit value in a variable
        try:
            u = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(u).perform()
            IPvalue = u.text.strip()
            print(IPvalue)
        except:
            print("No circuit ID inserted!")

        #IPServiceCircuit_value = IPvalue
        #------------------------------------------Initiate AP process---------------------------------------#
        time.sleep(2)
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_InAP,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()
        #------------------------------------------AP Description----------------------------------------------------------------------#
        Edit_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='pcEdit']")))
        a.move_to_element(Edit_but).click().perform()

        Desc_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='id_common_descr']")))
        current_text = Desc_but.get_attribute("value")
        Desc_but.clear()
        additional_text = f"\n{self.CustName}\n{self.circuitID}"
        updated_text = current_text + additional_text
        a.move_to_element(Desc_but)
        time.sleep(3)
        a.click().send_keys(updated_text).perform()

        try:
            Acceptance_Detail = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[3]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        except:
            print("No element found")
        else:        
            a.move_to_element(Acceptance_Detail).click().send_keys(IPvalue).perform()
            time.sleep(5)
            linkIP = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]/div[2]")))
            a.move_to_element(linkIP).click().click().perform()

        # Locate the dropdown element
        try:
            dropdown_element = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[2]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/select[1]")))
        except:
            print("No element found")
        else:
            # Initialize the Select class
            dropdown = Select(dropdown_element)
            # Get the available options
            available_options = [option.text for option in dropdown.options]
            # Check if "Yes" is in the available options
            if "Yes" in available_options:
                # Select the "Yes" option
                dropdown.select_by_visible_text("Yes")

        Vendorbut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[6]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        a.move_to_element(Vendorbut).click().send_keys(self.VendorName).perform()
        Vendorchoice = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                                                 "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]")))
        a.move_to_element(Vendorchoice).click().perform()
        update_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='theform_update']")))
        a.move_to_element(update_but).click().perform()
        time.sleep(3)
        #--------------------------------------------------Equipment Assignment-----------------------------------------------------------#
        equip = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[5]/a[1]")))
        a.move_to_element(equip).click(equip).perform()
        clickontick = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(clickontick).click().perform()
        clickonsave = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(clickonsave).click().perform()
        time.sleep(2)
        #-----------------------------------------upload attachments on Service Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        zx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[16]")))
        a.move_to_element(zx).click().perform()
        try:
            E7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E7).click().perform()
        except:
            print("No element found")
        else:
            A1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(A1).click().perform()
            F7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F7).click(F7).perform()
            time.sleep(3)
            G7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G7).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            K7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K7).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            EE7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE7).click().perform()
        except:
            print("No element found")
        else:
            A2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(A2).click().perform()
            FF7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF7).click().perform()
            time.sleep(3)
            GG7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG7).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            KK7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK7).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            EE12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE12).click().perform()
        except:
            print("No element found")
        else:
            A3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(A3).click().perform()
            FF12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF12).click().perform()
            time.sleep(3)
            GG12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG12).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            KK12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK12).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            EE22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE22).click().perform()
        except:
            print("No element found")
        else:
            A4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(A4).click().perform()
            FF22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF22).click().perform()
            time.sleep(3)
            GG22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG22).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            KK22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK22).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            box5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(box5).click().perform()
        except:
            print("No element found")
        else:
            A6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(A6).click().perform()
            box51 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box51).click().perform()
            time.sleep(3)
            box52 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box52).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(5)
            #click on Show Items per Page
            box53 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
            a.move_to_element(box53).click(box53).perform()
            time.sleep(3)
        #---------------------------------------------sixth box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            time.sleep(3)
            box6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(box6).click(box6).perform()       
        except:
            print("No element found")
        else:
            A7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(A7).click(A7).perform()    
            box61 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box61).click(box61).perform()
            time.sleep(3)
            box62 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box62).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            box63 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[16]")))
            a.move_to_element(box63).click(box63).perform()
            time.sleep(3)
        #click on Save button
        MM22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(MM22).click().perform()
        time.sleep(2)
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your AP in NCR system')

    def APCPEService(self):

        self.ID_InAP = self.usernameAP.get()
        self.Password_InAP = self.passwordAP.get()
        self.AO_InAP = self.acceptanceobjectAP.get()
        self.circuitID = self.circuitid1AP.get()
        self.APNAME_InAP = self.APnameAP.get()
        self.attachmentfile = self.attachment_Folder_Path.get()
        self.CustName = self.CustNameinput.get()
        self.VendorName = self.optionMenuVendor.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        url = "http://"
        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_InAP).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_InAP).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)
       #--------------------------------------search cct ID in NCR--------------------------------------------#
        servicecircuittab = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(servicecircuittab).perform()
        servicecircuittab2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search Service Circuit")))
        a.move_to_element(servicecircuittab2).click().perform()
        time.sleep(2)

        o = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/div[1]/table[1]/tbody[1]/tr[2]/td[1]/table[1]/tbody[1]/tr[2]/td[1]/div[1]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/input[1]")))
        a.move_to_element(o).click().send_keys(self.circuitID).send_keys(Keys.ENTER).perform()
        #q = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='pagerPageShow']")))
        #a.move_to_element(q).perform()
        filtericon = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                        "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/thead/tr/th[3]/table/tbody/tr/td[2]")))        # filter to In Service only
        a.move_to_element(filtericon).click(filtericon).perform()
        s = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//input[@value='9124352540013835640']")))  # checkbox for In Service
        a.move_to_element(s).click().perform()
        t = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")))  # button apply
        a.move_to_element(t).click().perform()
        time.sleep(5)

        #stored IP Service circuit value in a variable
        try:
            u = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(u).perform()
            IPvalue = u.text.strip()
            print(IPvalue)
        except:
            print("No circuit ID inserted!")

        #IPServiceCircuit_value = IPvalue
        #------------------------------------------Initiate AP process---------------------------------------#
        time.sleep(2)
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_InAP,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()
        #------------------------------------------AP Description----------------------------------------------------------------------#
        Edit_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='pcEdit']")))
        a.move_to_element(Edit_but).click().perform()

        Desc_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='id_common_descr']")))
        current_text = Desc_but.get_attribute("value")
        Desc_but.clear()
        additional_text = f"\n{self.CustName}\n{self.circuitID}"
        updated_text = current_text + additional_text
        a.move_to_element(Desc_but)
        time.sleep(3)
        a.click().send_keys(updated_text).perform()

        try:
            Acceptance_Detail = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[3]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        except:
            print("No element found")
        else:        
            a.move_to_element(Acceptance_Detail).click().send_keys(IPvalue).perform()
            time.sleep(5)
            linkIP = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]/div[2]")))
            a.move_to_element(linkIP).click().click().perform()

        # Locate the dropdown element
        try:
            dropdown_element = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[2]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/select[1]")))
        except:
            print("No element found")
        else:
            # Initialize the Select class
            dropdown = Select(dropdown_element)
            # Get the available options
            available_options = [option.text for option in dropdown.options]
            # Check if "Yes" is in the available options
            if "Yes" in available_options:
                # Select the "Yes" option
                dropdown.select_by_visible_text("Yes")

        Vendorbut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[6]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        a.move_to_element(Vendorbut).click().send_keys(self.VendorName).perform()
        Vendorchoice = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                                                 "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]")))
        a.move_to_element(Vendorchoice).click().perform()
        update_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='theform_update']")))
        a.move_to_element(update_but).click().perform()
        time.sleep(3)
        #--------------------------------------------------Equipment Assignment-----------------------------------------------------------#
        equip = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[5]/a[1]")))
        a.move_to_element(equip).click(equip).perform()
        clickontick = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(clickontick).click().perform()
        Savebut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(Savebut).click().perform()
        #-----------------------------------------Integration Acceptance Phase----------------------------------------------------------------------#
        #check to change AO Type first
        try:
            BOXAO1=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[4]")))
            a.move_to_element(BOXAO1).perform()
        except:
            print("No element found")
        else:
            BOXAO1=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[4]")))
            a.move_to_element(BOXAO1).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 2nd
        try:
            BOXAO2=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[4]")))
            a.move_to_element(BOXAO2).perform()
        except:
            print("No element found")
        else:
            BOXAO2=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[4]")))
            a.move_to_element(BOXAO2).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 3rd
        try:
            BOXAO3=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[4]")))
            a.move_to_element(BOXAO3).perform()
        except:
            print("No element found")
        else:
            BOXAO3=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[4]")))
            a.move_to_element(BOXAO3).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 4th
        try:
            BOXAO4=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[4]")))
            a.move_to_element(BOXAO4).perform()
        except:
            print("No element found")
        else:
            BOXAO4=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[4]")))
            a.move_to_element(BOXAO4).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[4]/div[1]/div[2]/img[1]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()
        #check to change AO Type 5th
        try:
            BOXAO5=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[4]")))
            a.move_to_element(BOXAO5).perform()
        except:
            print("No element found")
        else:
            BOXAO5=WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[4]")))
            a.move_to_element(BOXAO5).perform()
            routerbox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[4]/div[1]/div[2]")))
            a.move_to_element(routerbox).click().perform()
            RB2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(RB2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
            time.sleep(2)
            a.send_keys("Router").perform()
            time.sleep(3)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[2]")))
            a.move_to_element(RB3).click().perform()
            Dx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[6]")))
            a.move_to_element(Dx).click().perform()

        #click on Save button
        Save_int = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(Save_int).click().perform()
        Save_int2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(Save_int2).click().perform()
        time.sleep(3)
        #_________________________________________________________Integration__________________________________________________________________#
        #checkif there is any Integration Acceptance Phase or not
        try:
            FOP_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(FOP_AObox).perform()
        except:
            print("No element found")
        else:
            A11 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(A11).perform()
            FOP_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FOP_AOlist).click().perform()
            textarea = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D1).click().perform()
        #-----------------------------------------------------check if there is any input on 2nd box AO-------------------------------------------------------#
        try:
            TSTM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(TSTM_AObox).perform()
        except:
            print("No element found")
        else:
            A22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(A22).perform()
            TSTM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(TSTM_AOlist).click(TSTM_AOlist).perform()
            textarea2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 3rd box AO-------------------------------------------------------#
        try:
            CMB_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(CMB_AObox).perform()
        except:
            print("No element found")
        else:
            A33 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(A33).perform()
            CMB_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(CMB_AOlist).click().perform()
            textarea3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea3).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 4th box AO-------------------------------------------------------#
        try:
            SPM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(SPM_AObox).perform()
        except:
            print("No element found")
        else:
            A44 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(A44).perform()
            SPM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(SPM_AOlist).click().perform()
            textarea4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea4).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 5th box AO-------------------------------------------------------#
        try:
            FM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(FM_AObox).perform()
        except:
            print("No element found")
        else:
            A55 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(A55).perform()
            FM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FM_AOlist).click().perform()
            textarea5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea5).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #click on Save button
        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
        time.sleep(3)
        #-----------------------------------------upload attachments on Integration Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        z = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[17]")))
        a.move_to_element(z).click().perform()

        try:
            E = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E).click().perform()
        except:
            print("No element found")
        else:
            IA1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(IA1).click().perform()
            F = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F).click().perform()
            time.sleep(3)
            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #click on Show Items per Page
            K = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE).click().perform()
        except:
            print("No element found")
        else:
            IA2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(IA2).click().perform()
            FF = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF).click().perform()
            time.sleep(3)
            GG = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE1).click().perform()
        except:
            print("No element found")
        else:
            IA3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(IA3).click().perform()
            FF1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF1).click().perform()
            time.sleep(3)
            GG1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG1).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK1).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE2).click().perform()
        except:
            print("No element found")
        else:
            IA4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(IA4).click().perform()
            FF2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF2).click().perform()
            time.sleep(3)
            GG2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG2).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK2).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            EE3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(EE3).click().perform()
        except:
            print("No element found")
        else:
            IA5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(IA5).click().perform()
            FF3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF3).click().perform()
            time.sleep(3)
            GG3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG3).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)   
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK3).click().perform()
            time.sleep(3)

        #click on Save button
        L = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L).click().perform()

        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
 
        L1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L1).click().perform()
        #-----------------------------------------upload attachments on Service Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        zx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[16]")))
        a.move_to_element(zx).click().perform()
        try:
            E7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E7).click().perform()
        except:
            print("No element found")
        else:
            A1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(A1).click().perform()
            F7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F7).click().perform()
            time.sleep(3)
            G7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G7).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile, interval=0.01)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            K7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K7).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            EE7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE7).click().perform()
        except:
            print("No element found")
        else:
            EE7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE7).click().perform()
            FF7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF7).click().perform()
            time.sleep(3)
            GG7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG7).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            KK7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK7).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            EE12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE12).click().perform()
        except:
            print("No element found")
        else:
            EE12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE12).click().perform()
            FF12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF12).click().perform()
            time.sleep(3)
            GG12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG12).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            KK12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK12).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            EE22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE22).click().perform()
        except:
            print("No element found")
        else:
            EE22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE22).click().perform()
            FF22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF22).click().perform()
            time.sleep(3)
            GG22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG22).click().perform()
            time.sleep(3)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            KK22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK22).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            box5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(box5).click().perform()
        except:
            print("No element found")
        else:
            box5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(box5).click().perform()
            box51 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box51).click().perform()
            time.sleep(3)
            box52 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box52).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(5)
            #click on Show Items per Page
            box53 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
            a.move_to_element(box53).click(box53).perform()
            time.sleep(3)
        #---------------------------------------------sixth box attachment------------------------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        try:
            time.sleep(3)
            box6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(box6).click(box6).perform()       
        except:
            print("No element found")
        else:
            box6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(box6).click(box6).perform() 
            box61 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box61).click(box61).perform()
            time.sleep(3)
            box62 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box62).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            #click on Show Items per Page
            box63 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[16]")))
            a.move_to_element(box63).click(box63).perform()
            time.sleep(3)

        #click on Save button
        MM22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(MM22).click().perform()
        time.sleep(2)
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your AP in NCR system')

    def APSDWANService(self):  
        self.ID_InAP = self.usernameAP.get()
        self.Password_InAP = self.passwordAP.get()
        self.AO_InAP = self.acceptanceobjectAP.get()
        self.circuitID = self.circuitid1AP.get()
        self.APNAME_InAP = self.APnameAP.get()
        self.attachmentfile = self.attachment_Folder_Path.get()
        self.CustName = self.CustNameinput.get()
        self.VendorName = self.optionMenuVendor.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the Initiate AP program is finished running.')
        
        #bot = webdriver.Chrome(service=self.s, options=self.options)
        bot = webdriver.Chrome(service=self.s)
        bot.maximize_window() 
        url = "http://"
        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        a = ActionChains(bot)

        # LOGIN steps
        UserLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        a.move_to_element(UserLogIN).click().send_keys(self.ID_InAP).perform()
        PWLogIN = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        a.move_to_element(PWLogIN).click().send_keys(self.Password_InAP).perform()
        LogINButton = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "loginHiddenButton")))
        a.move_to_element(LogINButton).click().perform()
        time.sleep(5)
       #--------------------------------------search cct ID in NCR--------------------------------------------#
        servicecircuittab = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(servicecircuittab).perform()
        servicecircuittab2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search Service Circuit")))
        a.move_to_element(servicecircuittab2).click().perform()
        time.sleep(2)

        o = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[1]/div[1]/table[1]/tbody[1]/tr[2]/td[1]/table[1]/tbody[1]/tr[2]/td[1]/div[1]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[2]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/input[1]")))
        a.move_to_element(o).click().send_keys(self.circuitID).send_keys(Keys.ENTER).perform()
        #q = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='pagerPageShow']")))
        #a.move_to_element(q).perform()
        filtericon = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                        "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/thead/tr/th[3]/table/tbody/tr/td[2]")))        # filter to In Service only
        a.move_to_element(filtericon).click(filtericon).perform()
        s = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//input[@value='9124352540013835640']")))  # checkbox for In Service
        a.move_to_element(s).click().perform()
        t = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")))  # button apply
        a.move_to_element(t).click().perform()
        time.sleep(5)

        #stored IP Service circuit value in a variable
        try:
            u = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(u).perform()
            IPvalue = u.text.strip()
            print(IPvalue)
        except:
            print("No circuit ID inserted!")

        #IPServiceCircuit_value = IPvalue
        #------------------------------------------Initiate AP process---------------------------------------#
        time.sleep(2)
        m = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "fast_search_val")))
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.LINK_TEXT, "Search BIS process")))
        a.move_to_element(n).click().perform()

        #check and input AP number in NCR
        A = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(A).click().send_keys(self.APNAME_InAP,Keys.ENTER).perform()
        time.sleep(3)

        try:
            # open AP &
            B = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(B).click().perform()
            time.sleep(3)
        except:
            # if no visible APNAME enter by user, continue the process
            messagebox.showerror("Error", "No AP found in the NCR")
            bot.close()
        #------------------------------------------AP Description----------------------------------------------------------------------#
        Edit_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='pcEdit']")))
        a.move_to_element(Edit_but).click().perform()

        Desc_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='id_common_descr']")))
        current_text = Desc_but.get_attribute("value")
        Desc_but.clear()
        additional_text = f"\n{self.CustName}\n{self.circuitID}"
        updated_text = current_text + additional_text
        a.move_to_element(Desc_but)
        time.sleep(3)
        a.click().send_keys(updated_text).perform()

        try:
            Acceptance_Detail = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[3]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        except:
            print("No element found")
        else:        
            a.move_to_element(Acceptance_Detail).click().send_keys(IPvalue).perform()
            time.sleep(5)
            linkIP = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]/div[2]")))
            a.move_to_element(linkIP).click().click().perform()

        # Locate the dropdown element
        try:
            dropdown_element = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[2]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/select[1]")))
        except:
            print("No element found")
        else:
            # Initialize the Select class
            dropdown = Select(dropdown_element)
            # Get the available options
            available_options = [option.text for option in dropdown.options]
            # Check if "Yes" is in the available options
            if "Yes" in available_options:
                # Select the "Yes" option
                dropdown.select_by_visible_text("Yes")

        Vendorbut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[2]/tbody[6]/tr[6]/td[2]/span[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/input[1]")))
        a.move_to_element(Vendorbut).click().send_keys(self.VendorName).perform()
        Vendorchoice = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                                                 "/html[1]/body[1]/div[4]/div[2]/div[2]/div[1]/div[1]")))
        a.move_to_element(Vendorchoice).click().perform()
        update_but = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@id='theform_update']")))
        a.move_to_element(update_but).click().perform()
        time.sleep(3)
        #--------------------------------------------------Equipment Assignment-----------------------------------------------------------#
        equip = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[5]/a[1]")))
        a.move_to_element(equip).click(equip).perform()
        clickontick = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(clickontick).click().perform()
        Savebut = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(Savebut).click().perform()
        #_________________________________________________________Integration__________________________________________________________________#
        #checkif there is any Integration Acceptance Phase or not
        try:
            FOP_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(FOP_AObox).perform()
        except:
            print("No element found")
        else:
            A11 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            a.move_to_element(A11).click().perform()
            FOP_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FOP_AOlist).click().perform()
            textarea = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D1).click().perform()
        #-----------------------------------------------------check if there is any input on 2nd box AO-------------------------------------------------------#
        try:
            TSTM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(TSTM_AObox).perform()
        except:
            print("No element found")
        else:
            A22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]")))
            a.move_to_element(A22).perform()
            TSTM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(TSTM_AOlist).click(TSTM_AOlist).perform()
            textarea2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea2).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 3rd box AO-------------------------------------------------------#
        try:
            CMB_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(CMB_AObox).perform()
        except:
            print("No element found")
        else:
            A33 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]")))
            a.move_to_element(A33).perform()
            CMB_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(CMB_AOlist).click().perform()
            textarea3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea3).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 4th box AO-------------------------------------------------------#
        try:
            SPM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(SPM_AObox).perform()
        except:
            print("No element found")
        else:
            A44 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]")))
            a.move_to_element(A44).perform()
            SPM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(SPM_AOlist).click().perform()
            textarea4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea4).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #-----------------------------------------------------check if there is any input on 5th box AO-------------------------------------------------------#
        try:
            FM_AObox = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(FM_AObox).perform()
        except:
            print("No element found")
        else:
            A55 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]")))
            a.move_to_element(A55).perform()
            FM_AOlist = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[3]/div[1]/div[2]/img[1]")))
            a.move_to_element(FM_AOlist).click().perform()
            textarea5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//textarea[@id='jxctl_a_input']")))
            a.move_to_element(textarea5).click().key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
            time.sleep(2)
            a.send_keys(self.AO_InAP).perform()
            time.sleep(1)
            RB3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[2]/div[2]/div[1]/div[1]")))
            a.move_to_element(RB3).click().perform()
            D = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[4]/div[1]/div[2]/div/form[2]/table/tbody/tr/td/div[2]/table/tbody/tr[1]/td[5]")))
            a.move_to_element(D).click().perform()

        #click on Save button
        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
        time.sleep(3)
        #-----------------------------------------upload attachments on Integration Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        z = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[17]")))
        a.move_to_element(z).click().perform()

        try:
            E = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E).click().perform()
        except:
            print("No element found")
        else:
            IA1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(IA1).click().perform()
            F = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F).click().perform()
            time.sleep(3)
            G = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G).click().perform()
            time.sleep(5)
            # Type the file path
            pyautogui.typewrite(self.attachmentfile)
            # Press Enter
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(10)
            #click on Show Items per Page
            K = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE).click().perform()
        except:
            print("No element found")
        else:
            IA2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(IA2).click().perform()
            FF = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF).click().perform()
            time.sleep(3)
            GG = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE1).click().perform()
        except:
            print("No element found")
        else:
            IA3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(IA3).click().perform()
            FF1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF1).click().perform()
            time.sleep(3)
            GG1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG1).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK1).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE2).click().perform()
        except:
            print("No element found")
        else:
            IA4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(IA4).click().perform()
            FF2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF2).click().perform()
            time.sleep(3)
            GG2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG2).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK2).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            EE3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(EE3).click().perform()
        except:
            print("No element found")
        else:
            IA5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(IA5).click().perform()
            FF3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF3).click().perform()
            time.sleep(3)
            GG3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG3).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)   
            pyautogui.press('enter')
            time.sleep(10)
            WebDriverWait(bot, 240).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK3).click().perform()
            time.sleep(3)

        #click on Save button
        L = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L).click().perform()

        M = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(M).click().perform()
 
        L1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[1]/span[1]/label[1]/a[1]/img[1]")))
        a.move_to_element(L1).click().perform()
        #-----------------------------------------upload attachments on Service Acceptance Phase-------------------------------------------#
        #click on attachment to add HOD & BIS Checklist Template
        zx = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[16]")))
        a.move_to_element(zx).click().perform()
        try:
            E7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                    "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(E7).click().perform()
        except:
            print("No element found")
        else:
            A1 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]")))
            a.move_to_element(A1).click().perform()
            F7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(F7).click().perform()
            time.sleep(3)
            G7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(G7).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile, interval=0.01)
            pyautogui.press('enter')
            time.sleep(5)
            WebDriverWait(bot, 90).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            K7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(K7).click().perform()
            time.sleep(3)
        #---------------------------------------------second box attachment------------------------------------------------------------#
        try:
            EE7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(EE7).click().perform()
        except:
            print("No element found")
        else:
            A2 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]")))
            a.move_to_element(A2).click().perform()
            FF7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF7).click().perform()
            time.sleep(3)
            GG7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG7).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(5)
            WebDriverWait(bot, 90).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK7).click().perform()
            time.sleep(3)
        #---------------------------------------------third box attachment------------------------------------------------------------#
        try:
            EE12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(EE12).click().perform()
        except:
            print("No element found")
        else:
            A3 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]")))
            a.move_to_element(A3).click().perform()
            FF12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[3]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF12).click().perform()
            time.sleep(3)
            GG12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG12).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(5)
            WebDriverWait(bot, 90).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK12 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK12).click().perform()
            time.sleep(3)
        #---------------------------------------------fourth box attachment------------------------------------------------------------#
        try:
            EE22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(EE22).click().perform()
        except:
            print("No element found")
        else:
            A4 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]")))
            a.move_to_element(A4).click().perform()
            FF22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[4]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(FF22).click().perform()
            time.sleep(3)
            GG22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(GG22).click().perform()
            time.sleep(3)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(5)
            WebDriverWait(bot, 90).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            KK22 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[3]/div[3]")))
            a.move_to_element(KK22).click().perform()
            time.sleep(3)
        #---------------------------------------------fifth box attachment------------------------------------------------------------#
        try:
            box5 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(box5).click().perform()
        except:
            print("No element found")
        else:
            A6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]")))
            a.move_to_element(A6).click().perform()
            box51 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[5]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box51).click().perform()
            time.sleep(3)
            box52 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box52).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(5)
            WebDriverWait(bot, 90).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            time.sleep(5)
            box53 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
            a.move_to_element(box53).click(box53).perform()
            time.sleep(3)
        #---------------------------------------------sixth box attachment------------------------------------------------------------#
        try:
            time.sleep(3)
            box6 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(box6).click(box6).perform()       
        except:
            print("No element found")
        else:
            A7 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                            "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]")))
            a.move_to_element(A7).click(A7).perform()    
            box61 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[6]/td[15]/div[1]/div[2]/img[1]")))
            a.move_to_element(box61).click(box61).perform()
            time.sleep(3)
            box62 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//a[@title='Add or Drop File']")))
            a.move_to_element(box62).click().perform()
            time.sleep(5)
            pyautogui.typewrite(self.attachmentfile)
            pyautogui.press('enter')
            time.sleep(5)
            WebDriverWait(bot, 90).until(EC.invisibility_of_element_located((By.XPATH,  "//div[@class='nc-remaining-estimate nc-attach-row']")))
            box63 = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                                "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[2]/td[16]")))
            a.move_to_element(box63).click(box63).perform()
            time.sleep(3)
        clickonsave = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[4]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        a.move_to_element(clickonsave).click().perform()
        time.sleep(2)
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your AP in NCR system') 

    def Mainpage(self):

        self.username1 = StringVar()
        self.checklistform = StringVar()

        self.mainpageframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                    border_color="Gray", border_width=20, corner_radius=20)
        self.mainpageframe.place(x=40, y=40)

        self.titlelabel = customtkinter.CTkLabel(self.mainpageframe, text="Kindly register your name below",
                                                 font=customtkinter.CTkFont(size=40, weight="bold"))
        self.titlelabel.place(x=150, y=100)
        # ---------------------------------------username-----------------------------------------------------------
        user12 = CTkLabel(self.mainpageframe, width=120, text="USERNAME", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        user12.place(x=300, y=240)
        user11 = customtkinter.CTkEntry(master=self.mainpageframe,
                                        placeholder_text="USERNAME",
                                        width=320,
                                        height=25,
                                        border_width=2,
                                        corner_radius=10, textvariable=self.username1)
        user11.place(x=300, y=290)

        self.mainbutton = customtkinter.CTkButton(master=self.mainpageframe, width=120, height=32, border_width=0,
                                                  corner_radius=8, text="SUBMIT", command=self.MainForm)
        self.mainbutton.place(x=370, y=440)

    def MainForm(self):
        self.usernameform = self.username1.get()
        self.checklist = self.checklistform.get()

        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        self.bot.maximize_window()
        self.bot.get( 'https://forms')

        self.bot.find_element(By.XPATH,
                            "//*[@id=\"mG61Hd\"]/div[2]/div/div[2]/div[1]/div/div/div[2]/div/div[1]/div/div[1]/input").send_keys(
            self.usernameform)
        self.bot.implicitly_wait(20)
        self.bot.find_element(By.XPATH, "//*[@id=\"mG61Hd\"]/div[2]/div/div[3]/div[1]/div[1]/div/span/span").click()
        self.bot.implicitly_wait(5)

        time.sleep(20)
        messagebox.showinfo('information',
                            'Registered!')

    def HODTry(self):

        self.contentframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                   border_color="Gray", border_width=20, corner_radius=20)
        self.contentframe.place(x=40, y=40)

        net = customtkinter.CTkButton(self.contentframe, width=140, height=140, text="NCR", border_width=0,
                                      corner_radius=8, command=self.ncr_mainpagehod)
        net.place(x=300, y=220)

        ho = customtkinter.CTkButton(self.contentframe, width=140, height=140, text="HOD TEMPLATE", border_width=0,
                                     corner_radius=8, command=self.tab1)
        ho.place(x=500, y=220)

    def ncr_mainpagehod(self):

        # set default values
        self.mainframe1 = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                 border_color="Gray", border_width=20, corner_radius=20)
        self.mainframe1.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.mainframe1, image=my_image).place(x=60, y=120)

        self.username1 = StringVar()
        self.password1 = StringVar()
        self.circuitid1 = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head1 = CTkLabel(self.mainframe1, width=120, text="PLEASE ENTER THE DETAILS BELOW(NCR HOD)*",
                         font=("Tahoma", 15, "bold"))
        head1.place(x=305, y=40)
        # ---------------------------------------username-----------------------------------------------------------
        user12 = CTkLabel(self.mainframe1, width=120, text="USERNAME", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        user12.place(x=540, y=140)
        user11 = customtkinter.CTkEntry(master=self.mainframe1,
                                        placeholder_text="USERNAME",
                                        width=320,
                                        height=25,
                                        border_width=2,
                                        corner_radius=10, textvariable=self.username1)
        user11.place(x=550, y=190)
        # --------------------------------------password-------------------------------------------------------------
        kata11 = CTkLabel(self.mainframe1, width=120, text="PASSWORD", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        kata11.place(x=540, y=240)
        kata21 = customtkinter.CTkEntry(master=self.mainframe1,
                                        placeholder_text="USERNAME", show="*",
                                        width=320,
                                        height=25,
                                        border_width=2,
                                        corner_radius=10, textvariable=self.password1)
        kata21.place(x=550, y=280)
        # --------------------------------------circuitid-------------------------------------------------------------
        circuitid12 = CTkLabel(self.mainframe1, width=120, text="CIRCUIT ID", fg_color="transparent",
                               font=("Tahoma", 15, "bold"))
        circuitid12.place(x=540, y=330)
        circuitid13 = customtkinter.CTkEntry(master=self.mainframe1,
                                             placeholder_text="USERNAME",
                                             width=320,
                                             height=25,
                                             border_width=2,
                                             corner_radius=10, textvariable=self.circuitid1)
        circuitid13.place(x=550, y=380)
        # --------------------------------------BUTTON-------------------------------------------------------------
        CTkButton(self.mainframe1, width=100, height=30, text='Submit', font=("Tahoma", 10, "bold"),
                  command=self.ncrhod).place(
            x=750, y=500)
        CTkButton(self.mainframe1, text="Back", width=100, height=30, border_width=0, font=("Tahoma", 10, "bold"),
                  command=self.HODTry).place(x=620, y=500)



    def ncrhod(self):
        # NEW-PRKMO003, MCAME0067, anurulag, Newzealand@1234567890        
        self.ID = self.username1.get()
        self.Password = self.password1.get()
        self.CircuitID = self.circuitid1.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the NCR program is finished running.')
        
        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        self.bot.maximize_window()
        self.bot.get(
            'http://')

        # get element
        self.bot.find_element(By.ID, "user").send_keys(self.ID)
        self.bot.implicitly_wait(5)
        self.bot.find_element(By.ID, "pass").send_keys(self.Password)
        self.bot.implicitly_wait(5)
        self.bot.find_element(By.CLASS_NAME, "loginHiddenButton").click()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        a = ActionChains(self.bot)
        y = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "_v9127306460113165068")))
        a.move_to_element(y).click().send_keys(self.CircuitID).send_keys(Keys.ENTER).perform()
        time.sleep(5)

        nnyy = self.bot.find_element(By.XPATH, "//div[@class='pagerPageShow']")
        a.move_to_element(nnyy).perform()
        time.sleep(3)
        # screenshot on In Service Circuit
        self.bot.save_screenshot("Checklist Output\\NCR\\Circuit Status.png")
        time.sleep(5)

        # filter to In Service only
        yy = WebDriverWait(self.bot, 20).until(EC.element_to_be_clickable((By.XPATH,
                                                                           "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[3]/table[1]/tbody[1]/tr[1]/td[2]/div[1]/a[1]/img[1]")))
        a.move_to_element(yy).click().perform()
        z = self.bot.find_element(By.XPATH, "//input[@value='9124352540013835640']")  # checkbox for In Service
        a.move_to_element(z).click().perform()
        zz = self.bot.find_element(By.XPATH,
                                   "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")  # button apply
        a.move_to_element(zz).click().perform()
        time.sleep(5)

        # open path element
        o = self.bot.find_element(By.XPATH, '//*[@id="t4122361118013615427_9123528927413500157_t"]/tbody/tr[1]/td[2]/a')
        a.move_to_element(o).click().perform()
        self.bot.implicitly_wait(5)
        time.sleep(10)

        # 2nd Screenshot - Path Element (routers name/sequence/port no)
        self.bot.save_screenshot("Checklist Output\\NCR\\Path Element.png")
        time.sleep(8)

        # summary tab
        q = self.bot.find_element(By.XPATH, "//a[normalize-space()='Summary Page']")
        a.move_to_element(q).click().perform()
        time.sleep(5)

        # 5th Screenshot - Summary Tab first
        self.bot.save_screenshot("Checklist Output\\NCR\\Summary Page1.png")
        time.sleep(5)

        try:
            # SCROLL DOWN
            q2 = self.bot.find_element(By.XPATH, "//span[normalize-space()='Backup Details']")
        except:
            q22 = self.bot.find_element(By.XPATH, "//span[normalize-space()='Backup Availability']")
            a.move_to_element(q22).click().perform()
            time.sleep(5)
        else:
            a.move_to_element(q2).click().perform()
            time.sleep(5)

        # 6th Screenshot - Summary Tab first
        self.bot.save_screenshot("Checklist Output\\NCR\\Summary Page2.png")
        time.sleep(5)

        # scroll down
        pyautogui.press('down', presses=7)
        time.sleep(8)

        # click on ND
        r = WebDriverWait(self.bot, 20).until(EC.element_to_be_clickable((By.XPATH,
                                                                          '//*[@class="nc-attach-file"][@onclick="if (LoadingHook) { LoadingHook.suspend(); }"]')))
        a.move_to_element(r).click().perform()
        time.sleep(8)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCR and Download Files folders.')
        self.bot.close()
        time.sleep(5)

    def tab1(self):
        self.conframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.conframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.conframe, image=my_image).place(x=60, y=120)

        self.CustomerName11 = StringVar()
        self.CircuitID11 = StringVar()
        self.ServiceTypeBW11 = StringVar()
        self.UserName11 = StringVar()
        self.VendorName11 = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.conframe, width=120, text="PLEASE ENTER THE DETAILS BELOW*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        user2 = CTkLabel(self.conframe, width=120, text="CUSTOMER NAME", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=555, y=75)
        user1 = customtkinter.CTkEntry(master=self.conframe,
                                       placeholder_text="USERNAME",
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.CustomerName11)
        user1.place(x=550, y=115)
        # --------------------------------------Circuit ID-------------------------------------------------------------
        kata1 = CTkLabel(self.conframe, width=120, text="CIRCUIT ID (CAPITAL LETTERS)", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=555, y=155)
        kata = customtkinter.CTkEntry(master=self.conframe,
                                      placeholder_text="USERNAME",
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.CircuitID11)
        kata.place(x=550, y=195)
        # --------------------------------------Customer Name-------------------------------------------------------------
        apname2 = CTkLabel(self.conframe, width=120, text="SERVICE TYPE BW", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=555, y=245)
        apname3 = customtkinter.CTkEntry(master=self.conframe,
                                         placeholder_text="USERNAME",
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.ServiceTypeBW11)
        apname3.place(x=550, y=285)
        # --------------------------------------BW-------------------------------------------------------------
        circuitid2 = CTkLabel(self.conframe, width=120, text="PREPARED BY", fg_color="transparent",
                              font=("Tahoma", 15, "bold"))
        circuitid2.place(x=550, y=335)
        circuitid3 = customtkinter.CTkEntry(master=self.conframe,
                                            placeholder_text="USERNAME",
                                            width=320,
                                            height=25,
                                            border_width=2,
                                            corner_radius=10, textvariable=self.UserName11)
        circuitid3.place(x=550, y=375)
        # --------------------------------------Node-------------------------------------------------------------
        circuitid2 = CTkLabel(self.conframe, width=120, text="VENDOR NAME", fg_color="transparent",
                              font=("Tahoma", 15, "bold"))
        circuitid2.place(x=555, y=415)
        circuitid3 = customtkinter.CTkEntry(master=self.conframe,
                                            placeholder_text="USERNAME",
                                            width=320,
                                            height=25,
                                            border_width=2,
                                            corner_radius=10, textvariable=self.VendorName11)
        circuitid3.place(x=550, y=455)

        self.button1 = CTkButton(self.content_frame, text="Next", width=100, height=30, border_width=0, font=("Tahoma", 10, "bold"),
                                command=self.tab2)
        self.button1.place(x=750, y=550)
        CTkButton(self.content_frame, text="Back", width=100, height=30, border_width=0, font=("Tahoma", 10, "bold"),
            command=self.HODTry).place(x=620, y=550)

    def tab2(self):
        self.button1.destroy()

        self.level = StringVar()
        self.server = StringVar()
        self.racky = StringVar()

        self.content_frame2 = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                     border_color="Gray", border_width=20, corner_radius=20)
        self.content_frame2.place(x=40, y=40)

        my_image2 = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.content_frame2, image=my_image2).place(x=60, y=120)

        # ---------------------------------------label-----------------------------------------------------------
        head1 = CTkLabel(self.content_frame2, width=120, text="PLEASE ENTER THE DETAILS BELOW*",
                         font=("Tahoma", 15, "bold"))
        head1.place(x=305, y=30)
        # ---------------------------------------Level-----------------------------------------------------------
        level1 = CTkLabel(self.content_frame2, width=100, text="LEVEL", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        level1.place(x=530, y=155)
        level2 = customtkinter.CTkEntry(master=self.content_frame2,
                                        placeholder_text="USERNAME",
                                        width=320,
                                        height=25,
                                        border_width=2,
                                        corner_radius=10, textvariable=self.level)
        level2.place(x=550, y=195)
        # --------------------------------------MDF/Server Room-------------------------------------------------------------
        mdf = CTkLabel(self.content_frame2, width=100, text="MDF/SERVER ROOM", fg_color="transparent",
                       font=("Tahoma", 15, "bold"))
        mdf.place(x=560, y=245)
        mdf1 = customtkinter.CTkEntry(master=self.content_frame2,
                                      placeholder_text="USERNAME",
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.server)
        mdf1.place(x=550, y=285)
        # --------------------------------------Rack-------------------------------------------------------------
        rack = CTkLabel(self.content_frame2, width=120, text="LAT/LONGITUDE", fg_color="transparent",
                        font=("Tahoma", 15, "bold"))
        rack.place(x=555, y=330)
        rack1 = customtkinter.CTkEntry(master=self.content_frame2,
                                       placeholder_text="USERNAME",
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.racky)
        rack1.place(x=550, y=365)

        self.button3 = CTkButton(self.content_frame2, text='Next', font=('Times_New_Roman', 25), border_width=2,
                                 corner_radius=7,
                                 command=self.tab3)
        self.button3.place(x=600, y=450)

    def tab3(self):
        self.content_frame3 = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                     border_color="Gray", border_width=20, corner_radius=20)
        self.content_frame3.place(x=40, y=40)

        my_image3 = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.content_frame3, image=my_image3).place(x=60, y=120)

        self.branded = StringVar()
        self.mode = StringVar()
        self.number = StringVar()
        self.managed = StringVar()
        self.dual = StringVar()
        self.DC = StringVar()

        # ---------------------------------------Label-----------------------------------------------------------
        head2 = CTkLabel(self.content_frame3, width=120, text="PLEASE ENTER THE EQUIPMENT DETAILS BELOW*",
                         font=("Tahoma", 15, "bold"))
        head2.place(x=100, y=30)
        # ---------------------------------------Brand-----------------------------------------------------------
        brand = CTkLabel(self.content_frame3, width=120, text="BRAND", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        brand.place(x=525, y=35)
        brand1 = customtkinter.CTkEntry(master=self.content_frame3,
                                        placeholder_text="USERNAME",
                                        width=320,
                                        height=25,
                                        border_width=2,
                                        corner_radius=10, textvariable=self.branded)
        brand1.place(x=550, y=70)
        # --------------------------------------Model-------------------------------------------------------------
        model = CTkLabel(self.content_frame3, width=120, text="MODEL", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        model.place(x=525, y=115)
        model1 = customtkinter.CTkEntry(master=self.content_frame3,
                                        placeholder_text="USERNAME",
                                        width=320,
                                        height=25,
                                        border_width=2,
                                        corner_radius=10, textvariable=self.mode)
        model1.place(x=550, y=150)
        # --------------------------------------Serial Number-------------------------------------------------------------
        serial = CTkLabel(self.content_frame3, width=120, text="SERIAL NUMBER", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        serial.place(x=555, y=195)
        serial1 = customtkinter.CTkEntry(master=self.content_frame3,
                                         placeholder_text="USERNAME",
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.number)
        serial1.place(x=550, y=230)
        # --------------------------------------Managed Router Naming-------------------------------------------------------------
        naming = CTkLabel(self.content_frame3, width=120, text="NETWORK ELEMENT", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        naming.place(x=555, y=270)
        naming1 = customtkinter.CTkEntry(master=self.content_frame3,
                                         placeholder_text="USERNAME",
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.managed)
        naming1.place(x=550, y=305)
        # --------------------------------------Single-------------------------------------------------------------
        single = CTkLabel(self.content_frame3, width=120, text="SINGLE OR DUAL POWER", fg_color="transparent",
                          font=("Tahoma", 15, "bold"))
        single.place(x=555, y=345)
        single1 = customtkinter.CTkEntry(master=self.content_frame3,
                                         placeholder_text="USERNAME",
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.dual)
        single1.place(x=550, y=385)
        # --------------------------------------AC/DC-------------------------------------------------------------
        acdc = CTkLabel(self.content_frame3, width=120, text="AC/DC", fg_color="transparent",
                        font=("Tahoma", 15, "bold"))
        acdc.place(x=525, y=425)
        acdc1 = customtkinter.CTkEntry(master=self.content_frame3,
                                       placeholder_text="USERNAME",
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.DC)
        acdc1.place(x=550, y=465)
        # --------------------------------------BUTTON-------------------------------------------------------------
        self.button5 = CTkButton(self.content_frame3, text='SUBMIT', font=('Times_New_Roman', 20), border_width=2,
                                 corner_radius=7, command=self.create_template)
        self.button5.place(x=550, y=500)
        self.button6 = CTkButton(self.content_frame3, text='EMAIL', font=('Times_New_Roman', 20), border_width=2,
                                 corner_radius=7, command=self.send_email)
        self.button6.place(x=700, y=500)

    def send_email(self):
        # send output to email
        SP_file = max(
            glob.glob(os.path.join('C:\\Automate\\Checklist Output\\HOD', "HOD*")),
            key=os.path.getctime)
        print(SP_file)
        outlook = win32.Dispatch("outlook.application")
        email = outlook.CreateItem(0)
        #email.To = self.touser.get()
        # mail.CC = ""
        #email.Subject = self.subjectuser.get()
        #email.Body = self.messageuser.get()
        attachment_path = SP_file
        email.Attachments.Add(Source=attachment_path)
        # mail.Attachments.Add('C:\\Users\\ACER\\BIS\\Automate\\Checklist Output\\NCR\\Desc.png')
        email.Display()
        time.sleep(5)

        #messagebox.showinfo('information',
                            #'Hi! Your email have been successfully sent')
        # mail.Send()

    def create_template(self):
        # Get user input
        global latest_NDfile1, latest_NDfile2, latest_NDfile3, latest_NDfile4, latest_NDfile5
        CustomerName = self.CustomerName11.get()
        CircuitID2 = self.CircuitID11.get()
        ServiceTypeBW = self.ServiceTypeBW11.get()
        UserName = self.UserName11.get()
        VendorName = self.VendorName11.get()
        Level = self.level.get()
        MDF = self.server.get()
        Rack = self.racky.get()
        Brand = self.branded.get()
        Model = self.mode.get()
        SerialNumber = self.number.get()
        RouterNaming = self.managed.get()
        SinglePower = self.dual.get()
        AcDc = self.DC.get()

        # ---------------------------------------------First Page-----------------------------------------------------#
        # Create a new Microsoft Word document
        document = Document()

        # Add a header to the document
        header = document.sections[0].header

        # adjust the margin to be normal margin
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # Add a table to the header
        table = header.add_table(rows=1, cols=2, width=Cm(16))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'TableGrid'

        # Add some content to the table in header
        table.cell(0, 0).text = "HANDOVER DOCUMENT"
        paragraph1 = table.cell(0, 0).paragraphs[0]
        paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tc = table.cell(0, 0).paragraphs[0].runs
        tc[0].font.size = Pt(17)
        tc[0].font.name = 'Calibri Light (Headings)'
        tc[0].font.bold = True
        tc[0].font.color.rgb = RGBColor(5, 132, 132)

        # add picture to table in header
        paragraph = table.cell(0, 1).paragraphs[0]
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture('images\\maxisheader.png')

        # add paragraph space
        document.add_paragraph().paragraph_format.line_spacing = 10

        # Title - Insert user key in Customer Name
        para = document.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(CustomerName)
        font = run.font
        font.size = Pt(26)
        font.name = 'Times New Roman'
        font.bold = True

        # Title - Insert user key in Circuit ID
        para1 = document.add_paragraph()
        para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = para1.add_run(CircuitID2)
        font1 = run1.font
        font1.size = Pt(26)
        font1.name = 'Times New Roman'
        font1.bold = True

        # Title - Insert user key in Service Type & BW
        para2 = document.add_paragraph()
        para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = para2.add_run(ServiceTypeBW)
        font2 = run2.font
        font2.size = Pt(26)
        font2.name = 'Times New Roman'
        font2.bold = True
        document.add_paragraph().paragraph_format.line_spacing = 4

        # Insert user key in UserName for Prepared by:
        para3 = document.add_paragraph()
        para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run3 = para3.add_run('Prepared by: ' + UserName)
        font3 = run3.font
        font3.size = Pt(12)
        font3.name = 'Calibri (Body)'

        # Insert user key in Vendor Name:
        para4 = document.add_paragraph()
        # para4.paragraph_format.line_spacing = 1.75
        para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run4 = para4.add_run('Vendor Name: ' + VendorName)
        font4 = run4.font
        font4.size = Pt(12)
        font4.name = 'Calibri (Body)'

        # Auto insert date into new paragrah
        para5 = document.add_paragraph()
        para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        my_date = datetime.date.today()
        self.date = '{:%d-%b-%Y}'.format(my_date)  # format date to be 'ex:12-Jan-2023'
        run5 = para5.add_run('Date Prepared: ' + self.date)
        font5 = run5.font
        font5.size = Pt(12)
        font5.name = 'Calibri (Body)'
        document.add_paragraph().paragraph_format.line_spacing = 4

        # Add physical audit table to be signed by FOP after site audit (will review again)
        table1 = document.add_table(rows=5, cols=4)
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        table1.style = 'TableGrid'

        # Add some content to the table
        table1.cell(0, 1).text = "CONTRACTOR"
        table1.cell(0, 2).text = "MAXIS REPRESENTATIVES"
        table1.cell(0, 3).text = "MAXIS REPRESENTATIVES"
        table1.cell(1, 1).text = VendorName
        table1.cell(1, 2).text = "ESEA"
        table1.cell(1, 3).text = "FOP"
        table1.cell(2, 0).text = "SIGNATURE"
        table1.cell(3, 0).text = "NAME"
        table1.cell(4, 0).text = "DATE"

        # adjust all font style for text in table
        for row in table1.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(11)
                        font.name = 'Calibri (Body)'
                        font.bold = True

        # ---------------------------------------------Second Page--------------------------------------------------#
        # This second page consits of table of content

        document.add_page_break()
        document.add_paragraph().paragraph_format.line_spacing = 2
        firstpara = document.add_paragraph()
        firstpara.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        firstrun = firstpara.add_run("TABLE OF CONTENTS")
        firstfont = firstrun.font
        firstfont.size = Pt(12)
        firstfont.name = 'Times New Roman'
        firstfont.bold = True

        table2 = document.add_table(rows=11, cols=1)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add some content to the table
        table2.cell(1, 0).text = (
            "1.0    Background ........................................................................................................................3")
        table2.cell(2, 0).text = (
            "2.0    Overall Diagram .................................................................................................................3")
        table2.cell(3, 0).text = (
            "3.0    Current architecture - TP ....................................................................................................3")
        table2.cell(4, 0).text = (
            "4.0    Detail Design - DDD ..........................................................................................................3")
        table2.cell(5, 0).text = (
            "5.0    Configuration .....................................................................................................................4")
        table2.cell(6, 0).text = (
            "6.0    Physical Audit ....................................................................................................................4")
        table2.cell(7, 0).text = (
            "         6.1    Objectives ..................................................................................................................4")
        table2.cell(8, 0).text = (
            "         6.2    Scope of Work Description .......................................................................................4")
        table2.cell(9, 0).text = (
            "         6.3    Inventory Checklist ................................................................................................4-5")
        table2.cell(10, 0).text = (
            "         6.4    Router Location, Router Information and Other Information ...................................5")

        # adjust all font style for text in table
        for row in table2.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(12)
                        font.name = 'Times New Roman'

        # -----------------------------------------------Third Page--------------------------------------------------#
        # This third page consits of content in HOD template

        # 1.0 Background
        document.add_page_break()
        document.add_paragraph().paragraph_format.line_spacing = 2
        Page3A = document.add_paragraph()
        Page3A.style = 'Heading 1'  # Set the style of the first paragraph to Heading 1
        Page3Arun = Page3A.add_run("1.0    Background")
        Page3Afont = Page3Arun.font
        Page3Afont.name = 'Times New Roman'
        Page3Afont.color.rgb = RGBColor(0, 0, 0)
        Page3Afont.size = Pt(12)

        Page3Aa = document.add_paragraph()
        # Customer Name & Service Type BW value will be inserted here
        Page3Aarun = Page3Aa.add_run(
            "This solution is to provide " + CustomerName + " with service of " + ServiceTypeBW)
        Page3Aafont = Page3Aarun.font
        Page3Aafont.name = 'Times New Roman'
        Page3Aafont.size = Pt(12)
        document.add_paragraph()

        # 2.0 Overall Diagram
        Page3B = document.add_paragraph()
        Page3B.style = 'Heading 1'  # Set the style of the first paragraph to Heading 1
        Page3Brun = Page3B.add_run("2.0    Overall Diagram")
        Page3Bfont = Page3Brun.font
        Page3Bfont.name = 'Times New Roman'
        Page3Bfont.color.rgb = RGBColor(0, 0, 0)
        Page3Bfont.size = Pt(12)
        #document.add_paragraph().paragraph_format.line_spacing = 8

        PageND = document.add_paragraph()
        PageND.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   

        # Overall Diagram to be added from screenshot NCR/file from user
        ND_file = ''
        try:
            ND_file = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\HOD", "Network Diagram*")), key=os.path.getctime)
        except ValueError:
            print("No Network Diagram file found")

        if os.path.exists(ND_file):
            print(f"{ND_file} exists")
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{ND_file} not exists")
            PageNDrun = PageND.add_run("Not Applicable")
            self.set_font(PageNDrun, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #to add ND file name to insert ND file
        if os.path.exists(ND_file):
            PageND1 = document.add_paragraph()
            PageND1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
            PageND1run = PageND1.add_run("\nNetwork Diagram File\n")
            self.set_font(PageND1run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{ND_file} not exists")

        # 3.0 Current architecture - TP
        Page3C = document.add_paragraph()
        Page3C.style = 'Heading 1'  # Set the style of the first paragraph to Heading 1
        Page3Crun = Page3C.add_run("3.0    Current architecture - TP")
        Page3Cfont = Page3Crun.font
        Page3Cfont.name = 'Times New Roman'
        Page3Cfont.color.rgb = RGBColor(0, 0, 0)
        Page3Cfont.size = Pt(12)

        Page3Cc = document.add_paragraph()
        Page3Ccrun = Page3Cc.add_run("Please refer to attachment TP file")
        Page3Ccfont = Page3Ccrun.font
        Page3Ccfont.name = 'Times New Roman'
        Page3Ccfont.size = Pt(12)

        PageTP = document.add_paragraph()
        PageTP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    

        latest_TP_file = ''
        try:
            latest_TP_file = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\HOD", "TP*")), key=os.path.getctime)
        except ValueError:
            print("No TP file found")

        if os.path.exists(latest_TP_file):
            print(f"{latest_TP_file} exists")
            #document.add_paragraph().paragraph_format.line_spacing = 7
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_TP_file} not exists")
            PageTPrun = PageTP.add_run("Not Applicable")
            self.set_font(PageTPrun, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

         #to add TP file name to insert ND file
        if os.path.exists(latest_TP_file):
            PageTP1 = document.add_paragraph()
            PageTP1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
            PageTP1run = PageTP1.add_run("\nTP File\n")
            self.set_font(PageTP1run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_TP_file} not exists")       
        
        # 4.0 Detail Design – DDD
        Page3D = document.add_paragraph()
        Page3D.style = 'Heading 1'  # Set the style of the first paragraph to Heading 1
        Page3Drun = Page3D.add_run("4.0    Detail Design - DDD")
        Page3Dfont = Page3Drun.font
        Page3Dfont.name = 'Times New Roman'
        Page3Dfont.color.rgb = RGBColor(0, 0, 0)
        Page3Dfont.size = Pt(12)

        Page3Dd = document.add_paragraph()
        Page3Ddrun = Page3Dd.add_run("Please refer to attachment file")
        Page3Ddfont = Page3Ddrun.font
        Page3Ddfont.name = 'Times New Roman'
        Page3Ddfont.size = Pt(12)

        PageDDD = document.add_paragraph()
        PageDDD.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        DDD_file = ''
        try:
            DDD_file = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\HOD", "DDD*")), key=os.path.getctime)
        except ValueError:
            print("No DDD file found")

        if os.path.exists(DDD_file):
            print(f"{DDD_file} exists")
            #document.add_paragraph().paragraph_format.line_spacing = 7
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{DDD_file} not exists")
            PageDDDrun = PageDDD.add_run("Not Applicable")
            self.set_font(PageDDDrun, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

         #to add TP file name to insert ND file
        if os.path.exists(DDD_file):
            PageDDD1 = document.add_paragraph()
            PageDDD1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            PageDDD1run = PageDDD1.add_run("\nDDD File\n")
            self.set_font(PageDDD1run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{DDD_file} not exists")     

        # ----------------------------------------------Fourth Page--------------------------------------------------#
        # This fourth page consits of content in HOD template
        # 5.0	Configuration
        Page3E = document.add_paragraph()
        Page3E.style = 'Heading 1'  # Set the style of the first paragraph to Heading 1
        Page3Erun = Page3E.add_run("5.0    Port Configuration and Fiber Reading")
        Page3Efont = Page3Erun.font
        Page3Efont.name = 'Times New Roman'
        Page3Efont.color.rgb = RGBColor(0, 0, 0)
        Page3Efont.size = Pt(12)

        Page3Ee = document.add_paragraph()
        Page3Eerun = Page3Ee.add_run("Please refer to attachment file")
        Page3Eefont = Page3Eerun.font
        Page3Eefont.name = 'Times New Roman'
        Page3Eefont.size = Pt(12)

        PageConfig = document.add_paragraph()
        PageConfig.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        configuration_file = ''
        try:
            configuration_file = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\HOD", "Configuration*")), key=os.path.getctime)
        except ValueError:
            print("No Configuration file found")

        if os.path.exists(configuration_file):
            print(f"{configuration_file} exists")
            #document.add_paragraph().paragraph_format.line_spacing = 5
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{configuration_file} not exists")
            PageConfigrun = PageConfig.add_run("Not Applicable")
            self.set_font(PageConfigrun, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

         #to add configuration file name to insert configuration file
        if os.path.exists(configuration_file):
            PageConfig1 = document.add_paragraph()
            PageConfig1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            PageConfig1run = PageConfig1.add_run("\nConfiguration File\n")
            self.set_font(PageConfig1run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{configuration_file} not exists")    

        # 6.0	Physical Audit
        Page3F = document.add_paragraph()
        Page3F.style = 'Heading 1'  # Set the style of the first paragraph to Heading 1
        Page3Frun = Page3F.add_run("6.0    Physical Audit")
        Page3Ffont = Page3Frun.font
        Page3Ffont.name = 'Times New Roman'
        Page3Ffont.color.rgb = RGBColor(0, 0, 0)
        Page3Ffont.size = Pt(12)

        # 6.1 Objectives (P6B = Paragraph for 6.o heading)
        P6B = document.add_paragraph()
        P6B.paragraph_format.left_indent = Inches(0.5)
        P6B.style = 'Heading 2'  # Set the style of the first paragraph to Heading 1
        P6Brun = P6B.add_run("6.1    Objectives")
        P6Bfont = P6Brun.font
        P6Bfont.name = 'Times New Roman'
        P6Bfont.color.rgb = RGBColor(0, 0, 0)
        P6Bfont.size = Pt(12)

        P6Bb = document.add_paragraph()
        P6Bb.paragraph_format.left_indent = Inches(0.5)
        P6Bbrun = P6Bb.add_run(
            "•	This document describes the scope of work to be performed by Maxis for the customer's project.")
        P6Bbfont = P6Bbrun.font
        P6Bbfont.name = 'Times New Roman'
        P6Bbfont.size = Pt(12)

        # 6.2 Scope of Work Description (P6B = Paragraph for 6.o heading)
        P6C = document.add_paragraph()
        P6C.paragraph_format.left_indent = Inches(0.5)
        P6C.style = 'Heading 2'  # Set the style of the first paragraph to Heading 1
        P6Crun = P6C.add_run("6.2    Scope of Work Description")
        P6Cfont = P6Crun.font
        P6Cfont.name = 'Times New Roman'
        P6Cfont.color.rgb = RGBColor(0, 0, 0)
        P6Cfont.size = Pt(12)

        P6Cc = document.add_paragraph()
        P6Cc.paragraph_format.left_indent = Inches(0.5)
        P6Ccrun = P6Cc.add_run("•	To perform audit on managed router hardware")
        P6Ccfont = P6Ccrun.font
        P6Ccfont.name = 'Times New Roman'
        P6Ccfont.size = Pt(12)

        # 6.3 Inventory Checklist (P6B = Paragraph for 6.0 heading)
        P6D = document.add_paragraph()
        P6D.paragraph_format.left_indent = Inches(0.5)
        P6D.style = 'Heading 2'  # Set the style of the first paragraph to Heading 1
        P6Drun = P6D.add_run("6.3    Inventory Checklist")
        P6Dfont = P6Drun.font
        P6Dfont.name = 'Times New Roman'
        P6Dfont.color.rgb = RGBColor(0, 0, 0)
        P6Dfont.size = Pt(12)
        #document.add_paragraph().paragraph_format.line_spacing = 1

        PP1 = "C:\\Automate\\Checklist Output\\NCR\\Path Element.png"
        PP2 = "C:\\Automate\\Checklist Output\\NCR\\Summary Page1.png"
        PP3 = "C:\\Automate\\Checklist Output\\NCR\\Summary Page2.png"

        # Create a flag variable for each picture path
        PP1_inserted = False
        PP2_inserted = False
        PP3_inserted = False
        # Loop through all paths
        # insert screenshot summary page from NCR automation
        for path in [PP1, PP2, PP3]:
            if os.path.exists(path):
                print(f"{path} exists")
                document.add_paragraph().paragraph_format.line_spacing = 1
                # Check if the picture file exists
                if path == PP1 and not PP1_inserted:
                    # Insert the picture into the document
                    Picture2 = document.add_paragraph()
                    Picture2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    runPicture2 = Picture2.add_run()
                    runPicture2.add_picture(PP1, width=Cm(16),height=Cm(9))
                    PP1_inserted = True
                # Check if the picture file exists
                elif path == PP2 and not PP2_inserted:
                    # Insert the picture into the document
                    document.add_paragraph().paragraph_format.line_spacing = 1
                    Picture3 = document.add_paragraph()
                    Picture3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    runPicture3 = Picture3.add_run()
                    runPicture3.add_picture(PP2,width=Cm(16), height=Cm(9))
                    PP2_inserted = True
                elif path == PP3 and not PP3_inserted:
                    document.add_paragraph().paragraph_format.line_spacing = 1
                    Picture4 = document.add_paragraph()
                    Picture4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    runPicture4 = Picture4.add_run()
                    runPicture4.add_picture(PP3,width=Cm(16), height=Cm(9))

                    document.add_page_break()
                    document.add_paragraph().paragraph_format.line_spacing = 2
                    PP3_inserted = True
            else:
                    print('Picture not exist')

        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not PP1_inserted, not PP2_inserted, not PP3_inserted]):
            NotApp = document.add_paragraph()
            NotApp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            NotApprun = NotApp.add_run("Not Applicable")
            self.set_font(NotApprun, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #document.add_paragraph().paragraph_format.line_spacing = 1       
        # ----------------------------------------------Fifth Page--------------------------------------------------#
        # This fifth page consits of content in HOD template
        # Add a table to the new paragraph
        newtable = document.add_table(rows=2, cols=2)
        newtable.alignment = WD_TABLE_ALIGNMENT.CENTER
        newtable.style = 'TableGrid'
        # Add some content to the table
        newtable.cell(0, 0).text = "ROUTER LOCATION"
        # take value from user key in Lat/Longitude, Level, MDF/Serverroom)
        newtable.cell(0, 1).add_paragraph('Lat/Longitude : ' + Rack)
        newtable.cell(0, 1).add_paragraph()
        newtable.cell(0, 1).add_paragraph('Level: ' + Level)
        newtable.cell(0, 1).add_paragraph()
        newtable.cell(0, 1).add_paragraph('MDF/Server room : ' + MDF)
        newtable.cell(0, 1).add_paragraph()
        newtable.cell(1, 0).text = "ROUTER INFORMATION"
        # take value from user key in Brand/Model/Serial Num/Router Naming/Single or Dual Power/AC or DC Power)
        newtable.cell(1, 1).add_paragraph('Model: ' + Model)
        newtable.cell(1, 1).add_paragraph()
        newtable.cell(1, 1).add_paragraph('Brand : ' + Brand)
        newtable.cell(1, 1).add_paragraph()
        newtable.cell(1, 1).add_paragraph('Serial Number : ' + SerialNumber)
        newtable.cell(1, 1).add_paragraph()
        newtable.cell(1, 1).add_paragraph('Network Element: ' + RouterNaming)
        newtable.cell(1, 1).add_paragraph()
        newtable.cell(1, 1).add_paragraph('Single or Dual Power : ' + SinglePower)
        newtable.cell(1, 1).add_paragraph()
        newtable.cell(1, 1).add_paragraph('AC / DC Power :  ' + AcDc)
        newtable.cell(1, 1).add_paragraph()

        for cell in newtable.columns[0].cells:
            cell.width = Cm(5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell in newtable.columns[1].cells:
            cell.width = Cm(12)
            paragraphs = cell.paragraphs

        # adjust all font style for text in table
        for row in newtable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(11)
                        font.name = 'Times New Roman'

        document.add_paragraph().paragraph_format.line_spacing = 2

        # 6.4 Router Location, Router Information and Other Information (P6B = Paragraph for 6.o heading)
        P6E = document.add_paragraph()
        P6E.paragraph_format.left_indent = Inches(0.5)
        P6E.style = 'Heading 2'  # Set the style of the first paragraph to Heading 1
        P6Erun = P6E.add_run("6.4    Router Location, Router Information and Other Information")
        P6Efont = P6Erun.font
        P6Efont.name = 'Times New Roman'
        P6Efont.color.rgb = RGBColor(0, 0, 0)
        P6Efont.size = Pt(12)

        P6Ee = document.add_paragraph()
        P6Ee.paragraph_format.left_indent = Inches(0.5)
        P6Eerun = P6Ee.add_run("Please refer to attachment file")
        P6Eefont = P6Eerun.font
        P6Eefont.name = 'Times New Roman'
        P6Eefont.size = Pt(12)

        Pagerouter = document.add_paragraph()
        Pagerouter.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        router_file = ''
        try:
            router_file = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\HOD", "Router and Prewired info*")), key=os.path.getctime)
        except ValueError:
            print("No Router and Prewired info file found")

        if os.path.exists(router_file):
            print(f"{router_file} exists")
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{router_file} not exists")
            Pagerouterrun = Pagerouter.add_run("Not Applicable")
            self.set_font(Pagerouterrun, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

         #to add configuration file name to insert configuration file
        if os.path.exists(router_file):
            Pagerouter1 = document.add_paragraph()
            Pagerouter1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Pagerouter1run = Pagerouter1.add_run("\nRouter and Prewired info File\n")
            self.set_font(Pagerouter1run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{router_file} not exists")    

        # Save the document
        document.save('C:\\Automate\\Checklist Output\\HOD\\HOD Template.docx')

        # start a Word application instance
        word = win32com.client.gencache.EnsureDispatch('Word.Application')
        docpath = "C:\\Automate\\Checklist Output\\HOD\\HOD Template.docx"
        # create a new Word document
        document = word.Documents.Open(docpath)

        latest_NDfile1 = ''
        latest_DDDfile = ''
        latest_Configurationfile = ''
        latest_Routerfile = ''
        latest_TPfile = ''

        try:
            latest_NDfile1 = max(
                glob.glob(os.path.join('C:\\Automate\\Checklist Output\\HOD', "Network Diagram*")),
                key=os.path.getctime)
        except:
            print('No ND file')

        try:
            latest_DDDfile = max(
            glob.glob(os.path.join('C:\\Automate\\Checklist Output\\HOD', "DDD*")),
            key=os.path.getctime)
        except:
            print('No DDD file')

        try:
            latest_Configurationfile = max(
            glob.glob(os.path.join('C:\\Automate\\Checklist Output\\HOD', "Configuration*")),
            key=os.path.getctime)
        except:
            print('No Configuration file')


        try:
            latest_Routerfile = max(
            glob.glob(os.path.join('C:\\Automate\\Checklist Output\\HOD', "Router and Prewired info*")),
            key=os.path.getctime)
        except:
            print('No router & prewired info file')

        try:
            latest_TPfile = max(
            glob.glob(os.path.join('C:\\Automate\\Checklist Output\\HOD', "TP*")),
            key=os.path.getctime)
        except:
            print('No TP file')

        #to add ole object on each files
        # Loop through each paragraph of the document, and if the text 'Visio Diagram' is found, insert the file using the InlineShapes method
        for para in document.Paragraphs:
            if 'Network Diagram File' in para.Range.Text:
                if os.path.exists(latest_NDfile1):
                    print(f"{latest_NDfile1} exists")
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_NDfile1)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='Network Diagram')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_NDfile1} not exists")
                    break
        
        #TP file
        for para in document.Paragraphs:
            if 'TP File' in para.Range.Text:
                if os.path.exists(latest_TPfile):
                    print(f"{latest_TPfile} exists")
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_TPfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{AC76BA86-1033-FFFF-7760-BC15014EA700}\\_PDFFile.ico')),
                        DisplayAsIcon=1) #IconLabel='Network Diagram')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_TPfile} not exists")
                    break

        #DDD file
        for para in document.Paragraphs:
            if 'DDD File' in para.Range.Text:
                if os.path.exists(latest_DDDfile):
                    print(f"{latest_DDDfile} exists")
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_DDDfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\wordicon.exe')),
                        DisplayAsIcon=1) #IconLabel='Network Diagram')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_DDDfile} not exists")
                    break

        #DDD file
        for para in document.Paragraphs:
            if 'Configuration File' in para.Range.Text:
                if os.path.exists(latest_Configurationfile):
                    print(f"{latest_Configurationfile} exists")
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_Configurationfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\wordicon.exe')),
                        DisplayAsIcon=1) #IconLabel='Network Diagram')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_Configurationfile} not exists")
                    break

        #DDD file
        for para in document.Paragraphs:
            if 'Router and Prewired info File' in para.Range.Text:
                if os.path.exists(latest_Routerfile):
                    print(f"{latest_Routerfile} exists")
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_Routerfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\wordicon.exe')),
                        DisplayAsIcon=1) #IconLabel='Network Diagram')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_Routerfile} not exists")
                    break

        # save the document
        document.SaveAs("C:\\Automate\\Checklist Output\\HOD\\HOD Template.docx")
        document.Save()
        document.Close()

        # close the Word application instance
        word.Quit()

        messagebox.showinfo('information',
                            'Hi, HOD Template is successfully generated')

    def Automate(self):

        self.bisframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.bisframe.place(x=40, y=40)

        self.button_ncr = customtkinter.CTkButton(master=self.bisframe, width=120, height=100, border_width=0,
                                                  corner_radius=8, command=self.NCR_MainPage,
                                                  text="NCR")
        self.button_ncr.place(x=300, y=100)

        self.button_ncr_min = customtkinter.CTkButton(master=self.bisframe, width=120, height=100, border_width=0,
                                                      corner_radius=8,
                                                      text="NCR.MIN", command=self.NCR_MinPage)
        self.button_ncr_min.place(x=300, y=250)

        self.button_nms = customtkinter.CTkButton(master=self.bisframe, width=120, height=100, border_width=0,
                                                  corner_radius=8,
                                                  text="NMS", command=self.NMS_Mainpage)
        self.button_nms.place(x=300, y=400)

        self.button_meraki_main = customtkinter.CTkButton(master=self.bisframe, width=120, height=100, border_width=0,
                                                          corner_radius=8,
                                                          text="MERAKI MAIN", command=self.Meraki_MainPage)
        self.button_meraki_main.place(x=500, y=100)

        self.button_meraki_idm = customtkinter.CTkButton(master=self.bisframe, width=120, height=100, border_width=0,
                                                         corner_radius=8,
                                                         text="MERAKI IDM", command=self.MerakiIDMInput)
        self.button_meraki_idm.place(x=500, y=250)
        self.button_template = customtkinter.CTkButton(master=self.bisframe, width=120, height=100, border_width=0,
                                                       corner_radius=8, command=self.BIS_Template, text="TEMPLATE")
        self.button_template.place(x=500, y=400)

    def NCR_MainPage(self):
        self.ncrframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.ncrframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.content_frame, image=my_image).place(x=60, y=120)

        self.username11 = StringVar()
        self.password11 = StringVar()
        self.apname11 = StringVar()
        self.circuitid11 = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.ncrframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (NCR)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        user2 = CTkLabel(self.ncrframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=495, y=85)
        user1 = customtkinter.CTkEntry(master=self.ncrframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.username11)
        user1.place(x=500, y=125)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.ncrframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=495, y=165)
        self.password_entry = customtkinter.CTkEntry(master=self.ncrframe,
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, show="*", textvariable=self.password11)
        self.password_entry.place(x=500, y=205)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.ncrframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=202)
        # --------------------------------------APNAME-------------------------------------------------------------
        apname2 = CTkLabel(self.ncrframe, width=120, text="AP NAME", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=480, y=255)
        apname3 = customtkinter.CTkEntry(master=self.ncrframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.apname11)
        apname3.place(x=500, y=295)
        # --------------------------------------CIRCUITID-------------------------------------------------------------
        circuitid2 = CTkLabel(self.ncrframe, width=120, text="CIRCUIT ID*", fg_color="transparent",
                              font=("Tahoma", 15, "bold"))
        circuitid2.place(x=495, y=345)
        circuitid3 = customtkinter.CTkEntry(master=self.ncrframe,
                                            width=320,
                                            height=25,
                                            border_width=2,
                                            corner_radius=10, textvariable=self.circuitid11)
        circuitid3.place(x=500, y=385)

        CTkButton(self.ncrframe, text='Submit', width=100, height=30, border_width=0, 
                                 command=self.ncr).place(x=760, y=500)
        CTkButton(master=self.ncrframe, text="Back", width=100, height=30, border_width=0,
            command=self.Automate).place(x=650, y=500)

    def ncr(self):
    # NEW-PRKMO003, MCAME0067, anurulag, Lima@1234567890
        
        ID = self.username11.get()
        Password = self.password11.get()
        APNAME = self.apname11.get()
        CircuitID = self.circuitid11.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the NCR program is finished running.')

        # set the download directory and circuit ID
        download_dir = "C:\\Automate\\Checklist Output\\NCR"

        self.options = Options()

        # this parameter tells Chrome that
        # it should be run without UI (Headless)
        self.options.add_argument("--disable-infobars")
        self.options.add_argument("--start-maximized")
        self.options.add_argument("--disable-extensions")
        self.options.add_argument('--window-size=1920,1080')
        self.options.add_argument("--ignore-certificate-errors")
        self.options.add_argument('--ignore-ssl-errors=yes')
        prefs = {"download.default_directory": download_dir}
        self.options.add_experimental_option("prefs", prefs)
        self.options.add_argument("--headless=new")

        bot = webdriver.Chrome(service=self.s, options=self.options)
        bot.maximize_window()
        #bot.get('http://')
        # get IP Address
        url = "http://"
        #self.driver.get(f"http://{IPAddress1}")

        try:
            bot.get(url)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            bot.close()

        # get element
        bot.find_element(By.ID, "user").send_keys(ID)
        bot.implicitly_wait(8)
        bot.find_element(By.ID, "pass").send_keys(Password)
        bot.implicitly_wait(8)
        bot.find_element(By.CLASS_NAME, "loginHiddenButton").click()
        bot.implicitly_wait(8)
        time.sleep(8)

        a = ActionChains(bot)
        if APNAME: #if there is AP name entered by user, search the AP Name in the Search BIS Process page
            x = WebDriverWait(bot, 60).until(EC.element_to_be_clickable((By.ID, "_v_1")))
            a.move_to_element(x).click().send_keys(APNAME).send_keys(Keys.ENTER).perform()
            time.sleep(5)

            try:
                # open AP &
                bot.find_element(By.CLASS_NAME, "wrapLink").click()
                bot.implicitly_wait(8)
                time.sleep(8)
            except:
                # if no visible APNAME enter by user, continue the process
                print('No APNAME submitted by user. System continue to search using circuit ID')
            else:
                # 1st Screenshot - AP Description
                bot.save_screenshot("Checklist Output\\NCR\\AP Description.png")
                time.sleep(5)

                # navigate to bottom element to take screenshot for attachments
                cc = WebDriverWait(bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='itemsPP']")))
                a.move_to_element(cc).perform()
                time.sleep(8)

                # 2nd Screenshot - Attachments checking
                bot.save_screenshot("Checklist Output\\NCR\\Attachments Upload.png")
                time.sleep(5)

        m = bot.find_element(By.ID, "fast_search_val")
        a.move_to_element(m).perform()

        # Search service circuit - cct ID
        n = bot.find_element(By.LINK_TEXT, "Search Service Circuit")
        a.move_to_element(n).click().perform()
        bot.implicitly_wait(5)
        time.sleep(8)

        y = WebDriverWait(bot, 60).until(EC.element_to_be_clickable((By.ID, "_v9127306460113165068"))) #search the circuit ID in search box
        a.move_to_element(y).click().send_keys(CircuitID).send_keys(Keys.ENTER).perform()
        time.sleep(8)

        nnyy = bot.find_element(By.XPATH, "//div[@class='pagerPageShow']") #go to the Showing 1..2 of 2
        a.move_to_element(nnyy).perform()
        time.sleep(3)

        # screenshot on In Service Circuit
        bot.save_screenshot("Checklist Output\\NCR\\Circuit Status.png")
        time.sleep(5)

        # filter to In Service only
        yy = WebDriverWait(bot, 60).until(EC.element_to_be_clickable((By.XPATH,
             "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/thead[1]/tr[1]/th[3]/table[1]/tbody[1]/tr[1]/td[2]/div[1]/a[1]/img[1]"))) #filter icon
        a.move_to_element(yy).click().perform()
        z = bot.find_element(By.XPATH, "//input[@value='9124352540013835640']")  # checkbox for In Service
        a.move_to_element(z).click().perform()
        zz = bot.find_element(By.XPATH,
                                "//div[@class='buttonInner']//a[@role='button'][normalize-space()='Apply']")  # button apply
        a.move_to_element(zz).click().perform()
        time.sleep(8)

        # click on the serrvice circuit link to open path element
        o = bot.find_element(By.XPATH,
                                '//*[@id="t4122361118013615427_9123528927413500157_t"]/tbody/tr[1]/td[2]/a')
        a.move_to_element(o).click().perform()
        bot.implicitly_wait(5)
        time.sleep(10)

        # 2nd Screenshot - Path Element (routers name/sequence/port no)
        bot.save_screenshot("Checklist Output\\NCR\\Path Element.png")
        time.sleep(8)

        # Click Main Tab
        p = bot.find_element(By.XPATH, "//a[normalize-space()='Main']")
        a.move_to_element(p).click().perform()
        time.sleep(5)

        # 3rd Screenshot - Main Tab (Circuit ID)
        bot.save_screenshot("Checklist Output\\NCR\\Main Tab.png")
        time.sleep(8)

        # Parameters tab
        pq = bot.find_element(By.XPATH, "//a[normalize-space()='Parameters']")
        a.move_to_element(pq).click().perform()
        time.sleep(5)

        # 4th Screenshot - Parameters Tab (Circuit ID)
        bot.save_screenshot("Checklist Output\\NCR\\Parameters Tab.png")
        time.sleep(8)

        # summary tab
        q = bot.find_element(By.XPATH, "//a[normalize-space()='Summary Page']")
        a.move_to_element(q).click().perform()
        time.sleep(5)

        # 5th Screenshot - Summary Tab first
        bot.save_screenshot("Checklist Output\\NCR\\Summary Page1.png")
        time.sleep(8)

        # SCROLL DOWN & click Backup details
        q2 = bot.find_element(By.XPATH, "//span[normalize-space()='Backup Details']")

        a.move_to_element(q2).click().perform()
        time.sleep(5)

        # 6th Screenshot - Summary Tab first
        bot.save_screenshot("Checklist Output\\NCR\\Summary Page2.png")
        time.sleep(5)

        # scroll down
        pyautogui.press('down', presses=7)
        time.sleep(8)

        # click on ND
        r = WebDriverWait(bot, 60).until(EC.element_to_be_clickable((By.XPATH,
                                                                        '//*[@class="nc-attach-file"][@onclick="if (LoadingHook) { LoadingHook.suspend(); }"]')))
        a.move_to_element(r).click().perform()
        time.sleep(20)

        bot.refresh()
        time.sleep(5)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCR folders.')
        bot.close()
        time.sleep(5)
    

    def NCR_MinPage(self):
        self.ncrminframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                  border_color="Gray", border_width=20, corner_radius=20)
        self.ncrminframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.ncrminframe, image=my_image).place(x=60, y=120)

        self.usernameMIN11 = StringVar()
        self.passwordMIN11 = StringVar()
        self.apnameMIN11 = StringVar()
        self.AOMIN11 = StringVar()
        self.MINAO22 = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.ncrminframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (NCR MIN)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        user2 = CTkLabel(self.ncrminframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=495, y=85)
        user1 = customtkinter.CTkEntry(master=self.ncrminframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.usernameMIN11)
        user1.place(x=500, y=125)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.ncrminframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=490, y=165)
        self.password_entry = customtkinter.CTkEntry(master=self.ncrminframe,
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, show="*", textvariable=self.passwordMIN11)
        self.password_entry.place(x=500, y=205)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.ncrminframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=202)
        # --------------------------------------APNAME-------------------------------------------------------------
        apname2 = CTkLabel(self.ncrminframe, width=120, text="AP NAME", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=480, y=255)
        apname3 = customtkinter.CTkEntry(master=self.ncrminframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.apnameMIN11)
        apname3.place(x=500, y=295)
        # --------------------------------------MIN Acceptance Object 1-------------------------------------------------------------
        MINAO1 = CTkLabel(self.ncrminframe, width=120, text="NETWORK ELEMENT MIN", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        MINAO1.place(x=505, y=345)
        MINAO1 = customtkinter.CTkEntry(master=self.ncrminframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.AOMIN11)
        MINAO1.place(x=500, y=385)
        # --------------------------------------MIN Acceptance Object 2-------------------------------------------------------------
        MINAO2 = CTkLabel(self.ncrminframe, width=120, text="NETWORK ELEMENT MIN", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        MINAO2.place(x=505, y=435)
        MINAO2 = customtkinter.CTkEntry(master=self.ncrminframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.MINAO22)
        MINAO2.place(x=500, y=475)

        self.button1 = CTkButton(self.ncrminframe, text='Submit',  width=100, height=30, border_width=0, command=self.ncr_min)
        self.button1.place(x=760, y=520)
        CTkButton(master=self.ncrminframe, text="Back", width=100, height=30, border_width=0,
        command=self.Automate).place(x=650, y=520)

    def ncr_min(self):
        ID = self.usernameMIN11.get()
        Password = self.passwordMIN11.get()
        APNAME = self.apnameMIN11.get()
        MINAObject1 = self.AOMIN11.get()
        MINAObject2 = self.MINAO22.get()

        messagebox.showinfo('information',
                    'Hi! Your BIS Automation process is currently in progress.Please wait until the NCR MIN program is finished running.')

        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        self.bot.maximize_window()
        self.bot.get(
            'http://')

        # get element
        self.bot.find_element(By.ID, "user").send_keys(ID)
        self.bot.implicitly_wait(5)
        self.bot.find_element(By.ID, "pass").send_keys(Password)
        self.bot.implicitly_wait(5)
        self.bot.find_element(By.CLASS_NAME, "loginHiddenButton").click()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        a = ActionChains(self.bot)
        #check if there is AP-NAME enter by users

        b = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "_v_1")))
        a.move_to_element(b).click().send_keys(APNAME).send_keys(Keys.ENTER).perform()
        time.sleep(5)

        try:
            # open AP &
            c = WebDriverWait(self.bot, 20).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink")))
            a.move_to_element(c).click().perform()
            self.bot.implicitly_wait(30)
            time.sleep(5)
        except:
            #put all network element process here
            aMIN = WebDriverWait(self.bot, 20).until(EC.element_to_be_clickable((By.ID, "fast_search_val"))) #clisk on the search box
            a.move_to_element(aMIN).perform()
            time.sleep(2)

            # Search service circuit - cct ID
            bMIN = self.bot.find_element(By.LINK_TEXT, "Search Network Elements") #search in the dropdown searchbox
            a.move_to_element(bMIN).click().perform()
            self.bot.implicitly_wait(5)
            time.sleep(8)

            cMIN = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "_v_1"))) #click on the search box
            a.move_to_element(cMIN).click().send_keys(MINAObject1).send_keys(Keys.ENTER).perform()
            time.sleep(8)

            dMIN = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.CLASS_NAME, "wrapLink"))) #click on the MIN product link appear on the search results
            a.move_to_element(dMIN).click().perform()
            time.sleep(8)

            # Screenshot - MIN AP Info
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\Info Top Page.png")
            time.sleep(5)

            # search on Sharing Info - SIM Info
            eMIN = WebDriverWait(self.bot, 50).until(
            EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Sharing Information']")))
            a.move_to_element(eMIN).perform()
            time.sleep(8)

            # 4th Screenshot - SIM Info
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\SIM Info.png")
            time.sleep(5)

            # move to Name in Management System in order to ss bottom page
            fMIN = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Name in Management System']")))
            a.move_to_element(fMIN).perform()
            time.sleep(5)

            # 5th Screenshot - MIN AP Info
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\Info Bottom Page.png")
            time.sleep(5)

            gMIN = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Name in Management System']"))) #go to Name in Management System element
            a.move_to_element(gMIN).perform()
            time.sleep(5)

            # Go back to the previous webpage
            self.bot.back()
            time.sleep(10)

            hMIN = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "_v_1")))
            hMIN.clear()
            a.move_to_element(hMIN).click().send_keys(MINAObject2).send_keys(Keys.ENTER).perform()
            time.sleep(8)

            try:
                # Find the element with the given ID
                iMIN = self.bot.find_element(By.CLASS_NAME, "wrapLink")
                a.move_to_element(iMIN).click().perform()
                time.sleep(8)
            except:
                messagebox.showinfo('information',
                                    'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCR>MIN TEMPLATE folder.')
                self.bot.close()
                time.sleep(3)
            else:
                # Screenshot - MIN AP Info
                self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\Info Top Page.png")
                time.sleep(5)

                # search on Sharing Info - SIM Info
                jMIN = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Sharing Information']")))
                a.move_to_element(jMIN).perform()
                time.sleep(8)

                # 4th Screenshot - SIM Info
                self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\SIM Info.png")
                time.sleep(5)

                # move to Name in Management System in order to ss bottom page
                kMIN = WebDriverWait(self.bot, 60).until(
                    EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Name in Management System']")))
                a.move_to_element(kMIN).perform()
                time.sleep(5)

                # 5th Screenshot - MIN AP Info
                self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\Info Bottom Page.png")
                time.sleep(5)

                messagebox.showinfo('information',
                'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCR>MIN TEMPLATE folder.')

                self.bot.close()
                time.sleep(3)

        else:
            # 1st Screenshot - AP Description
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\AP Description.png")
            time.sleep(5)

            # navigate to bottom element to take screenshot for attachments
            cc = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='itemsPP']")))
            a.move_to_element(cc).perform()
            time.sleep(5)

            # 2nd Screenshot - Attachments checking
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\Attachments Upload.png")
            time.sleep(5)

            # click on first acceptance object found
            d = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/"
                                                    "table[1]/tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/a[1]")))
            a.move_to_element(d).click().perform()
            self.bot.implicitly_wait(5)
            time.sleep(8)

            # 3rd Screenshot - MIN AP Info
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\Info Top Page.png")
            time.sleep(5)

            # search on Sharing Info - SIM Info
            dd = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Sharing Information']")))
            a.move_to_element(dd).perform()
            self.bot.implicitly_wait(30)
            time.sleep(8)

            # 4th Screenshot - SIM Info
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\SIM Info.png")
            time.sleep(5)

            # move to Name in Management System in order to ss bottom page
            e = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Name in Management System']")))
            a.move_to_element(e).perform()
            time.sleep(5)

            # 5th Screenshot - MIN AP Info
            self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\Info Bottom Page.png")
            time.sleep(5)

            # move to Name in Management System in order to ss bottom page
            ee = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Name in Management System']")))
            a.move_to_element(ee).perform()
            time.sleep(5)

            # Go back to the previous webpage
            self.bot.back()
            time.sleep(10)

            # check visibility of element MIN 2 node
            try:
                # Find the element with the given ID
                element = self.bot.find_element(By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[2]/div[1]/form[2]/table[1]/"
                                                        "tbody[1]/tr[1]/td[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/a[2]/span[1]")
            except:
                messagebox.showinfo('information',
                                    'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCR>MIN TEMPLATE folder.')
                self.bot.close()
                time.sleep(3)
            else:
                # Element exists, do something
                a.move_to_element(element).click().perform()
                time.sleep(5)

                # 6th Screenshot - MIN AP Info
                self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\Info Top Page.png")
                time.sleep(5)

                # search on Sharing Info - SIM Info
                ddd = WebDriverWait(self.bot, 60).until(
                    EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Sharing Information']")))
                a.move_to_element(ddd).perform()
                self.bot.implicitly_wait(30)
                time.sleep(5)

                # 7th Screenshot - SIM Info
                self.bot.save_screenshot("Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\SIM Info.png")
                time.sleep(5)

                # move to Name in Management System in order to ss bottom page
                g = WebDriverWait(self.bot, 60).until(
                    EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Name in Management System']")))
                a.move_to_element(g).perform()
                time.sleep(5)

                # 6th Screenshot - MIN AP Info
                self.bot.save_screenshot(
                    "Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\Info Bottom Page.png")
                time.sleep(5)

                messagebox.showinfo('information',
                                    'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCR>MIN TEMPLATE folder.')  # message box appear

                self.bot.close()
                time.sleep(3)

    def BIS_Template(self):

        self.bistemplate = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                  border_color="Gray", border_width=20, corner_radius=20)
        self.bistemplate.place(x=40, y=40)

        # Label
        heading = CTkLabel(self.bistemplate,
                           text='This template option enable user to create template consits of all the automation outputs.',
                           font=('Tahoma', 18, 'bold'))
        heading.place(x=50, y=40)
        # Label
        heading = CTkLabel(self.bistemplate,
                           text='User need to choose one or two AP template from dropdown and key-in all required items.',
                           font=('Tahoma', 18, 'bold'))
        heading.place(x=50, y=80)

        #global self.combo1
       #global self.combobox
        self.combo1 = StringVar()
        self.combobox = StringVar()
        # Label
        heading = CTkLabel(self.bistemplate, text='Choose AP Template', font=('Tahoma', 18, 'bold'))
        heading.place(x=100, y=200)
        #global combo1
        option2 = [" ", "Template 1", "Template 2", "Template 3", "Template 4", "Template 5", "Template 6"]
        self.combo1 = CTkOptionMenu(self.content_frame, values=option2)
        self.combo1.place(x=150, y=280)

        # Label
        heading = CTkLabel(self.bistemplate, text='Choose AP Template', font=('Tahoma', 18, 'bold'))
        heading.place(x=100, y=300)
        #global combobox
        options1 = [" ", "Template 1", "Template 2", "Template 3", "Template 4", "Template 5", "Template 6"]
        self.combobox = CTkOptionMenu(self.content_frame, values=options1)
        self.combobox.place(x=150, y=380)

        # Label
        heading = CTkLabel(self.bistemplate, text='Please key-in information below',
                           font=('Tahoma', 18, 'bold'))
        heading.place(x=500, y=200)

        self.ProjectCode = StringVar()
        self.CircuitID = StringVar()
        self.PreparedBy = StringVar()
        self.OtherInfo = StringVar()
        self.Department = StringVar()


        # Label
        heading = CTkLabel(self.bistemplate, text='Project Code', font=('Tahoma', 18))
        heading.place(x=500, y=250)
        # Label
        heading = CTkLabel(self.bistemplate, text='Circuit ID', font=('Tahoma', 18))
        heading.place(x=500, y=280)
        # Label
        heading = CTkLabel(self.bistemplate, text='Prepared by', font=('Tahoma', 18))
        heading.place(x=500, y=310)
        # Label
        heading = CTkLabel(self.bistemplate, text='Department', font=('Tahoma', 18))
        heading.place(x=500, y=340)
        # Label
        heading = CTkLabel(self.bistemplate, text='Other Info', font=('Tahoma', 18))
        heading.place(x=500, y=370)
        heading = CTkLabel(self.bistemplate, text='(SDWAN ID/Customer Name etc.)', font=('Tahoma', 13))
        heading.place(x=500, y=395)

        Entry1 = CTkEntry(self.bistemplate, textvariable=self.ProjectCode, font=("Tahoma", 19))
        Entry1.place(x=650, y=250)

        Entry2 = CTkEntry(self.bistemplate, textvariable=self.CircuitID, font=("Tahoma", 19))
        Entry2.place(x=650, y=280)

        Entry3 = CTkEntry(self.bistemplate, textvariable=self.PreparedBy, font=("Tahoma", 19))
        Entry3.place(x=650, y=310)

        Entry4 = CTkEntry(self.bistemplate, textvariable=self.Department, font=("Tahoma", 19))
        Entry4.place(x=650, y=340)

        Entry4 = CTkEntry(self.bistemplate, textvariable=self.OtherInfo, font=("Tahoma", 19))
        Entry4.place(x=650, y=370)

        CTkButton(self.bistemplate, text='Submit',command=self.create_templateBIS, width=100, height=30, 
                  border_width=0).place(x=470, y=500)
        CTkButton(self.bistemplate, text="Back", width=100, height=30, border_width=0,
        command=self.Automate).place(x=350, y=500)

    # BIS CHecklist Template
    def set_font(self, run, size, name, bold=False, italic=False, color=None):
        font = run.font
        font.size = Pt(size)
        font.name = name
        font.bold = bold
        font.italic = italic
        if color:
            font.color.rgb = color

    def create_templateBIS(self):
        messagebox.showinfo('information',
                    'Hi! Your BIS Checklist Template is currently being generated.Please wait for a moment')
        
        # Get user input
        self.ProjectCodeInput = self.ProjectCode.get()
        self.CircuitIDInput = self.CircuitID.get()
        self.UserName = self.PreparedBy.get()
        self.OtherInfoInput = self.OtherInfo.get()
        self.DepartmentInput = self.Department.get()
        self.Selected_Value = self.combo1.get()
        self.Selected_Value2 = self.combobox.get()
        # ---------------------------------------------First Page-----------------------------------------------------#
        # Create a new Microsoft Word document

        document = Document()

        # Add a header to the document
        header = document.sections[0].header

        # adjust the margin to be normal margin
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # Add a table to the header
        table = header.add_table(rows=1, cols=2, width=Cm(16))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'TableGrid'

        # Add some content to the table in header
        table.cell(0, 0).text = ("BIS CHECKLIST DOCUMENT")
        paragraph1 = table.cell(0, 0).paragraphs[0]
        paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tc = table.cell(0, 0).paragraphs[0].runs
        self.set_font(tc[0], 18, 'Times New Roman', bold=True, color=RGBColor(5, 132, 132))

        # add picture to table in header
        paragraph = table.cell(0, 1).paragraphs[0]
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture('images\\maxisheader.png')

        # add paragraph space
        document.add_paragraph().paragraph_format.line_spacing = 12

        # Title - Insert user key in Project Code/ AP Name
        para = document.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pararun = para.add_run(self.ProjectCodeInput)
        self.set_font(pararun, 28, 'Times New Roman', bold=True)

        # Title - Insert user key in Circuit ID
        para1 = document.add_paragraph()
        para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = para1.add_run(self.CircuitIDInput)
        self.set_font(run1, 28, 'Times New Roman', bold=True)

        # add paragraph space
        document.add_paragraph().paragraph_format.line_spacing = 10

        # Add table for user key in
        table1 = document.add_table(rows=5, cols=2)
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        table1.style = 'TableGrid'

        # Add some content to the table
        table1.cell(0, 0).text = ("Prepared By:\n")
        table1.cell(0, 1).text = (f"{self.UserName}\n")
        table1.cell(1, 0).text = ("AP Template:\n")
        table1.cell(1, 1).text = (f"{self.Selected_Value}\n{self.Selected_Value2}")
        table1.cell(2, 0).text = ("Other Info:\n")
        table1.cell(2, 1).text = (f"{self.OtherInfoInput}\n")
        table1.cell(3, 0).text = ("Department/Team:\n")
        table1.cell(3, 1).text = (f"{self.DepartmentInput}\n")
        table1.cell(4, 0).text = ("Date prepared:\n")

        global date 
        my_date = datetime.date.today()
        self.date = '{:%d-%b-%Y}'.format(my_date) # format date to be 'ex:12-Jan-2023'
        #date='{:%d-%b-%Y}'.format(date.today()) # format date to be 'ex:12-Jan-2023'
        table1.cell(4, 1).text = (f"{self.date}\n")

        
        for cell in table1.columns[0].cells:
            cell.width = Cm(5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell in table1.columns[1].cells:
            cell.width = Cm(16)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # adjust all font style for text in table
        for row in table1.rows:
            row.height = Cm(1)
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        self.set_font(run, 14, 'Times New Roman')

        #____________________________________________Page 3 & 4____________________________________________________#
        #------------------------------------------------NCR-------------------------------------------------------#
        document.add_page_break()
        document.add_paragraph()
        P1 = document.add_paragraph()
        P1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P1run = P1.add_run("Template Overview")
        self.set_font(P1run, 12, 'Times New Roman', bold=True)

        P2 = document.add_paragraph()
        P2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        P2run = P2.add_run("This template includes BIS Checklist items that are required for the acceptance of ETHD, SPM and TSTM processes.")
        self.set_font(P2run, 12, 'Times New Roman')

        P3 = document.add_paragraph()
        P3.style = 'Heading 1' # Set the style of the paragraph to Heading 1
        P3run = P3.add_run("1.0    Inventory - NCR\n")
        self.set_font(P3run, 12, 'Times New Roman', bold=True, color=RGBColor(0, 0, 0))

        P4 = document.add_paragraph()
        P4.paragraph_format.left_indent = Inches(0.37)
        P4.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P4run = P4.add_run("1.1    AP Description, Circuit Status & Circuit ID tagging\n")
        self.set_font(P4run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P5 = document.add_paragraph()
        P5run = P5.add_run("The AP description must contain Customer Name and Circuit ID. All required documents must be uploaded in the attachment section.")
        self.set_font(P5run, 12, 'Times New Roman')
        document.add_paragraph().paragraph_format.line_spacing = 1 

        # insert screenshot attachments from NCR automation
        P6 = document.add_paragraph()
        P6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP6 = P6.add_run()

        #path_dir_NCR = ""
        picture_path = "C:\\Automate\\Checklist Output\\NCR\\AP Description.png"
        print(picture_path)
        # Create a flag variable for each picture path
        picture_path_inserted = False

        # Loop through all paths
        for path in [picture_path]:
            if os.path.exists(path):
                print(f"{path} exists")
                if path == picture_path and not picture_path_inserted:
                    runP6.add_picture(picture_path, width=Cm(16), height=Cm(9))
                    runP6 = P6.add_run('\nAP Description in NCR\n')
                    picture_path_inserted = True       
            else:
                    runP6 = P6.add_run('Not Applicable')
                    self.set_font(runP6, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #--------------------------------------------------Summary Page---------------------------------------------#
        
        P7 = document.add_paragraph()
        P7.paragraph_format.left_indent = Inches(0.37)
        P7.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P7run = P7.add_run("1.2    Summary Page Information\n")
        self.set_font(P7run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P8 = document.add_paragraph()
        P8run = P8.add_run("All required Summary Page sections must be filled with valid information. Network Diagram file must be inserted at its section.")
        self.set_font(P8run, 12, 'Times New Roman')

        #---------------------------------------------Summary Page--------------------------------------------------#
        # insert screenshot attachments from NCR automation
        P9 = document.add_paragraph()
        P9.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP9 = P9.add_run()

        PP3 = "C:\\Automate\\Checklist Output\\NCR\\Summary Page1.png"
        PP4 = "C:\\Automate\\Checklist Output\\NCR\\Summary Page2.png"

        PP3_inserted = False
        PP4_inserted = False

        for path in [PP3, PP4]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check if both picture files exist
                if path == PP3 and not PP3_inserted:
                    runP9.add_picture(PP3, width=Cm(16), height=Cm(9))
                    runP9 = P9.add_run('\nSummary Page\n')
                    PP3_inserted = True
                elif path == PP4 and not PP4_inserted:
                    # Insert the picture into the document
                    runP9.add_picture(PP4, width=Cm(16), height=Cm(9))
                    runP9 = P9.add_run('\nSummary Page')
                    PP4_inserted = True
            else:
                print(f"{path} does not exist")

        if all([not PP3_inserted, not PP4_inserted]):
            # If neither picture file exists, insert 'Not Applicable'
            runP9 = P9.add_run('Not Applicable')
            self.set_font(runP9, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #--------------------------------------------Network Diagram-------------------------------------------------#
        document.add_paragraph().paragraph_format.line_spacing = 1 #CHECK SPACE!!!!!!!!!!
        # Visio diagram from Work Order file
        P11 = document.add_paragraph()
        P11.paragraph_format.left_indent = Inches(0.37)
        P11.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P11run = P11.add_run("1.3    Network Diagram (Visio)\n")
        self.set_font(P11run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P12 = document.add_paragraph()
        P12run = P12.add_run("Visio diagram must be insert in the Network Diagram file.\n")
        self.set_font(P12run, 12, 'Times New Roman')

        latest_NDfile1 = ''
        try:
            latest_NDfile1 = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\NCR", "*NWO*")) + glob.glob(os.path.join("C:\\Automate\\Checklist Output\\NCR", "*MWO*")), key=os.path.getctime)
        except ValueError:
            print("No Network Diagram file found")

        if os.path.exists(latest_NDfile1):
            print(f"{latest_NDfile1} exists")
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NDfile1} not exists")
            P121 = document.add_paragraph()
            P121.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runP121 = P121.add_run('Not Applicable')
            self.set_font(runP121, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))
        
        if os.path.exists(latest_NDfile1):
            P122 = document.add_paragraph()
            P122.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P122run = P122.add_run("\nNetwork Diagram File\n")
            self.set_font(P122run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NDfile1} not exists")

    #-----------------------------------------------Network Element NCR------------------------------------------#
        #document.add_paragraph('\n').paragraph_format.line_spacing = 2
        P13 = document.add_paragraph()
        P13.paragraph_format.left_indent = Inches(0.37)
        P13.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P13run = P13.add_run("\n1.4    Network Element - Circuit ID & Parameters\n")
        self.set_font(P13run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P14 = document.add_paragraph()
        P14run = P14.add_run("To ensure that all required equipment has the appropriate Circuit ID tagging, the Description section must include the Circuit ID.")
        self.set_font(P14run, 12, 'Times New Roman')

        P15 = document.add_paragraph()
        P15.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP15 = P15.add_run()

        PP5 = "C:\\Automate\\Checklist Output\\NCR\\Main Tab.png"
        PP6 = "C:\\Automate\\Checklist Output\\NCR\\Parameters Tab.png"

        # Create a flag variable for each picture path
        PP5_inserted = False
        PP6_inserted = False

        # Loop through all paths
        for path in [PP5, PP6]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check if the picture file exists
                if path == PP5 and not PP5_inserted:
                    # Insert the picture into the document
                    runP15.add_picture(PP5, width=Cm(16), height=Cm(9))
                    runP15 = P15.add_run('\nCircuit ID Tagging on all Nodes/Equipments\n')
                    PP5_inserted = True
                # Check if the picture file exists
                elif path == PP6 and not PP6_inserted:
                    # Insert the picture into the document
                    runP15.add_picture(PP6, width=Cm(16), height=Cm(9))
                    runP15 = P15.add_run('\nParameters Tab')
                    PP6_inserted = True
            else:
                    print('Picture not exist')

        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not PP5_inserted, not PP6_inserted]):
            runP15 = P15.add_run('Not Applicable')
            self.set_font(runP15, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))


        #-------------------------------------------------Path Element NCR----------------------------------------#
        P17 = document.add_paragraph()
        P17.paragraph_format.left_indent = Inches(0.37)
        P17.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P17run = P17.add_run("1.5    Path Element\n")
        self.set_font(P17run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P18 = document.add_paragraph()
        P18run = P18.add_run("Order of equipment’s in Path Element must in sequent and tally from maxis cloud/customer end.\nAll port up link/ downlink must match with labelling in devices.")
        self.set_font(P18run, 12, 'Times New Roman')
        #document.add_paragraph().paragraph_format.line_spacing = 1 

        P19 = document.add_paragraph()
        P19.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP19 = P19.add_run()

        PP7 = "C:\\Automate\\Checklist Output\\NCR\\Path Element.png"

        # Check if the picture file exists
        if os.path.exists(PP7):
            # Insert the picture into the document
            runP19.add_picture(PP7, width=Cm(16), height=Cm(9))
            runP19 = P19.add_run('\nPath Element in NCR')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            runP19 = P19.add_run('Not Applicable')
            self.set_font(runP19, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

    
        #----------------------------------------------NCR MIN IOT-------------------------------------------------#
        #document.add_page_break()
        #document.add_paragraph().paragraph_format.line_spacing = 1 
        
        P20 = document.add_paragraph()
        P20.style = 'Heading 1' # Set the style of the paragraph to Heading 1
        P20run = P20.add_run("2.0    Inventory - NCR for AP MIN IOT Template\n")
        self.set_font(P20run, 12, 'Times New Roman', bold=True, color=RGBColor(0, 0, 0))
        
        #-------------------------------------AP Description & Attachments Upload----------------------------------#
        P21 = document.add_paragraph()
        P21.paragraph_format.left_indent = Inches(0.37)
        P21.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P21run = P21.add_run("2.1    AP Description & Attachments Upload\n")
        self.set_font(P21run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P23 = document.add_paragraph()
        P23.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP23 = P23.add_run()

        PP8="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\AP Description.png"
        # Create a flag variable for each picture path
        PP8_inserted = False

        for path in [PP8]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check if the picture file exists
                if path == PP8 and not PP8_inserted:
                    # Insert the picture into the document
                    runP23.add_picture(PP8, width=Cm(16), height=Cm(9))
                    runP23 = P23.add_run('\nAP Description\n')
                    PP8_inserted = True             
            else:
                print(f"{path} does not exist")
        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not PP8_inserted]):
            runP23 = P23.add_run('Not Applicable')
            self.set_font(runP23, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))
        #-----------------------------------------------MIN Product 1----------------------------------------------#
        P24 = document.add_paragraph()
        P24.paragraph_format.left_indent = Inches(0.37)
        P24.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P24run = P24.add_run("2.2    MIN Product 1\n")
        self.set_font(P24run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P25 = document.add_paragraph()
        P25.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP25 = P25.add_run()

        # Picture path = screenshots from BIS Automation
        PP10="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\Info Top Page.png"
        PP11="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\SIM Info.png"
        PP12="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 1\\Info Bottom Page.png"
        # Create a flag variable for each picture path
        PP10_inserted = False
        PP11_inserted = False
        PP12_inserted = False

        for path in [PP10, PP11, PP12]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check if the picture file exists
                if path == PP10 and not PP10_inserted:
                    # Insert the picture into the document
                    runP25.add_picture(PP10, width=Cm(16), height=Cm(9))
                    runP25 = P25.add_run('\nInfo Top Page\n')
                    PP10_inserted = True
                elif path == PP11 and not PP11_inserted:
                    # Insert the picture into the document
                    runP25.add_picture(PP11, width=Cm(16), height=Cm(9))
                    runP25 = P25.add_run('\nSIM Info\n')
                    PP11_inserted = True
                elif path == PP12 and not PP12_inserted:
                    # Insert the picture into the document
                    runP25.add_picture(PP12, width=Cm(16), height=Cm(9))
                    runP25 = P25.add_run('\nInfo Bottom Page')
                    PP12_inserted = True
            else:
                print(f"{path} does not exist")
        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not PP10_inserted, not PP11_inserted, not PP12_inserted]):
            runP25 = P25.add_run('Not Applicable')
            self.set_font(runP25, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #-----------------------------------------------MIN Product 2----------------------------------------------#
        P27 = document.add_paragraph()
        P27.paragraph_format.left_indent = Inches(0.37)
        P27.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P27run = P27.add_run("2.2    MIN Product 2\n")
        self.set_font(P27run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P28 = document.add_paragraph()
        P28.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP28 = P28.add_run()

        # Picture path = screenshots from BIS Automation
        PP13="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\Info Top Page.png"
        PP14="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\SIM Info.png"
        PP15="C:\\Automate\\Checklist Output\\NCR-MIN TEMPLATE\\MIN Acceptance Object 2\\Info Bottom Page.png"

        # Create a flag variable for each picture path
        PP13_inserted = False
        PP14_inserted = False
        PP15_inserted = False

        for path in [PP13, PP14, PP15]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check if the picture file exists
                if path == PP13 and not PP13_inserted:
                    # Insert the picture into the document
                    runP28.add_picture(PP13, width=Cm(16), height=Cm(9))
                    runP28 = P28.add_run('\nInfo Top Page\n')
                    PP13_inserted = True
                elif path == PP14 and not PP14_inserted:
                    # Insert the picture into the document
                    runP28.add_picture(PP14, width=Cm(16), height=Cm(9))
                    runP28 = P28.add_run('\nSIM Info\n')
                    PP14_inserted = True
                elif path == PP15 and not PP15_inserted:
                    # Insert the picture into the document
                    runP28.add_picture(PP15, width=Cm(16), height=Cm(9))
                    runP28 = P28.add_run('\nInfo Bottom Page')
                    PP15_inserted = True
            else:
                print(f"{path} does not exist")
        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not PP13_inserted, not PP14_inserted, not PP15_inserted]):
            runP28 = P28.add_run('Not Applicable')
            self.set_font(runP28, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #-----------------------------------------------EM Performance---------------------------------------------#
        P29 = document.add_paragraph()
        P29.style = 'Heading 1' # Set the style of the paragraph to Heading 1
        P29run = P29.add_run("3.0    Element Manager (EM) - Performance")
        self.set_font(P29run, 12, 'Times New Roman', bold=True, color=RGBColor(0, 0, 0))
        #------------------------------------------------CNMAESTRO--------------------------------------------------#
        P30 = document.add_paragraph()
        P30.paragraph_format.left_indent = Inches(0.37)
        P30.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P30run = P30.add_run("3.1    CNMAESTRO (FWA Equipment)\n")
        self.set_font(P30run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P31 = document.add_paragraph()
        P31.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP31 = P31.add_run()

        # Picture path = screenshots from BIS Automation
        PP16="C:\\Automate\\Checklist Output\\CNMAESTRO\\Dashboard1.png"
        PP17="C:\\Automate\\Checklist Output\\CNMAESTRO\\Dashboard2.png"
        PP18="C:\\Automate\\Checklist Output\\CNMAESTRO\\Performance1.png"
        PP19="C:\\Automate\\Checklist Output\\CNMAESTRO\\Performance2.png"

        # Create a flag variable for each picture path
        PP16_inserted = False
        PP17_inserted = False
        PP18_inserted = False
        PP19_inserted = False

        for path in [PP16, PP17, PP18, PP19]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check if the picture file exists
                if path == PP16 and not PP16_inserted:
                    # Insert the picture into the document
                    runP31.add_picture(PP16, width=Cm(16), height=Cm(9))
                    runP31 = P31.add_run('\nDashboard1\n')
                    PP16_inserted = True
                elif path == PP17 and not PP17_inserted:
                    # Insert the picture into the document
                    runP31.add_picture(PP17, width=Cm(16), height=Cm(9))
                    runP31 = P31.add_run('\nDashboard2\n')
                    PP17_inserted = True
                elif path == PP18 and not PP18_inserted:
                    # Insert the picture into the document
                    runP31.add_picture(PP18, width=Cm(16), height=Cm(9))
                    runP31 = P31.add_run('\nPerformance1\n')
                    PP18_inserted = True
                elif path == PP19 and not PP19_inserted:
                    # Insert the picture into the document
                    runP31.add_picture(PP19, width=Cm(16), height=Cm(9))
                    runP31 = P31.add_run('\nPerformance2')
                    PP19_inserted = True
            else:
                print(f"{path} does not exist")
        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not PP16_inserted, not PP17_inserted, not PP18_inserted, not PP19_inserted]):
            runP31 = P31.add_run('Not Applicable')
            self.set_font(runP31, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #-----------------------------------------------BREEZE AIR-------------------------------------------------#
        P32 = document.add_paragraph()
        P32.paragraph_format.left_indent = Inches(0.37)
        P32.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P32run = P32.add_run("3.2    BREEZE AIR (NLOS Equipment)\n")
        self.set_font(P32run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        latest_NLOSfile1 = ''
        try:
            latest_NLOSfile1 = max(glob.glob(os.path.join("C:\\Automate\\Checklist Output\\BREEZE AIR", "radios*")), key=os.path.getctime)
        except ValueError:
            print("No Network Diagram file found")

        if os.path.exists(latest_NLOSfile1):
            print(f"{latest_NLOSfile1} exists")
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NLOSfile1} not exists")
            P321 = document.add_paragraph()
            P321.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runP321 = P321.add_run('Not Applicable')
            self.set_font(runP321, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        if os.path.exists(latest_NLOSfile1):
            P322 = document.add_paragraph()
            P322.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P322run = P322.add_run("\nNLOS FWA File\n")
            self.set_font(P322run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NLOSfile1} not exists")

        #-----------------------------------------------EXFO-------------------------------------------------#
        P33 = document.add_paragraph()
        P33.paragraph_format.left_indent = Inches(0.37)
        P33.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        runP33 = P33.add_run("3.3    EXFO (CPE/OZRO/NID/ETHERNET SWITCH Equipment)\n")
        self.set_font(runP33, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P34 = document.add_paragraph()
        PP34run = P34.add_run("Interface description tally with format labelling document and no alarm at logical physical interface.\n")
        self.set_font(PP34run, 12, 'Times New Roman')

        path_dirDF = "C:\\Automate\\Checklist Output\\EXFO"
        latest_EXFOfile1 = ''

        try:
            latest_EXFOfile1 = max(glob.glob(os.path.join(path_dirDF, "*monitors*")), key=os.path.getctime)
        except ValueError:
            print("No monitors file found")

        if os.path.exists(latest_EXFOfile1):
            print(f"{latest_EXFOfile1} exists")
            #document.add_paragraph('\n')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_EXFOfile1} not exists")
            P341 = document.add_paragraph()
            P341.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runP341 = P341.add_run('Not Applicable')
            self.set_font(runP341, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        if os.path.exists(latest_EXFOfile1):
            P342 = document.add_paragraph()
            P342.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P342run = P342.add_run("\nEXFO Files\n")
            self.set_font(P342run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_EXFOfile1} not exists")
        
        EXFOsswarning = ''
        try:
            EXFOsswarning = max(glob.glob(os.path.join(path_dirDF, "*Warning Alert_*")), key=os.path.getctime)
        except ValueError:
            print("No Warning Alert screenshot found")

        if os.path.exists(EXFOsswarning):
            print(f"{EXFOsswarning} exists")
            P3423 = document.add_paragraph()
            P3423.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P3423run = P3423.add_run()
            P3423run.add_picture(EXFOsswarning, width=Cm(16), height=Cm(9))
            P3423run = P3423.add_run('\nWarning Alert')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            print('NO warning alert ss')

        #--------------------------------------------------NCM-----------------------------------------------------#
        P35 = document.add_paragraph()
        P35.paragraph_format.left_indent = Inches(0.37)
        P35.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        runP35 = P35.add_run("3.4    NCM SOLARWIND (CPE/OZRO Equipment)\n")
        self.set_font(runP35, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P36 = document.add_paragraph()
        P36run = P36.add_run("Node is visualized in NCM SolarWinds")
        self.set_font(P36run, 12, 'Times New Roman')

        P361 = document.add_paragraph()
        P361.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P361run = P361.add_run()
        self.set_font(P361run, 12, 'Times New Roman')

        path_dirNCM = "C:\\Automate\\Checklist Output\\NCM"
        fileNCM = glob.glob(os.path.join(path_dirNCM, "NCM*"))
        sorted_fileNCM = sorted(fileNCM, key=os.path.getctime, reverse=False)
        latest_NCMfile = None
        second_latest_NCMfile = None

        if len(sorted_fileNCM) > 0:
            # If there are two or more EXFO files, insert the latest two
            if len(sorted_fileNCM) >= 2:
                latest_NCMfile = sorted_fileNCM[0]
                second_latest_NCMfile = sorted_fileNCM[1]
            # If there is only one EXFO file, insert it
            else:
                latest_NCMfile = sorted_fileNCM[0]
                second_latest_NCMfile = None
        else:
            # Handle the case where the directory is empty
            print("Error: the NCM directory is empty")

        # Check if the picture file exists
        if latest_NCMfile is not None and os.path.exists(latest_NCMfile):
            # Insert the picture into the document
            P361run.add_picture(latest_NCMfile, width=Cm(16), height=Cm(9))
            P361run = P361.add_run('\nNCM Visibility\n')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NCMfile} not exists")

        if second_latest_NCMfile is not None and os.path.exists(second_latest_NCMfile):
            # Insert the picture into the document
            P361run.add_picture(second_latest_NCMfile, width=Cm(16), height=Cm(9))
            P361run = P361.add_run('\nNCM Visibility\n')
        elif len(sorted_fileNCM) > 1:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{sorted_fileNCM} not exists")
        #Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not latest_NCMfile, not second_latest_NCMfile]):
            P361run = P361.add_run('Not Applicable')
            self.set_font(P361run, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #-------------------------------------------------MERAKI----------------------------------------------------#
        P37 = document.add_paragraph()
        P37.paragraph_format.left_indent = Inches(0.37)
        P37.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P37run = P37.add_run("3.5    MERAKI (SDWAN Equipment)\n")
        self.set_font(P37run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P38 = document.add_paragraph()
        P38.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP38 = P38.add_run()

        # Picture path = screenshots from BIS Automation
        PP23="C:\\Automate\\Checklist Output\\MERAKI\\DAP Device Connectivity.png"
        PP24="C:\\Automate\\Checklist Output\\MERAKI\\DSW Device Connectivity.png"
        PP25="C:\\Automate\\Checklist Output\\MERAKI\\MerakiTags.png"
        PP26="C:\\Automate\\Checklist Output\\MERAKI\\DRO Device Configuration.png"
        PP27="C:\\Automate\\Checklist Output\\MERAKI\\DRO Device Live Data.png"
        PP27_DCG1="C:\\Automate\\Checklist Output\\MERAKI\\DCG Device.png"
        PP27_DCG2="C:\\Automate\\Checklist Output\\MERAKI\\DCG Device Configuration.png"
        PP27_DCG3="C:\\Automate\\Checklist Output\\MERAKI\\DCG Device Live Data.png"
        PP27_DCG4="C:\\Automate\\Checklist Output\\MERAKI\\DCG Device Latency & Loss.png"

        # Create a flag variable for each picture path
        PP23_inserted = False
        PP24_inserted = False
        PP25_inserted = False
        PP26_inserted = False
        PP27_inserted = False
        PP27_DCG1_inserted = False
        PP27_DCG2_inserted = False
        PP27_DCG3_inserted = False
        PP27_DCG4_inserted = False

        # Loop through all paths
        for path in [PP23, PP24, PP25, PP26, PP27, PP27_DCG1, PP27_DCG2, PP27_DCG3, PP27_DCG4]:
            if os.path.exists(path):
                print(f"{path} exists")
                # Check which path it is, and whether it has already been inserted
                if path == PP23 and not PP23_inserted:
                    runP38.add_picture(PP23, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nAccess Points\n')
                    PP23_inserted = True
                elif path == PP24 and not PP24_inserted:
                    runP38.add_picture(PP24, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nSwitches\n')
                    PP24_inserted = True
                elif path == PP25 and not PP25_inserted:
                    runP38.add_picture(PP25, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nInterface Description\n')
                    PP25_inserted = True
                elif path == PP26 and not PP26_inserted:
                    runP38.add_picture(PP26, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nDRO Device Configuration\n')
                    PP26_inserted = True
                elif path == PP27 and not PP27_inserted:
                    runP38.add_picture(PP27, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nDRO Device Live Data\n')
                    PP27_inserted = True
                elif path == PP27_DCG1 and not PP27_DCG1_inserted:
                    runP38.add_picture(PP27_DCG1, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nDCG Device\n')
                    PP27_DCG1_inserted = True
                elif path == PP27_DCG2 and not PP27_DCG2_inserted:
                    runP38.add_picture(PP27_DCG2, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nDCG Device Configuration\n')
                    PP27_DCG2_inserted = True
                elif path == PP27_DCG3 and not PP27_DCG3_inserted:
                    runP38.add_picture(PP27_DCG3, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nDCG Device Live Data\n')
                    PP27_DCG3_inserted = True
                elif path == PP27_DCG4 and not PP27_DCG4_inserted:
                    runP38.add_picture(PP27_DCG4, width=Cm(16), height=Cm(9))
                    runP38 = P38.add_run('\nDCG Device Latency & Loss\n')
                    PP27_DCG4_inserted = True
            else:
                print(f"{path} does not exist")

        # Check whether all pictures were not applicable, and insert "Not Applicable" if necessary
        if all([not PP23_inserted, not PP24_inserted, not PP25_inserted, not PP26_inserted, not PP27_inserted, not PP27_DCG1_inserted, not PP27_DCG2_inserted,
                not PP27_DCG3_inserted, not PP27_DCG4_inserted]):
            document.add_paragraph()
            runP38 = P38.add_run('Not Applicable')
            self.set_font(runP38, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))


        #--------------------------------------------------NCE-IP---------------------------------------------------#
        P39 = document.add_paragraph()
        P39.paragraph_format.left_indent = Inches(0.37)
        P39.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        runP39 = P39.add_run("3.6    NCE-IP (POC3/POC2/ SPE NGBB)\n")
        self.set_font(runP39, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P40 = document.add_paragraph()
        P40run = P40.add_run("Node is visualized in NCE-IP\n")
        self.set_font(P40run, 12, 'Times New Roman')

        P401 = document.add_paragraph()
        P401.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP401 = P401.add_run()

        PP28="C:\\Automate\\Checklist Output\\NCE-IP\\Node Visibility.png"
        if os.path.exists(PP28):
            print(f"{PP28} exists")
            runP401.add_picture(PP28, width=Cm(16), height=Cm(9))
            runP401 = P401.add_run('\nPerformance')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            runP401 = P401.add_run('Not Applicable')
            self.set_font(runP401, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        path_dirNCEIPCONFIG = "C:\\Automate\\Checklist Output\\CONFIG\\POC3POC2POC1"
        latest_NCEIPCONFIG = ''

        try:
            latest_NCEIPCONFIG = max(glob.glob(os.path.join(path_dirNCEIPCONFIG, "*ConfigOutputNCEIP*")), key=os.path.getctime)
        except ValueError:
            print("No monitors file found")

        if os.path.exists(latest_NCEIPCONFIG):
            print(f"{latest_NCEIPCONFIG} exists")
            #document.add_paragraph('\n')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NCEIPCONFIG} not exists")
            P402 = document.add_paragraph()
            P402.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runP402 = P402.add_run('Not Applicable')
            self.set_font(runP402, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        if os.path.exists(latest_NCEIPCONFIG):
            P402 = document.add_paragraph()
            P402.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P402run = P402.add_run("\nConfig Description NCE-IP\n")
            self.set_font(P402run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_NCEIPCONFIG} not exists")
        #--------------------------------------------------NCE-FAN--------------------------------------------------#
        P41 = document.add_paragraph()
        P41.paragraph_format.left_indent = Inches(0.37)
        P41.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        runP41 = P41.add_run("3.7    NCE-FAN (MXU/OLT)\n")
        self.set_font(runP41, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P42 = document.add_paragraph()
        P42run = P42.add_run("Node is visualized in NCE-FAN\n")
        self.set_font(P42run, 12, 'Times New Roman')

        P421 = document.add_paragraph()
        P421.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP421 = P421.add_run()

        PP29="C:\\Automate\\Checklist Output\\NCE-FAN\\Node Visibility.png"
        if os.path.exists(PP29):
            print(f"{PP29} exists")
            runP421.add_picture(PP29, width=Cm(16), height=Cm(9))
            runP421 = P421.add_run('\nPerformance')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            runP421 = P421.add_run('Not Applicable')
            self.set_font(runP421, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #-------------------------------------------------PRTG----------------------------------------------------#
        P43 = document.add_paragraph()
        P43.paragraph_format.left_indent = Inches(0.37)
        P43.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        P43run = P43.add_run("3.8    PRTG (MIN Equipment)\n")
        self.set_font(P43run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P44 = document.add_paragraph()
        P44.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP44 = P44.add_run()

        path_dirMIN = "C:\\Automate\\Checklist Output\\PRTG MIN"
        fileMIN = glob.glob(os.path.join(path_dirMIN, "MIN*"))
        sorted_fileMIN = sorted(fileMIN, key=os.path.getctime, reverse=False)
        latest_MINfile = None
        second_latest_MINfile = None
        third_latest_MINfile = None
        fourth_latest_MINfile = None

        if len(sorted_fileMIN) > 0:
            # If there are two or more EXFO files, insert the latest two
            if len(sorted_fileMIN) >= 2:
                latest_MINfile = sorted_fileMIN[0]
                second_latest_MINfile = sorted_fileMIN[1]
                third_latest_MINfile = sorted_fileMIN[2] if len(sorted_fileMIN) >= 3 else None
                fourth_latest_MINfile = sorted_fileMIN[3] if len(sorted_fileMIN) >= 4 else None
            # If there is only one EXFO file, insert it
            else:
                latest_MINfile = sorted_fileMIN[0]
                second_latest_MINfile = None
                third_latest_MINfile = None
                fourth_latest_MINfile = None
        else:
            # Handle the case where the directory is empty
            print("Error: the PRTG MIN directory is empty")

        # Check if the picture file exists
        if latest_MINfile is not None and os.path.exists(latest_MINfile):
            # Insert the picture into the document
            runP44.add_picture(latest_MINfile, width=Cm(16), height=Cm(9))
            runP44 = P44.add_run('\nPRTG MIN SENSOR\n')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_MINfile} not exists")

        if second_latest_MINfile is not None and os.path.exists(second_latest_MINfile):
            # Insert the picture into the document
            runP44.add_picture(second_latest_MINfile, width=Cm(16), height=Cm(9))
            runP44 = P44.add_run('\nPRTG MIN SENSOR\n')
        elif len(sorted_fileMIN) > 1:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{second_latest_MINfile} not exists")

        if third_latest_MINfile is not None and os.path.exists(third_latest_MINfile):
            # Insert the picture into the document
            runP44.add_picture(third_latest_MINfile, width=Cm(16), height=Cm(9))
            runP44 = P44.add_run('\nPRTG MIN SENSOR\n')
        elif len(sorted_fileMIN) > 2:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{third_latest_MINfile} not exists")

        if fourth_latest_MINfile is not None and os.path.exists(fourth_latest_MINfile):
            # Insert the picture into the document
            runP44.add_picture(fourth_latest_MINfile, width=Cm(16), height=Cm(9))
            runP44 = P44.add_run('\nPRTG MIN SENSOR')
        elif len(sorted_fileMIN) > 3:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{fourth_latest_MINfile} not exists")

        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not latest_MINfile, not second_latest_MINfile, not third_latest_MINfile, not fourth_latest_MINfile]):
            runP44 = P44.add_run('Not Applicable')
            self.set_font(runP44, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        #--------------------------------------------------PDH MW--------------------------------------------------#
        P45 = document.add_paragraph()
        P45.paragraph_format.left_indent = Inches(0.37)
        P45.style = 'Heading 2' # Set the style of the paragraph to Heading 2
        runP45 = P45.add_run("3.9    PDH (MW Equipment)\n")
        self.set_font(runP45, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        P46 = document.add_paragraph()
        P46run = P46.add_run("Node is visualized in iPASOLINK")
        self.set_font(P46run, 12, 'Times New Roman')

        P461 = document.add_paragraph()
        P461.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runP461 = P461.add_run()

        path_dirPDH = "C:\\Automate\\Checklist Output\\PDH MW"
        filePDH = glob.glob(os.path.join(path_dirPDH, "PDH*"))
        sorted_filePDH = sorted(filePDH, key=os.path.getctime, reverse=False)
        
        latest_PDHfile = None
        second_latest_PDHfile = None
        third_latest_PDHfile = None
        fourth_latest_PDHfile = None

        if len(sorted_filePDH) > 0:
            # If there are two or more EXFO files, insert the latest two
            if len(sorted_filePDH) >= 2:
                latest_PDHfile = sorted_filePDH[0]
                second_latest_PDHfile = sorted_filePDH[1]
                third_latest_PDHfile = sorted_filePDH[2] if len(sorted_filePDH) >= 3 else None
                fourth_latest_PDHfile = sorted_filePDH[3] if len(sorted_filePDH) >= 4 else None
            # If there is only one EXFO file, insert it
            else:
                latest_PDHfile = sorted_filePDH[0]
                second_latest_PDHfile = None
                third_latest_PDHfile = None
                fourth_latest_PDHfile = None
        else:
            # Handle the case where the directory is empty
            print("Error: the PDH MW directory is empty")

        # Check if the picture file exists
        if latest_PDHfile is not None and os.path.exists(latest_PDHfile):
            # Insert the picture into the document
            runP461.add_picture(latest_PDHfile, width=Cm(16), height=Cm(9))
            runP461 = P461.add_run('\nPDH MW\n')
        else:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_PDHfile} not exists")

        if second_latest_PDHfile is not None and os.path.exists(second_latest_PDHfile):
            # Insert the picture into the document
            runP461.add_picture(second_latest_PDHfile, width=Cm(16), height=Cm(9))
            runP461 = P461.add_run('\nPDH MW\n')
        elif len(sorted_filePDH) > 1:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{second_latest_PDHfile} not exists")

        if third_latest_PDHfile is not None and os.path.exists(third_latest_PDHfile):
            # Insert the picture into the document
            runP461.add_picture(third_latest_PDHfile, width=Cm(16), height=Cm(9))
            runP461 = P461.add_run('\nPDH MW\n')
        elif len(sorted_filePDH) > 2:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{third_latest_PDHfile} not exists")

        if fourth_latest_PDHfile is not None and os.path.exists(fourth_latest_PDHfile):
            # Insert the picture into the document
            runP461.add_picture(fourth_latest_PDHfile, width=Cm(16), height=Cm(9))
            runP461 = P461.add_run('\nPDH MW')
        elif len(sorted_filePDH) > 3:
            # If the picture file does not exist, insert 'Not Applicable'
            print(f"{fourth_latest_PDHfile} not exists")

        # Check the flag variable after the loop, and insert 'Not Applicable' if it is still True
        if all([not latest_PDHfile, not second_latest_PDHfile, not third_latest_PDHfile, not fourth_latest_PDHfile]):
            runP461 = P461.add_run('Not Applicable')
            self.set_font(runP461, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))    

        #------------------------------------------------Config-----------------------------------------------------#
        P47 = document.add_paragraph()
        P47.style = 'Heading 1' # Set the style of the paragraph to Heading 1
        P47run = P47.add_run("4.0    Configuration File ")
        self.set_font(P47run, 12, 'Times New Roman', bold=True, color=RGBColor(0, 0, 0))

        P48 = document.add_paragraph()
        P48run = P48.add_run("Please refer to file below")
        self.set_font(P48run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))
        document.add_paragraph().paragraph_format.line_spacing = 1 

        path_dirConfig = "C:\\Automate\\Checklist Output\\CONFIG"
        latest_Configfile1 = ''

        try:
            latest_Configfile1 = max(glob.glob(os.path.join(path_dirConfig, "ConfigOutput*")), key=os.path.getctime)
        except ValueError:
            print("No monitors file found")

        if os.path.exists(latest_Configfile1):
            print(f"{latest_Configfile1} exists")
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_Configfile1} not exists")
            P481 = document.add_paragraph()
            P481.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runP481 = P481.add_run('Not Applicable')
            self.set_font(runP481, 12, 'Times New Roman',bold=True, color=RGBColor(255, 0, 0))

        if os.path.exists(latest_Configfile1):
            P482 = document.add_paragraph()
            P482.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P482run = P482.add_run("\nConfig Files\n")
            self.set_font(P482run, 12, 'Times New Roman')
        else:
        # If the picture file does not exist, insert 'Not Applicable'
            print(f"{latest_Configfile1} not exists")
        #------------------------------------------------Other Info------------------------------------------------#
        P47 = document.add_paragraph()
        P47.style = 'Heading 1' # Set the style of the paragraph to Heading 1
        P47run = P47.add_run("5.0    Other Information")
        self.set_font(P47run, 12, 'Times New Roman', bold=True, color=RGBColor(0, 0, 0))

        P48 = document.add_paragraph()
        P48run = P48.add_run("Please add any additional information here")
        self.set_font(P48run, 12, 'Times New Roman', color=RGBColor(0, 0, 0))

        #--------------------------------------------Save Document-------------------------------------------------#
        # Save the document
        document.save('C:\\Automate\\Checklist Output\\BIS Checklist Template.docx')

        #------------------------------------------find latest NWO/ND file-----------------------------------------#
        path_dir_EXFO = "C:\\Automate\\Checklist Output\\EXFO"
        path_dir_ND = "C:\\Automate\\Checklist Output\\NCR"
        path_dir_NLOS = "C:\\Automate\\Checklist Output\\BREEZE AIR"
        path_dir2 = "C:\\Automate\\Checklist Output\\CONFIG"
        path_NCEIP = "C:\\Automate\\Checklist Output\\CONFIG\\POC3POC2POC1"

        # Find the latest file starting with 'NWO'
        latest_NDfile = ''
        latest_NLOSfile = ''
        latest_EXFOfile = ''
        latest_Configfile = ''
        latest_ConfigNCEIP = ''
        second_latest_fileEXFO = None 
        second_latest_Configfile = None
        second_latest_ConfigNCEIP = None

        #------------------------------------Add file in BIS Template----------------------------------------------#
        time.sleep(5)
        word = win32com.client.gencache.EnsureDispatch('Word.Application')
        docpath = os.path.abspath("C:\\Automate\\Checklist Output\\BIS Checklist Template.docx")
        # create a new Word document
        document = word.Documents.Open(docpath)
        #document1 = Document(docpath)
        # Find the latest Network Diagram file
        try:
            latest_NDfile = max(glob.glob(os.path.join(path_dir_ND, "*NWO*")) + glob.glob(os.path.join(path_dir_ND, "*MWO*")), key=os.path.getctime)
            latest_NLOSfile = max(glob.glob(os.path.join(path_dir_NLOS, "radios*")), key=os.path.getctime)
        except ValueError:
            print("Files not found")

        fileNCEIP = glob.glob(os.path.join(path_NCEIP, "*ConfigOutputNCEIP*"))
        sorted_fileNCEIP = sorted(fileNCEIP, key=os.path.getctime, reverse=False)

        if len(sorted_fileNCEIP) > 0:
            # If there are two or more NCE_IP files, insert the latest two
            if len(sorted_fileNCEIP) >= 2:
                latest_ConfigNCEIP = sorted_fileNCEIP[0]
                second_latest_ConfigNCEIP = sorted_fileNCEIP[1]
                third_latest_ConfigNCEIP = sorted_fileNCEIP[2] if len(sorted_fileNCEIP) >= 3 else None
                fourth_latest_ConfigNCEIP = sorted_fileNCEIP[3] if len(sorted_fileNCEIP) >= 4 else None
            # If there is only one NCE-IP file, insert it
            else:
                latest_ConfigNCEIP = sorted_fileNCEIP[0]
                second_latest_ConfigNCEIP = None
                third_latest_ConfigNCEIP = None
                fourth_latest_ConfigNCEIP = None
        else:
            # Handle the case where the directory is empty
            print("Error: the directory is empty")

        fileEXFO2 = glob.glob(os.path.join(path_dir_EXFO, "*monitors*"))
        sorted_fileEXFO2 = sorted(fileEXFO2, key=os.path.getctime, reverse=False)

        if len(sorted_fileEXFO2) > 0:
            # If there are two or more EXFO files, insert the latest two
            if len(sorted_fileEXFO2) >= 2:
                latest_EXFOfile = sorted_fileEXFO2[0]
                second_latest_fileEXFO = sorted_fileEXFO2[1]
                third_latest_fileEXFO = sorted_fileEXFO2[2] if len(sorted_fileEXFO2) >= 3 else None
                fourth_latest_fileEXFO = sorted_fileEXFO2[3] if len(sorted_fileEXFO2) >= 4 else None
            # If there is only one EXFO file, insert it
            else:
                latest_EXFOfile = sorted_fileEXFO2[0]
                second_latest_fileEXFO = None
                third_latest_fileEXFO = None
                fourth_latest_fileEXFO = None
        else:
            # Handle the case where the directory is empty
            print("Error: the directory is empty")


        files2 = glob.glob(os.path.join(path_dir2, "ConfigOutput*"))
        sorted_files2 = sorted(files2, key=os.path.getctime, reverse=False)

        if len(sorted_files2) > 0:
            # If there are two or more EXFO files, insert the latest two
            if len(sorted_files2) >= 2:
                latest_Configfile = sorted_files2[0]
                second_latest_Configfile = sorted_files2[1]
                third_latest_Configfile = sorted_files2[2] if len(sorted_files2) >= 3 else None
                fourth_latest_Configfile = sorted_files2[3] if len(sorted_files2) >= 4 else None
            # If there is only one EXFO file, insert it
            else:
                latest_Configfile = sorted_files2[0]
                second_latest_Configfile = None
                third_latest_Configfile = None
                fourth_latest_Configfile = None
        else:
            # Handle the case where the directory is empty
            print("Error: the directory is empty")

  
        # Loop through each paragraph of the document, and if the text 'Visio Diagram' is found, insert the file using the InlineShapes method
        for para in document.Paragraphs:
            if 'Network Diagram File' in para.Range.Text:
                if os.path.exists(latest_NDfile):
                    print(f"{latest_NDfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_NDfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='Network Diagram')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_NDfile} not exists")
                    break


        # Loop through each paragraph of the document, and if the text 'radios.csv' is found, insert the file using the InlineShapes method
        for para in document.Paragraphs:
            if 'NLOS FWA File' in para.Range.Text:
                if os.path.exists(latest_NLOSfile):
                    print(f"{latest_NLOSfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_NLOSfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='NLOS FWA')
                    break
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_NLOSfile} not exists")
                    break

        # Loop through each paragraph of the document, and if the text 'Visio Diagram' is found, insert the file using the InlineShapes method
        for para in document.Paragraphs:
            if 'EXFO Files' in para.Range.Text:
                if latest_EXFOfile is not None and os.path.exists(latest_EXFOfile):
                    print(f"{latest_EXFOfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_EXFOfile)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_EXFOfile} not exists")

                if second_latest_fileEXFO is not None and os.path.exists(second_latest_fileEXFO):
                    print(f"{second_latest_fileEXFO} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(second_latest_fileEXFO)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_fileEXFO2) > 1:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{second_latest_fileEXFO} not exists")

                if third_latest_fileEXFO is not None and os.path.exists(third_latest_fileEXFO):
                    print(f"{third_latest_fileEXFO} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(third_latest_fileEXFO)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_fileEXFO2) > 2:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{third_latest_fileEXFO} not exists")

                if fourth_latest_fileEXFO is not None and os.path.exists(fourth_latest_fileEXFO):
                    print(f"{fourth_latest_fileEXFO} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(fourth_latest_fileEXFO)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_fileEXFO2) > 3:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{fourth_latest_fileEXFO} not exists")

        # Loop through each paragraph of the document, and if the text 'Config Description NCE-IP' is found, insert the file using the InlineShapes method
        for para in document.Paragraphs:
            if 'Config Description NCE-IP' in para.Range.Text:
                if latest_ConfigNCEIP is not None and os.path.exists(latest_ConfigNCEIP):
                    print(f"{latest_ConfigNCEIP} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_ConfigNCEIP)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_ConfigNCEIP} not exists")

                if second_latest_ConfigNCEIP is not None and os.path.exists(second_latest_ConfigNCEIP):
                    print(f"{second_latest_ConfigNCEIP} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(second_latest_ConfigNCEIP)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_fileNCEIP) > 1:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{second_latest_ConfigNCEIP} not exists")

                if third_latest_ConfigNCEIP is not None and os.path.exists(third_latest_ConfigNCEIP):
                    print(f"{third_latest_ConfigNCEIP} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(third_latest_ConfigNCEIP)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_fileNCEIP) > 2:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{third_latest_ConfigNCEIP} not exists")

                if fourth_latest_ConfigNCEIP is not None and os.path.exists(fourth_latest_ConfigNCEIP):
                    print(f"{fourth_latest_ConfigNCEIP} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(fourth_latest_ConfigNCEIP)),
                    IconFileName=str(Path('C:\\Windows\\Installer\\{90160000-000F-0000-1000-0000000FF1CE}\\xlicons.exe')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_fileNCEIP) > 3:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{fourth_latest_ConfigNCEIP} not exists")

                # Loop through each paragraph of the document, and if the text 'Visio Diagram' is found, insert the file using the InlineShapes method
        for para in document.Paragraphs:
            if 'Config Files' in para.Range.Text:
                if latest_Configfile is not None and os.path.exists(latest_Configfile):
                    print(f"{latest_Configfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(latest_Configfile)),
                    IconFileName=str(Path('C:\\Windows\\System32\\packager.dll')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                else:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{latest_Configfile} not exists")

                if second_latest_Configfile is not None and os.path.exists(second_latest_Configfile):
                    print(f"{second_latest_Configfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(second_latest_Configfile)),
                    IconFileName=str(Path('C:\\Windows\\System32\\packager.dll')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_files2) > 1:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{second_latest_Configfile} not exists")

                if third_latest_Configfile is not None and os.path.exists(third_latest_Configfile):
                    print(f"{third_latest_Configfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(third_latest_Configfile)),
                    IconFileName=str(Path('C:\\Windows\\System32\\packager.dll')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_files2) > 2:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{third_latest_Configfile} not exists")

                if fourth_latest_Configfile is not None and os.path.exists(fourth_latest_Configfile):
                    print(f"{fourth_latest_Configfile} exists")
                    time.sleep(3)
                    para.Range.InlineShapes.AddOLEObject(FileName=str(Path(fourth_latest_Configfile)),
                    IconFileName=str(Path('C:\\Windows\\System32\\packager.dll')),
                        DisplayAsIcon=1) #IconLabel='EXFO')
                elif len(sorted_files2) > 3:
                    # If the picture file does not exist, insert 'Not Applicable'
                    print(f"{fourth_latest_Configfile} not exists")

        # save the document
        time.sleep(3)
        document.Save()
        document.Close()
        word.Quit()
        messagebox.showinfo('information',
                            'Hi! Your BIS Checklist Template is done.Please refer to your Checklist Output>BIS Checklist Template.docx file')

    def NMS_Mainpage(self):

        self.nmspage = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                              border_color="Gray", border_width=20, corner_radius=20)
        self.nmspage.place(x=40, y=40)

        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="CNMAESTRO", command=self.cnmaestro_login).place(relx=0.3, rely=0.2, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="BREEZE AIR", command=self.nmsbreezeair).place(relx=0.3, rely=0.4, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="EXFO", command=self.exfo_login).place(relx=0.3, rely=0.6, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="NCM", command=self.ncm_login).place(relx=0.3, rely=0.8, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="NCE-IP", command=self.NCEIP_GUI).place(relx=0.5, rely=0.2, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="NCE-FAN",command=self.NCEFAN_GUI).place(relx=0.5, rely=0.4, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="PRTG", command=self.PRTGLogin).place(relx=0.5, rely=0.6, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="PDH MW", command=self.PDH_MWGUI).place(relx=0.5, rely=0.8, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="CONFIG DESC.", command=self.Configs).place(relx=0.7, rely=0.2, anchor=tkinter.CENTER)
        CTkButton(master=self.nmspage, width=120, height=100, border_width=0,
                  corner_radius=8,
                  text="CONFIG DESC.", command=self.ConfigNCE_GUI).place(relx=0.7, rely=0.4, anchor=tkinter.CENTER)
        CTkButton(self.nmspage, text="Back", width=100, height=30, border_width=0,
                  command=self.Automate).place(x=780, y=540)
        

    def Meraki_MainPage(self):

        self.codeusername = StringVar()
        self.codepassword = StringVar()
        self.codesdwanid = StringVar()
        self.combomeraki = StringVar()

        # set default values
        self.meraframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                border_color="Gray", border_width=20, corner_radius=20)
        self.meraframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.meraframe, image=my_image).place(x=20, y=120)

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.meraframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (MERAKI MAIN)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=70)
        # ---------------------------------------USERNAME-----------------------------------------------------------
        user2 = CTkLabel(self.meraframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=495, y=150)
        user1 = customtkinter.CTkEntry(master=self.meraframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusername)
        user1.place(x=500, y=190)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.meraframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=495, y=245)
        self.password_entry = customtkinter.CTkEntry(master=self.meraframe,
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, show="*", textvariable=self.codepassword)
        self.password_entry.place(x=500, y=285)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.meraframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=282)
        # -------------------------------------SDWANID-------------------------------------------------------------
        apname2 = CTkLabel(self.meraframe, width=120, text="SDWAN ID or CIRCUIT ID*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=500, y=335)
        apname3 = customtkinter.CTkEntry(master=self.meraframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codesdwanid)
        apname3.place(x=500, y=375)

        # create the Combobox widget and populate it with some values
        optionmeraki = ["Option 1", "Option 2", "Option 3", "Option 4", "Option 5", "Option 6", "Option 7"]
        self.combomeraki = CTkComboBox(master=self.meraframe, values=optionmeraki, width=200)
        self.combomeraki.place(x=550, y=430)
        CTkLabel(master=self.meraframe, text="")
        CTkButton(master=self.meraframe, text="Submit", width=100, command=self.MerakiMain, height=30,
                  border_width=0).place(x=760, y=500)
        CTkButton(master=self.meraframe, text="Back", width=100, height=30, border_width=0,
        command=self.Automate).place(x=650, y=500)

    def MerakiMain(self):
                    
        ID = self.codeusername.get()
        PASSWORD = self.codepassword.get()
        SDWANID = self.codesdwanid.get()
        Selected_Valuemeraki = self.combomeraki.get()

        messagebox.showinfo('information',
                    'Hi! Your BIS Automation process is currently in progress.Please wait until the MERAKI program is finished running.')

        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        #self.bot = webdriver.Chrome(service=self.s)
        self.bot.get("https://")
        self.bot.maximize_window()

        # log in
        a = WebDriverWait(self.bot, 10).until(EC.element_to_be_clickable((By.ID, "email")))
        ActionChains(self.bot).move_to_element(a).click(a).send_keys(ID).perform()
        self.bot.implicitly_wait(5)

        b = WebDriverWait(self.bot, 10).until(EC.element_to_be_clickable((By.ID, "next-btn")))
        ActionChains(self.bot).move_to_element(b).click(b).perform()
        self.bot.implicitly_wait(5)

        c = WebDriverWait(self.bot, 10).until(EC.element_to_be_clickable((By.ID, "password")))
        ActionChains(self.bot).move_to_element(c).click(c).send_keys(PASSWORD).perform()
        self.bot.implicitly_wait(5)

        self.bot.find_element(By.ID, "remember_user").click()
        self.bot.implicitly_wait(5)

        self.bot.find_element(By.ID, "login-btn").click()

        value = simpledialog.askstring("Enter Verification Code", "Enter the Verification Code:")

        verifybox = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "code")))
        ActionChains(self.bot).move_to_element(verifybox).click(verifybox).send_keys(value, Keys.ENTER).perform()
        time.sleep(5)

        #----------------------------------------------------------NEW STEP---------------------------------------------------------------------#
        # Choose Organization - Button Manual click
        NewStep = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@class='mds-flex mds-global-nav-select-button']")))
        ActionChains(self.bot).move_to_element(NewStep).click(NewStep).perform()
        time.sleep(5)

        # auto choose organization from combobox enter by user
        organization_Tab= WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, f"//span[normalize-space()='{Selected_Valuemeraki}']")))
        ActionChains(self.bot).move_to_element(organization_Tab).click(organization_Tab).perform()
        time.sleep(5)

        #click on overview-----------
        tab_xpaths = ["//span[@class='tabTitle']", "//span[normalize-space()='Organization']", 
                      "//div[@class='mds-flex mds-global-nav-item mds-global-nav-item-active']"]

        for xpath in tab_xpaths:
            try:
                tab_element = WebDriverWait(self.bot, 20).until(
                    EC.visibility_of_element_located((By.XPATH, xpath)))
                tab_element.click()
                break  # Exit the loop if a visible element is found
            except:
                continue  # Continue to the next xpath if the current one fails

        # Organization --> Overview tabs
        other_tab_xpaths = ["//span[normalize-space()='Overview']", "//a[@class='mds-global-nav-menu-link mds-global-nav-menu-link-active']//div[@class='mds-flex']//div[@class='mds-flex'][normalize-space()='Overview']",
                            "//a[@data-gtm-menu-navigation='organization > monitor > overview']//div[@class='mds-flex']//div[@class='mds-flex'][normalize-space()='Overview']"]

        for xpath in other_tab_xpaths:
            try:
                other_tab_element = WebDriverWait(self.bot, 20).until(
                    EC.visibility_of_element_located((By.XPATH, xpath)))
                other_tab_element.click()
                break  # Exit the loop if a visible element is found
            except:
                continue  # Continue to the next xpath if the current one fails

        # Search SDWAN ID
        g = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "search_holder")))
        ActionChains(self.bot).move_to_element(g).click(g).send_keys(SDWANID).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # first screenshot
        self.bot.save_screenshot("Checklist Output\\MERAKI\\MerakiTags.png")
        self.bot.implicitly_wait(5)
        time.sleep(5)

        try:
            # SDWAN device link on Tags Column
            #h = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[1]/div[3]/div[1]/div[3]/div[1]/div[2]/div/div[2]/table/tbody/tr/td[3]")))
            h = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[1]/div[3]/div[1]/div[5]/div[1]/div[2]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            ActionChains(self.bot).move_to_element(h).click(h).perform()
            time.sleep(5)
        except:
            print("No element detected!")


        try:
            # Cellular Gateways (DCG Device)
            DCG_tab = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Cellular Gateway']")))
        except:
            print('No element for DCG Tab')
        else:
            ActionChains(self.bot).move_to_element(DCG_tab).click(DCG_tab).perform()
            time.sleep(2)

            # Cellular Gateways (DCG Device)
            DCG = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Cellular gateways')]")))
            ActionChains(self.bot).move_to_element(DCG).click(DCG).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device.png")
            time.sleep(5)

            #----------------------------------------------------------NEW STEP---------------------------------------------------------------------#
            # click DCG Device link
            try:
                DCG_link = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable(
                    (By.XPATH, "/html[1]/body[1]/div[1]/div[3]/div[1]/div[6]/div[1]/div[1]/div[1]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/a[1]")))
                ActionChains(self.bot).move_to_element(DCG_link).click(DCG_link).perform()
                time.sleep(5)
            except:
                print('No DCG link')

            # click DCG Uplink
            DCG_Uplink = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//a[normalize-space()='Uplink']")))
            ActionChains(self.bot).move_to_element(DCG_Uplink).click(DCG_Uplink).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device Configuration.png")
            time.sleep(5)


            LiveData = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//div[normalize-space()='Connectivity to']"))) #go to Connectivity to
            ActionChains(self.bot).move_to_element(LiveData).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device Live Data.png")
            time.sleep(5)

            LiveData2 = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//p[normalize-space()='Last login']")))
            ActionChains(self.bot).move_to_element(LiveData2).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device Latency & Loss.png")
            time.sleep(5)

        # Security & SDWAN
        i = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Security & SD-WAN']")))
        ActionChains(self.bot).move_to_element(i).click(i).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Security & SDWAN > Appliance status
        ii = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Appliance status')]")))
        ActionChains(self.bot).move_to_element(ii).click(ii).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Security & SDWAN > Appliance status > Uplink
        iii = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//a[normalize-space()='Uplink']")))
        ActionChains(self.bot).move_to_element(iii).click(iii).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # second screenshot
        self.bot.save_screenshot("Checklist Output\\MERAKI\\DRO Device Configuration.png")
        time.sleep(5)

        newstep = self.bot.find_element(By.XPATH, "//p[normalize-space()='Last login']")
        ActionChains(self.bot).move_to_element(newstep).perform()
        time.sleep(5)


        # third screenshot
        self.bot.save_screenshot("Checklist Output\\MERAKI\\DRO Device Live Data.png")
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # check if the device has Wireless or not
        try:
            Wirelesselement = WebDriverWait(self.bot, 20).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Wireless']")))
        except:
            # eleement does not exist, do something else
            print('No Wireless Devices in MERAKI')
        else:
            # Access Points
            ActionChains(self.bot).move_to_element(Wirelesselement).click(Wirelesselement).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)

            n = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable(
                (By.XPATH, "//div[contains(text(),'Access points')]")))
            ActionChains(self.bot).move_to_element(n).click(n).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)


            # fifth screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DAP Device Connectivity.png")
            self.bot.implicitly_wait(5)
            time.sleep(5)

        try:
            # switch tab
            k = WebDriverWait(self.bot, 20).until(EC.element_to_be_clickable(
                (By.XPATH, "//span[normalize-space()='Switching']")))

        except:
            # eleement does not exist, do something else
            messagebox.showinfo('information',
                                'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>MERAKI folder.')
            self.bot.close()
            time.sleep(3)

        else:
            ActionChains(self.bot).move_to_element(k).click(k).perform()
            self.bot.implicitly_wait(5)
            time.sleep(8)

            # Switches tab
            l = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH,
                                                                              "//div[contains(text(),'Switches')]")))
            ActionChains(self.bot).move_to_element(l).click(l).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)

            # fourth screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DSW Device Connectivity.png")
            self.bot.implicitly_wait(2)
            time.sleep(2)

            # message box appear
            messagebox.showinfo('information',
                                'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>MERAKI folder.')
            self.bot.close()
            time.sleep(5)

    def MerakiIDMInput(self):

        self.codeusername = StringVar()
        self.codepassword = StringVar()
        self.codesdwanid = StringVar()
        self.combomeraki = StringVar()

        # set default values
        self.idmframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.idmframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.idmframe, image=my_image).place(x=20, y=120)

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.idmframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (MERAKI IDM)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=70)
        # ---------------------------------------USERNAME-----------------------------------------------------------
        user2 = CTkLabel(self.idmframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=490, y=150)
        user1 = customtkinter.CTkEntry(master=self.idmframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusername)
        user1.place(x=500, y=190)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.idmframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=490, y=245)
        self.password_entry = customtkinter.CTkEntry(master=self.idmframe,
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, show="*", textvariable=self.codepassword)
        self.password_entry.place(x=500, y=285)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.idmframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=282)
        # -------------------------------------SDWANID-------------------------------------------------------------
        apname2 = CTkLabel(self.idmframe, width=120, text="SDWAN ID or CIRCUIT ID*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=500, y=335)
        apname3 = customtkinter.CTkEntry(master=self.idmframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codesdwanid)
        apname3.place(x=500, y=375)
        # create the Combobox widget and populate it with some values
        optionmeraki = ["Option 1", "Option 2", "Option 3", "Option 4", "Option 5", "Option 6", "Option 7"]
        self.combomeraki = CTkComboBox(master=self.idmframe, values=optionmeraki, width=200)
        self.combomeraki.place(x=550, y=430)
        CTkLabel(master=self.idmframe, text="")
        CTkButton(master=self.idmframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.MerakiIDM).place(x=760, y=500)        
        CTkButton(master=self.idmframe, text="Back", width=100, height=30, border_width=0,
        command=self.Automate).place(x=650, y=500)

    def MerakiIDM(self):
        ID = self.codeusername.get()
        PASSWORD = self.codepassword.get()
        SDWANID = self.codesdwanid.get()
        Selected_ValuemerakiIDM = self.combomeraki.get()

        messagebox.showinfo('information',
                    'Hi! Your BIS Automation process is currently in progress.Please wait until the MERAKI program is finished running.')
        
        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        self.bot.get(
            'https://')
        self.bot.maximize_window()

        # log in
        a = WebDriverWait(self.bot, 10).until(EC.element_to_be_clickable((By.ID, "Ecom_User_ID")))
        ActionChains(self.bot).move_to_element(a).click(a).send_keys(ID).perform()
        self.bot.implicitly_wait(5)

        b = WebDriverWait(self.bot, 10).until(EC.element_to_be_clickable((By.ID, "Ecom_Password")))
        ActionChains(self.bot).move_to_element(b).click(b).send_keys(PASSWORD).perform()
        self.bot.implicitly_wait(5)

        self.bot.find_element(By.ID, "loginButton2").click()
        messagebox.showwarning('Warning',
                            'Please click Accept on your smartphone NETIQ Authenticator')
        time.sleep(30)

        # Choose Organization - Button Manual click
        d = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@class='mds-flex mds-global-nav-select-button']")))
        ActionChains(self.bot).move_to_element(d).click(d).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # auto choose organization from combobox enter by user
        organization_Tab= WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, f"//span[normalize-space()='{Selected_ValuemerakiIDM}']")))
        ActionChains(self.bot).move_to_element(organization_Tab).click(organization_Tab).perform()
        time.sleep(10)

        #click on overview-----------
        tab_xpaths = ["//span[@class='tabTitle']", "//span[normalize-space()='Organization']", 
                      "//div[@class='mds-flex mds-global-nav-item mds-global-nav-item-active']"]

        for xpath in tab_xpaths:
            try:
                tab_element = WebDriverWait(self.bot, 20).until(
                    EC.visibility_of_element_located((By.XPATH, xpath)))
                tab_element.click()
                break  # Exit the loop if a visible element is found
            except:
                continue  # Continue to the next xpath if the current one fails

        # Organization --> Overview tabs
        other_tab_xpaths = ["//span[normalize-space()='Overview']", "//a[@class='mds-global-nav-menu-link mds-global-nav-menu-link-active']//div[@class='mds-flex']//div[@class='mds-flex'][normalize-space()='Overview']",
                            "//a[@data-gtm-menu-navigation='organization > monitor > overview']//div[@class='mds-flex']//div[@class='mds-flex'][normalize-space()='Overview']"]

        for xpath in other_tab_xpaths:
            try:
                other_tab_element = WebDriverWait(self.bot, 20).until(
                    EC.visibility_of_element_located((By.XPATH, xpath)))
                other_tab_element.click()
                break  # Exit the loop if a visible element is found
            except:
                continue  # Continue to the next xpath if the current one fails

        # Search SDWAN ID
        g = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.ID, "search_holder")))
        ActionChains(self.bot).move_to_element(g).click(g).send_keys(SDWANID).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # first screenshot
        self.bot.save_screenshot("Checklist Output\\MERAKI\\MerakiTags.png")
        self.bot.implicitly_wait(5)
        time.sleep(5)

        try:
            # SDWAN device link on Tags Column
            #h = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[1]/div[3]/div[1]/div[3]/div[1]/div[2]/div/div[2]/table/tbody/tr/td[3]")))
            h = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[1]/div[3]/div[1]/div[5]/div[1]/div[2]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]")))
            ActionChains(self.bot).move_to_element(h).click(h).perform()
            time.sleep(5)
        except:
            print("No element detected!")

        try:
            # Cellular Gateways (DCG Device)
            DCG_tab = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Cellular Gateway']")))
        except:
            print('No element for DCG Tab')
        else:
            ActionChains(self.bot).move_to_element(DCG_tab).click(DCG_tab).perform()
            time.sleep(2)

            # Cellular Gateways (DCG Device)
            DCG = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Cellular gateways')]")))
            ActionChains(self.bot).move_to_element(DCG).click(DCG).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device.png")
            time.sleep(5)

            #----------------------------------------------------------NEW STEP---------------------------------------------------------------------#
            # DCG Device link
            try:
                DCG_link = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable(
                    (By.XPATH, "/html[1]/body[1]/div[1]/div[3]/div[1]/div[6]/div[1]/div[1]/div[1]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/a[1]")))
                ActionChains(self.bot).move_to_element(DCG_link).click(DCG_link).perform()
                time.sleep(5)
            except:
                print('No DCG link')

            try:
                DCG_link = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable(
                    (By.XPATH, "/html[1]/body[1]/div[1]/div[3]/div[1]/div[4]/div[1]/div[1]/div[1]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/a[1]")))
                ActionChains(self.bot).move_to_element(DCG_link).click(DCG_link).perform()
                time.sleep(5)
            except:
                DCG_link2 = WebDriverWait(self.bot, 60).until(EC.element_to_be_clickable(
                    (By.XPATH, "/html[1]/body[1]/div[1]/div[3]/div[1]/div[5]/div[1]/div[1]/div[1]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[3]/a[1]")))
                ActionChains(self.bot).move_to_element(DCG_link2).click(DCG_link2).perform()
                time.sleep(5)

            # DCG Uplink
            DCG_Uplink = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//a[normalize-space()='Uplink']")))
            ActionChains(self.bot).move_to_element(DCG_Uplink).click(DCG_Uplink).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device Configuration.png")
            time.sleep(5)


            LiveData = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//div[normalize-space()='Connectivity to']")))
            ActionChains(self.bot).move_to_element(LiveData).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device Live Data.png")
            time.sleep(5)

            LiveData2 = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//p[normalize-space()='Last login']")))
            ActionChains(self.bot).move_to_element(LiveData2).perform()
            time.sleep(5)

            # DCG screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DCG Device Latency & Loss.png")
            time.sleep(5)

        # Security & SDWAN
        i = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Security & SD-WAN']")))
        ActionChains(self.bot).move_to_element(i).click(i).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Security & SDWAN > Appliance status
        ii = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Appliance status')]")))
        ActionChains(self.bot).move_to_element(ii).click(ii).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Security & SDWAN > Appliance status > Uplink
        iii = WebDriverWait(self.bot, 60).until(
            EC.element_to_be_clickable((By.XPATH, "//a[normalize-space()='Uplink']")))
        ActionChains(self.bot).move_to_element(iii).click(iii).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # second screenshot
        self.bot.save_screenshot("Checklist Output\\MERAKI\\DRO Device Configuration.png")
        self.bot.implicitly_wait(5)
        time.sleep(5)

        newstep = self.bot.find_element(By.XPATH, "//p[normalize-space()='Last login']")
        ActionChains(self.bot).move_to_element(newstep).perform()
        time.sleep(5)

        # third screenshot
        self.bot.save_screenshot("Checklist Output\\MERAKI\\DRO Device Live Data.png")
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # check if the device has Wireless or not
        try:
            # Wireless Tab
            m = WebDriverWait(self.bot, 20).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Wireless']")))
        except:
            # eleement does not exist, do something else
            print('No Wireless Devices in MERAKI')
        else:
            ActionChains(self.bot).move_to_element(m).click(m).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)

            # Access point
            n = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Access points')]")))
            ActionChains(self.bot).move_to_element(n).click(n).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)

            # fifth screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DAP Device Connectivity.png")
            self.bot.implicitly_wait(5)
            time.sleep(5)

        try:
            # switch tab
            k = WebDriverWait(self.bot, 40).until(
                EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Switching']")))
        except:
            # eleement does not exist, do something else
            messagebox.showinfo('information',
                                'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>MERAKI folder.')
            self.bot.close()
            time.sleep(3)
        else:
            ActionChains(self.bot).move_to_element(k).click(k).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)

            # Switches tab
            l = WebDriverWait(self.bot, 60).until(
                EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Switches')]")))
            ActionChains(self.bot).move_to_element(l).click(l).perform()
            self.bot.implicitly_wait(5)
            time.sleep(5)

            # fourth screenshot
            self.bot.save_screenshot("Checklist Output\\MERAKI\\DSW Device Connectivity.png")
            self.bot.implicitly_wait(5)
            time.sleep(5)

            # message box appear
            messagebox.showinfo('information',
                                'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>MERAKI folder.')
            self.bot.close()
            time.sleep(5)

    # NMS Login
    def cnmaestro_login(self):

        self.lrd1 = StringVar()
        # set default values
        self.cnframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                              border_color="Gray", border_width=20, corner_radius=20)
        self.cnframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.cnframe, image=my_image).place(x=60, y=120)

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.cnframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (CNMAESTRO)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=70)
        # ---------------------------------------LRD-----------------------------------------------------------
        user2 = CTkLabel(self.cnframe, width=120, text="LRD", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=510, y=250)
        user1 = customtkinter.CTkEntry(master=self.cnframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.lrd1)
        user1.place(x=550, y=300)

        CTkButton(master=self.cnframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.cnmaestro).place(x=750, y=500)
        CTkButton(master=self.cnframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)

    def cnmaestro(self):

        ID = "xxxxxx"
        PASSWORD = "xxxxxx"
        LRD = self.lrd1.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the CNMAESTRO program is finished running.')

        bot = webdriver.Chrome(service=self.s, options=self.options)
        bot.maximize_window()
        bot.get('https://')

        # LOGIN
        user = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input[type=\"text\" i]")))
        user.send_keys(ID)
        bot.find_element(By.CSS_SELECTOR, "input[type=\"password\" i] ").send_keys(PASSWORD)
        bot.find_element(By.CSS_SELECTOR, "[type=submit]").click()
        bot.implicitly_wait(5)

        # Click Moniter and Managed - Search LRD
        a = bot.find_element(By.XPATH, "//*[@id=\"app\"]/nav/div/div/nav/ul/li[3]/a/span[2]")
        bot.implicitly_wait(5)
        ActionChains(bot).move_to_element(a).click(a).perform()
        time.sleep(10)

        bot.find_element(By.CSS_SELECTOR, "[type=search]").send_keys(LRD)
        bot.find_element(By.CSS_SELECTOR, "[type=search]").send_keys(Keys.ENTER) #paste & search the LRD insert by user in GUI
        bot.implicitly_wait(5)
        time.sleep(8)

        # Dropdown 1 (1st box)
        b = bot.find_element(By.XPATH,
                             "/html/body/cnssng-app/section/ui-view/div/div/div[2]/div/aside/div[1]/div/ul/li[1]/a")
        bot.implicitly_wait(10)
        ActionChains(bot).move_to_element(b).click(b).perform()
        bot.implicitly_wait(50)
        time.sleep(8)

        # Screenshot 1
        bot.save_screenshot("Checklist Output\\CNMAESTRO\\Dashboard1.png")
        time.sleep(5)

        # click Performance tab
        c = bot.find_element(By.XPATH,
                             "//uib-tab-heading[normalize-space()='Performance']")
        bot.implicitly_wait(5)
        ActionChains(bot).move_to_element(c).click(c).perform()
        bot.implicitly_wait(20)
        time.sleep(8)

        # Screenshot 2
        bot.save_screenshot("Checklist Output\\CNMAESTRO\\Performance1.png")
        time.sleep(5)

        # Click search box and enter the LRD again
        d = WebDriverWait(bot, 60).until(EC.element_to_be_clickable((By.XPATH, "//input[@aria-label='Search']")))
        bot.implicitly_wait(5)
        ActionChains(bot).move_to_element(d).click(d).perform()

        bot.find_element(By.CSS_SELECTOR, "[type=search]").send_keys(LRD)
        bot.find_element(By.CSS_SELECTOR, "[type=search]").send_keys(Keys.ENTER)
        bot.implicitly_wait(5)
        time.sleep(3)

        # Dropdown 2 (2nd box)
        e = bot.find_element(By.XPATH,
                             "/html/body/cnssng-app/section/ui-view/div/div/div[2]/div/aside/div[1]/div/ul/li[2]/a")
        bot.implicitly_wait(10)
        ActionChains(bot).move_to_element(e).click(e).perform()
        bot.implicitly_wait(8)
        time.sleep(5)

        # Screenshot 2
        bot.save_screenshot("Checklist Output\\CNMAESTRO\\Performance2.png")
        bot.implicitly_wait(5)
        time.sleep(2)

        # Dashboard
        f = WebDriverWait(bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                     "//uib-tab-heading[normalize-space()='Dashboard']")))
        bot.implicitly_wait(5)
        ActionChains(bot).move_to_element(f).click(f).perform()
        bot.implicitly_wait(8)
        time.sleep(5)

        # Screenshot 3
        bot.save_screenshot("Checklist Output\\CNMAESTRO\\Dashboard2.png")
        bot.implicitly_wait(5)
        time.sleep(5)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>CNMAESTRO folder.')
        bot.close()
        time.sleep(5)


    def exfo_login(self):

        self.nodename1 = StringVar()
        # set default values
        self.exfoframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                border_color="Gray", border_width=20, corner_radius=20)
        self.exfoframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.exfoframe, image=my_image).place(x=60, y=120)

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.exfoframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (EXFO)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=70)
        # ---------------------------------------nodename-----------------------------------------------------------
        user2 = CTkLabel(self.exfoframe, width=125, text="NODENAME or CIRCUIT ID", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=550, y=250)
        user1 = customtkinter.CTkEntry(master=self.exfoframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.nodename1)
        user1.place(x=550, y=300)

        CTkButton(master=self.exfoframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.exfo).place(x=750, y=500)
        CTkButton(master=self.exfoframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)

    def exfo(self):

        ID = "xxxxx"
        PASSWORD = "xxxxx"
        NodeNAME = self.nodename1.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the EXFO program is finished running.')

        # set the download directory and circuit ID
        download_dir_EXFO = "C:\\Automate\\Checklist Output\\EXFO"

        self.options = Options()

        # this parameter tells Chrome that
        # it should be run without UI (Headless)
        self.options.add_argument("--disable-infobars")
        self.options.add_argument("--start-maximized")
        self.options.add_argument("--disable-extensions")
        self.options.add_argument('--window-size=1920,1080')
        self.options.add_argument("--ignore-certificate-errors")
        self.options.add_argument('--ignore-ssl-errors=yes')
        prefs = {"download.default_directory": download_dir_EXFO}
        self.options.add_experimental_option("prefs", prefs)
        self.options.add_argument("--headless=new")

        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        #self.bot = webdriver.Chrome(service=self.s)
        self.bot.maximize_window()
        urlexfo= "https://"

        
        try:
            self.bot.get(urlexfo)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            self.bot.close()

        # log in
        user = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.ID, "user")))
        ActionChains(self.bot).move_to_element(user).click(user).send_keys(ID).perform()

        pw = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.ID, "pass")))
        ActionChains(self.bot).move_to_element(pw).click(pw).send_keys(PASSWORD).perform()

        self.bot.implicitly_wait(8)

        self.bot.find_element(By.ID, "login-btn").click()
        self.bot.implicitly_wait(8)
        time.sleep(8)

        self.bot.set_page_load_timeout(8)

        #--------------------------#NEW ADDED#click on Element button--------------------------------------------#
        #click on the Element tab
        first = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "/html[1]/body[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[1]/div[2]/div[1]/a[1]/span[1]/span[1]/span[2]/span[1]")))
        ActionChains(self.bot).move_to_element(first).click(first).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # Element Manager - click type
        a = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
              '/html[1]/body[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[1]/input[1]')))
        ActionChains(self.bot).move_to_element(a).click(a).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # select Group from dropdown
        b = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                    "//li[normalize-space()='Group']")))
        ActionChains(self.bot).move_to_element(b).click(b).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Click on anywhere - LAN
        c = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//div[@title='LAN']")))
        ActionChains(self.bot).move_to_element(c).click(c).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # select All from dropdown
        d = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
             "//input[@id='ext-232-cep-panel-widget-dashboard-Id-8641-289-cep-drop-target-container-0-0cep-widget-chart-status-org-group-selection-inputEl']")))
        ActionChains(self.bot).move_to_element(d).click(d).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # click on checkbox - All
        e = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='ALL']")))
        ActionChains(self.bot).move_to_element(e).click(e).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Click on View
        f = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
             "//input[@id='ext-232-cep-panel-widget-dashboard-Id-8641-289-cep-drop-target-container-0-0cep-widget-chart-agent-viewmode-inputEl']")))
        ActionChains(self.bot).move_to_element(f).click(f).perform()
        self.bot.implicitly_wait(5)

        # Click on View : Grid
        g = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//li[normalize-space()='Grid']")))
        ActionChains(self.bot).move_to_element(g).click(g).perform()
        self.bot.implicitly_wait(5)
        time.sleep(8)

        # Click on next arrow
        h = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//div[@id='tool-1110-toolEl']")))
        ActionChains(self.bot).move_to_element(h).click(h).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        # Click on search box
        hh = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
            "//input[@id='ext-232-cep-panel-widget-dashboard-Id-8641-289-cep-drop-target-container-0-0cep-widget-chart-agent-search-inputEl']")))
        ActionChains(self.bot).move_to_element(hh).click(hh).send_keys(NodeNAME).perform()
        time.sleep(10)
        self.bot.find_element(By.XPATH,
                        "//input[@id='ext-232-cep-panel-widget-dashboard-Id-8641-289-cep-drop-target-container-0-0cep-widget-chart-agent-search-inputEl']").send_keys(Keys.ENTER)
        time.sleep(10)

        # export report based on node
        report = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                        "/html[1]/body[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/div[1]/div[3]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[1]/div[3]/div[3]/div[1]/div[1]/div[1]/div[1]/div[1]/div[5]/div[1]/div[1]/div[1]/div[1]/div[1]/a[8]/span[1]/span[1]/span[2]")))
        ActionChains(self.bot).move_to_element(report).click(report).perform()
        self.bot.implicitly_wait(5)
        time.sleep(5)

        #--------------------------#NEW ADDED#click on ALERT tab-------------------------------------------#
        ALERT = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                "/html[1]/body[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[1]/div[2]/div[1]/a[2]/span[1]/span[1]/span[2]/span[1]")))
        ActionChains(self.bot).move_to_element(ALERT).click(ALERT).perform()
        time.sleep(5)  
        
        # click on search field
        searchalert = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                "/html/body/div[1]/div/div[3]/div/div/div[2]/div/div/div/div[2]/div[2]/div/div[1]/div/div/div/div/div/table/tbody/tr/td/div/div/div[3]/div/div[1]/div/div/div[7]/div")))
        ActionChains(self.bot).move_to_element(searchalert).click(searchalert).send_keys(NodeNAME, Keys.ENTER).perform()
        #self.bot.find_element(By.XPATH, "/html/body/div[1]/div/div[3]/div/div/div[2]/div/div/div/div[2]/div[2]/div/div[1]/div/div/div/div/div/table/tbody/tr/td/div/div/div[3]/div/div[1]/div/div/div[7]/div").send_keys(Keys.ENTER)
        time.sleep(5)

        # Screenshot ALERT WARNING
        self.bot.save_screenshot(f"Checklist Output\\EXFO\\Warning Alert_{NodeNAME}.png")
        self.bot.implicitly_wait(5)
        time.sleep(3)

        #--------------------------#NEW ADDED#click on ELEMENT tab again-------------------------------------------#
        ELEMENT = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                "/html[1]/body[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[1]/div[2]/div[1]/a[1]/span[1]/span[1]/span[2]/span[1]")))
        ActionChains(self.bot).move_to_element(ELEMENT).click(ELEMENT).perform()
        time.sleep(5)    


        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to Checklist Output>EXFO>monitors.csv file.')
        self.bot.close()
        time.sleep(5)

    def ncm_login(self):

        self.nodename = StringVar()
        # set default values
        self.ncmframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.ncmframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(500, 400))
        CTkLabel(self.ncmframe, image=my_image).place(x=60, y=120)

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.ncmframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (NCM)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=70)
        # ---------------------------------------nodename-----------------------------------------------------------
        user2 = CTkLabel(self.ncmframe, width=120, text="NODENAME", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=550, y=250)
        user1 = customtkinter.CTkEntry(master=self.ncmframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.nodename)
        user1.place(x=550, y=300)

        CTkButton(master=self.ncmframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.ncm).place(x=750, y=500)
        CTkButton(master=self.ncmframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)
        

    def ncm(self):

        ID = "xxxxx"
        PASSWORD = "xxxxxx"
        NodeNAME = self.nodename.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the NCM Solarwind program is finished running.')

        # create webdriver object
        self.driver = webdriver.Chrome(service=self.s, options=self.options)
        self.driver.maximize_window()

        # get ncm
        self.driver.get("https://")

        # log in
        a = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "ctl00_BodyContent_Username")))
        ActionChains(self.driver).move_to_element(a).click(a).send_keys(ID).perform()
        self.driver.implicitly_wait(5)

        b = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "ctl00_BodyContent_Password")))
        ActionChains(self.driver).move_to_element(b).click(b).send_keys(PASSWORD).perform()
        self.driver.implicitly_wait(5)

        self.driver.find_element(By.ID, "ctl00_BodyContent_LoginButton").click()
        self.driver.implicitly_wait(5)
        time.sleep(8)

        # Search node in NCM - click on the search button
        c = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable(
            (By.XPATH, "//*[@id=\"pageHeader\"]/sw-mega-menu/div/div/sw-orion-search-component")))
        ActionChains(self.driver).move_to_element(c).click(c).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)
        #paste yhe node name in the searchbox
        d = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                             "/html/body/div[1]/div[1]/div/sw-mega-menu/div/div/sw-orion-search-component/div/div/div/div[1]/div")))
        ActionChains(self.driver).move_to_element(d).send_keys(NodeNAME).send_keys(Keys.ENTER).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)
        #click on the link rrsult
        e = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                             "/html/body/div/div/div/div/div/div[2]/div[2]/div/div/div/div/div[2]/xui-filtered-list-v2/div/div[2]/div/div[2]/div/div/div/div[1]/div[2]/div[2]/div[2]/div[1]/div/form/ul/li/div/ng-include/div/div/div/div/div[1]/div[1]/a/span[2]")))
        ActionChains(self.driver).move_to_element(e).click(e).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)

        # Screenshot
        self.driver.save_screenshot(f'Checklist Output\\NCM\\NCM_{NodeNAME}.png')
        self.driver.implicitly_wait(5)
        time.sleep(5)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>NCM folder.')
        self.driver.close()
        time.sleep(5)


    def PRTGLogin(self):
        self.prtgminframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                   border_color="Gray", border_width=20, corner_radius=20)
        self.prtgminframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.prtgminframe, image=my_image).place(x=60, y=120)

        self.codeusernamemin = StringVar()
        self.codepasswordmin = StringVar()
        self.codemin1 = StringVar()
        self.codemin2 = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.prtgminframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (PRTG)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        user2 = CTkLabel(self.prtgminframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=495, y=85)
        user1 = customtkinter.CTkEntry(master=self.prtgminframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusernamemin)
        user1.place(x=500, y=125)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.prtgminframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=490, y=165)
        self.password_entry = customtkinter.CTkEntry(master=self.prtgminframe, show="*",
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.codepasswordmin)
        self.password_entry.place(x=500, y=205)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.prtgminframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=202)
        # --------------------------------------MIN Node 1-------------------------------------------------------------
        apname2 = CTkLabel(self.prtgminframe, width=120, text="MIN Node 1", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=490, y=255)
        apname3 = customtkinter.CTkEntry(master=self.prtgminframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codemin1)
        apname3.place(x=500, y=295)
        # --------------------------------------MIN Node 2-------------------------------------------------------------
        apname2 = CTkLabel(self.prtgminframe, width=120, text="MIN Node 2", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=490, y=340)
        apname3 = customtkinter.CTkEntry(master=self.prtgminframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codemin2)
        apname3.place(x=500, y=385)

        CTkButton(master=self.prtgminframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.prtgmin).place(x=750, y=500)
        CTkButton(master=self.prtgminframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)

    def prtgmin(self):

        # configure chrome driver & input items
        self.s = Service("images\\chromedriver.exe")
        ID = self.codeusernamemin.get()
        PASSWORD = self.codepasswordmin.get()
        MIN1 = self.codemin1.get()
        MIN2 = self.codemin2.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the PRTG MIN program is finished running.')

        # create webdriver object
        self.driver = webdriver.Chrome(service=self.s, options=self.options)
        self.driver.maximize_window()

        # get PRTG MIN
        self.driver.get("https://")

        # log in
        username = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//input[@id='loginusername']")))
        ActionChains(self.driver).move_to_element(username).click(username).send_keys(ID).perform()
        self.driver.implicitly_wait(5)
        password = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//input[@id='loginpassword']")))
        ActionChains(self.driver).move_to_element(password).click(password).send_keys(PASSWORD).perform()
        self.driver.implicitly_wait(5)
        submitbutton = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@type='submit']")))
        ActionChains(self.driver).move_to_element(submitbutton).click(submitbutton).perform()
        self.driver.implicitly_wait(5)
        time.sleep(30)

        # click pn search field
        b = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH, "//input[@id='searchtext']")))
        ActionChains(self.driver).move_to_element(b).click(b).send_keys(MIN1).send_keys(Keys.ENTER).perform()
        self.driver.implicitly_wait(5)
        time.sleep(10)

        # click on visible MIN equipment
        c = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable(
            (By.XPATH, "/html[1]/body[1]/div[4]/div[1]/div[1]/div[1]/div[1]/div[2]/div[2]/div[1]/div[3]/div[1]"
                       "/form[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
        ActionChains(self.driver).move_to_element(c).click(c).perform()
        self.driver.implicitly_wait(5)
        time.sleep(10)

        # 6th Screenshot - first sensors list
        self.driver.save_screenshot(f"Checklist Output\\PRTG MIN\\{MIN1}.png")
        time.sleep(5)

        # scroll to the bottom page
        d = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//span[@class='tablenavigation']")))
        ActionChains(self.driver).move_to_element(d).perform()
        time.sleep(3)

        # 6th Screenshot - Summary Tab first
        self.driver.save_screenshot(f"Checklist Output\\PRTG MIN\\{MIN1} bottom page.png")
        time.sleep(5)

        # check if MIN2 has been provided by the user
        if MIN2:
            # click on search field
            bb = WebDriverWait(self.driver, 90).until(
                EC.element_to_be_clickable((By.XPATH, "//input[@id='searchtext']")))
            bb.clear()
            bb1 = WebDriverWait(self.driver, 90).until(
                EC.element_to_be_clickable((By.XPATH, "//input[@id='searchtext']")))
            ActionChains(self.driver).move_to_element(bb1).click(bb1).send_keys(MIN2).send_keys(Keys.ENTER).perform()
            self.driver.implicitly_wait(5)
            time.sleep(10)

            # check if visible MIN2 equipment exists
            try:
                c = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH,
                                                                                     "/html[1]/body[1]/div[4]/div[1]/div[1]/div[1]/div[1]/div[2]/div[2]/div[1]/div[3]/div[1]/form[1]/table[1]/tbody[1]/tr[1]/td[1]/a[1]")))
                ActionChains(self.driver).move_to_element(c).click(c).perform()
                self.driver.implicitly_wait(5)
                time.sleep(10)

                self.driver.save_screenshot(f"Checklist Output\\PRTG MIN\\{MIN2}.png")
                time.sleep(5)

                # scroll to the bottom page
                dd = WebDriverWait(self.driver, 90).until(
                    EC.element_to_be_clickable((By.XPATH, "//span[@class='tablenavigation']")))
                ActionChains(self.driver).move_to_element(dd).perform()
                time.sleep(3)

                # 6th Screenshot - bottom page sensors
                self.driver.save_screenshot(f"Checklist Output\\PRTG MIN\\{MIN2} bottom page.png")
                time.sleep(5)

            except:
                # if no visible MIN2 equipment, display message

                messagebox.showinfo('information',
                                    'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>PRTG MIN folder.')  # message box appear

                self.driver.close()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>PRTG MIN folder.')  # message box appear

        self.driver.close()

    def PDH_MWGUI(self):
        self.pdhframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.pdhframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.pdhframe, image=my_image).place(x=60, y=120)

        self.codeusernamepdh = StringVar()
        self.codepasswordpdh = StringVar()
        self.codepdhaddress1 = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.pdhframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (PDH)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        user2 = CTkLabel(self.pdhframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=495, y=155)
        user1 = customtkinter.CTkEntry(master=self.pdhframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusernamepdh)
        user1.place(x=500, y=195)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.pdhframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=490, y=245)
        self.password_entry = customtkinter.CTkEntry(master=self.pdhframe,show="*",
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.codepasswordpdh)
        self.password_entry.place(x=500, y=285)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.pdhframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=282)
        # --------------------------------------codepdhaddress1-------------------------------------------------------------
        apname2 = CTkLabel(self.pdhframe, width=120, text="IP Address*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=490, y=325)
        apname3 = customtkinter.CTkEntry(master=self.pdhframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codepdhaddress1)
        apname3.place(x=500, y=365)

        CTkButton(master=self.pdhframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.pdhmw).place(x=750, y=500)
        CTkButton(master=self.pdhframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)
        
    def pdhmw(self):

        USERNAME = self.codeusernamepdh.get()
        PASSWORD = self.codepasswordpdh.get()
        IPAddress1 = self.codepdhaddress1.get()

        messagebox.showinfo('information',
                        'Hi! Your BIS Automation process is currently in progress.Please wait until the PDH MW program is finished running.')
        # create webdriver object
        #, options=self.options
        #self.driver = webdriver.Chrome(service=self.s)
        self.driver = webdriver.Chrome(service=self.s, options=self.options)
        self.driver.maximize_window()

        # get IP Address
        url = f"https://{IPAddress1}"
        url2 = f"http://{IPAddress1}"
        #self.driver.get(f"http://{IPAddress1}")

        try:
            try:
                self.driver.get(url)
            except:
                self.driver.get(url2)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            self.driver.close()
        # username & password to login (there are at least 9 type of pdhmw & each of them have different xpath)
        try:
            tab_element_ID_1 = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "UserName")))
            ActionChains(self.driver).move_to_element(tab_element_ID_1).click(tab_element_ID_1).send_keys(USERNAME).perform()
            time.sleep(5)
        except:
            tab_element_ID_2 = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "UserName-inputEl")))
            ActionChains(self.driver).move_to_element(tab_element_ID_2).click(tab_element_ID_2).send_keys(USERNAME).perform()
            time.sleep(5)

        try:
            tab_element_PW = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "PassWord")))
            ActionChains(self.driver).move_to_element(tab_element_PW).click(tab_element_PW).send_keys(PASSWORD).perform()
            time.sleep(5)
        except:
            tab_element_PW_2 = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "PassWord-inputEl")))
            ActionChains(self.driver).move_to_element(tab_element_PW_2).click(tab_element_PW_2).send_keys(PASSWORD).perform()
            time.sleep(5)

        #click OK oif there is pop-up box appear
        tab_xpaths_OK = ["/html[1]/body[1]/div[1]/div[1]/div[1]/div[1]/table[1]/tbody[1]/tr[2]/td[3]/div[1]/div[1]/div[1]/div[2]/div[1]/div[1]/div[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[1]/td[1]/table[1]/tbody[1]/tr[2]/td[2]",
                            "/html[1]/body[1]/div[1]/div[1]/div[1]/div[1]/div[1]/table[1]/tbody[1]/tr[3]/td[2]/div[1]/div[1]/div[2]/div[2]/div[1]/div[1]/a[1]/span[1]/span[1]"]
        for xpath in tab_xpaths_OK:
            try:
                tab_element_OK = WebDriverWait(self.driver, 90).until(EC.visibility_of_element_located((By.XPATH, xpath)))
                ActionChains(self.driver).move_to_element(tab_element_OK).click(tab_element_OK).perform()
                time.sleep(5)
                break  # Exit the loop if a visible element is found
            except:
                continue  # Continue to the next xpath if the current one fails
            
        # Switch to the new tab
        self.driver.switch_to.window(self.driver.window_handles[1])
        time.sleep(20)

        #go to Provisioning tab
        tab_xpaths_Provisioning = ["/html/body/div[1]/div[2]/div[2]/div/div/div/div/div/div/div[1]/div/div/ul/div/li[4]",
                                    "/html[1]/body[1]/div[1]/div[3]/div[1]/div[1]/div[1]/div[1]/div[1]/div[2]/div[1]/div[1]/table[4]/tbody[1]/tr[1]/td[1]/div[1]/span[1]"]
        for xpath in tab_xpaths_Provisioning:
            try:
                tab_xpaths_Provisioning = WebDriverWait(self.driver, 90).until(EC.visibility_of_element_located((By.XPATH, xpath)))
                ActionChains(self.driver).move_to_element(tab_xpaths_Provisioning).click(tab_xpaths_Provisioning).perform()
                break  # Exit the loop if a visible element is found
            except:
                continue  # Continue to the next xpath if the current one fails

        #go to ETH Function Setting tab
        d = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='ETH Function Setting']")))
        ActionChains(self.driver).move_to_element(d).click(d).perform()
        time.sleep(5)

        #go to ETH Port Setting tab
        e = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='ETH Port Setting']")))
        ActionChains(self.driver).move_to_element(e).click(e).perform()
        time.sleep(8)
        
        # Screenshot
        self.driver.save_screenshot(f"Checklist Output\\PDH MW\\PDH_{IPAddress1}.png")
        time.sleep(5)
        
        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to Checklist Output>PDH MW folder.')
        self.driver.close()
        time.sleep(5)


    def NCEIP_GUI(self):
        self.nceframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.nceframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.nceframe, image=my_image).place(x=60, y=120)

        self.codeusernamenceip = StringVar()
        self.codepasswordnceip = StringVar()
        self.codenodenamenceip = StringVar()
        #self.codeportnonceip = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.nceframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (NCEIP)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        user2 = CTkLabel(self.nceframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=490, y=165)
        user1 = customtkinter.CTkEntry(master=self.nceframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusernamenceip)
        user1.place(x=500, y=205)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.nceframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=490, y=245)
        self.password_entry = customtkinter.CTkEntry(master=self.nceframe,
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.codepasswordnceip)
        self.password_entry.place(x=500, y=285)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.nceframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=282)
        # --------------------------------------Node Name-------------------------------------------------------------
        apname2 = CTkLabel(self.nceframe, width=120, text="Node Name*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=490, y=335)
        apname3 = customtkinter.CTkEntry(master=self.nceframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codenodenamenceip)
        apname3.place(x=500, y=375)

        CTkButton(master=self.nceframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.nceip).place(x=750, y=500)
        CTkButton(master=self.nceframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)

    def nceip(self):
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the NCE-IP program is finished running.')

        ID = self.codeusernamenceip.get()
        PASSWORD = self.codepasswordnceip.get()
        NodeNAME = self.codenodenamenceip.get()
        #PortNum = self.codeportnonceip.get()

        # create webdriver object
        self.driver = webdriver.Chrome(service=self.s, options=self.options)
        self.driver.maximize_window()

        # get NCE-IP
        self.driver.get("https://")

        # log in
        a = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "usernameInput")))
        ActionChains(self.driver).move_to_element(a).click(a).send_keys(ID).perform()
        self.driver.implicitly_wait(5)

        b = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "/html/body/div/div[2]/div[2]/div/div[3]/div[1]/div[3]/div[3]")))
        ActionChains(self.driver).move_to_element(b).click(b).send_keys(PASSWORD).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)
        c = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "btn_outerverify")))
        ActionChains(self.driver).move_to_element(c).click(c).perform()
        self.driver.implicitly_wait(5)
        time.sleep(40)

        self.driver.set_page_load_timeout(10)

        # NCE-IP page - Agree Button
        d = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "login_warn_confirm")))
        ActionChains(self.driver).move_to_element(d).click(d).perform()
        self.driver.implicitly_wait(5)
        time.sleep(10)

        # Network Management
        e = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "/html/body/div[2]/div/div/div[2]/div/div[2]/div[4]/div[1]")))
        ActionChains(self.driver).move_to_element(e).click(e).perform()
        time.sleep(10)

        # Switch to the new tab
        self.driver.switch_to.window(self.driver.window_handles[1])
        time.sleep(60)

        # num_iframes = driver.execute_script('return document.getElementsByTagName("iframe").length')
        # print(f'Number of iframes on the page: {num_iframes}')

        # switch to iframe element
        self.driver.switch_to.frame(0)
        time.sleep(8)

        # click on iframe link
        f = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//input[@id='topo-main-search-input']")))
        ActionChains(self.driver).move_to_element(f).click().send_keys(NodeNAME).perform()
        time.sleep(10)

        # click on search button
        g = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="topo-search-action-button"]')))
        ActionChains(self.driver).move_to_element(g).click().perform()
        time.sleep(5)

        # click on node link
        h = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//span[@class='topo-text-ellipsis']")))
        ActionChains(self.driver).move_to_element(h).click().perform()
        time.sleep(5)

        # 1st Screenshot - Node visible or not in system
        self.driver.save_screenshot("Checklist Output\\NCE-IP\\Node Visibility.png")
        time.sleep(5)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to Checklist Output>NCE-IP folder.')
        self.driver.close()

    def NCEFAN_GUI(self):
        self.ncefanframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                               border_color="Gray", border_width=20, corner_radius=20)
        self.ncefanframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.ncefanframe, image=my_image).place(x=60, y=120)

        self.codeusernamencefan = StringVar()
        self.codepasswordncefan = StringVar()
        self.codenodenamencefan = StringVar()
        #self.codeportnoncefan = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.ncefanframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (NCE-FAN)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------Project Name label-----------------------------------------------------------
        CTkLabel(self.ncefanframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold")).place(x=490, y=165)
        CTkLabel(self.ncefanframe, width=120, text="(Please add @maxis)", fg_color="transparent",
                         font=("Tahoma", 10)).place(x=630, y=165)
        customtkinter.CTkEntry(master=self.ncefanframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusernamencefan).place(x=500, y=205)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        CTkLabel(self.ncefanframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold")).place(x=490, y=245)
        self.password_entry = customtkinter.CTkEntry(master=self.ncefanframe,
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.codepasswordncefan)
        self.password_entry.place(x=500, y=285)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.ncefanframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=282)
        # --------------------------------------Node Name-------------------------------------------------------------
        CTkLabel(self.ncefanframe, width=120, text="Node Name*", fg_color="transparent",
                           font=("Tahoma", 15, "bold")).place(x=490, y=335)
        customtkinter.CTkEntry(master=self.ncefanframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codenodenamencefan).place(x=500, y=375)

        CTkButton(master=self.ncefanframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.ncefan).place(x=750, y=500)
        CTkButton(master=self.ncefanframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)
        
    def ncefan(self):
        ID = self.codeusernamencefan.get()
        PASSWORD = self.codepasswordncefan.get()
        NodeNAME = self.codenodenamencefan.get()
        #PortNum = self.codeportnonceip.get()
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the NCE-FAN program is finished running.')
        # create webdriver object
        self.driver = webdriver.Chrome(service=self.s, options=self.options)
        #self.driver = webdriver.Chrome(service=self.s)
        self.driver.maximize_window()

        # get NCE-FAN
        self.driver.get("https://")

        # log in
        a = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "usernameInput")))
        ActionChains(self.driver).move_to_element(a).click(a).send_keys(ID).perform()
        self.driver.implicitly_wait(5)

        b = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//input[@id='value']")))
        ActionChains(self.driver).move_to_element(b).click(b).send_keys(PASSWORD).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)
        c = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "btn_outerverify")))
        ActionChains(self.driver).move_to_element(c).click(c).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)

        self.driver.set_page_load_timeout(10)

        # NCE-FAN page - Agree Button
        d = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.ID, "login_warn_confirm")))
        ActionChains(self.driver).move_to_element(d).click(d).perform()
        self.driver.implicitly_wait(5)
        time.sleep(5)

        # Network Management
        e = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//div[@title='Network Management']")))
        ActionChains(self.driver).move_to_element(e).click(e).perform()
        time.sleep(5)

        # NCE-FAN page - Agree Button
        try:
            ddd = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH, "/html[1]/body[1]/div[8]/div[1]/div[1]/span[1]")))
            ActionChains(self.driver).move_to_element(ddd).click(ddd).perform()
            self.driver.implicitly_wait(5)
            time.sleep(10)
        except:
            print("Continue next step")

        # Switch to the new tab
        self.driver.switch_to.window(self.driver.window_handles[1])
        time.sleep(60)

        # switch to iframe element
        self.driver.switch_to.frame(0)
        time.sleep(8)

        # click on iframe link
        f = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//input[@id='topo_main_search_input']")))
        ActionChains(self.driver).move_to_element(f).click().send_keys(NodeNAME).perform()
        time.sleep(10)

        # click on search button
        g = WebDriverWait(self.driver, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//div[@id='topo_search_action_button_icon']")))
        ActionChains(self.driver).move_to_element(g).click().perform()
        time.sleep(10)

        # click on node link
        try:
            h = WebDriverWait(self.driver, 90).until(EC.element_to_be_clickable((By.XPATH, 
                        "/html[1]/body[1]/div[2]/div[1]/div[1]/div[6]/div[1]/div[2]/div[1]/div[1]/div[2]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[2]/table[1]/tbody[1]/tr[1]/td[1]/div[1]/span[1]")))
            ActionChains(self.driver).move_to_element(h).click().perform()
            time.sleep(10)
        except:
            print("not found!")
        
        # 1st Screenshot - Node visible or not in system
        self.driver.save_screenshot("Checklist Output\\NCE-FAN\\Node Visibility.png")
        time.sleep(5)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to Checklist Output>NCE-FAN folder.')
        self.driver.close()

    def ConfigNCE_GUI(self):
        self.configframence = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                  border_color="Gray", border_width=20, corner_radius=20)
        self.configframence.place(x=40, y=40)

        #my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        #CTkLabel(self.configframe, image=my_image).place(x=60, y=120)

        self.codeusernameConfigNCE = StringVar()
        self.codepasswordConfigNCE = StringVar()
        self.codeIPAddNodeConfigNCE = StringVar()
        self.codePortNumberConfigNCE = StringVar()
        self.codePortNoVLANIDConfigNCE = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.configframence, width=120, text="PLEASE ENTER THE DETAILS BELOW (CONFIGURATION)",
                        font=("Tahoma", 15, "bold"))
        head.place(x=280, y=30)
        CTkLabel(self.configframence,text="This system only to check config for router only", 
                 fg_color="transparent",font=("Tahoma", 10)).place(x=320, y=50)
        # ---------------------------------------USERNAME-----------------------------------------------------------
        user2 = CTkLabel(self.configframence, width=120, text="Username*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=85, y=155)
        CTkLabel(self.configframence,text="for xxxx, please add xxxx", fg_color="transparent",font=("Tahoma", 9)).place(x=100, y=175)
        user1 = customtkinter.CTkEntry(master=self.configframence,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusernameConfigNCE)
        user1.place(x=95, y=195)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.configframence, width=120, text="Password*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=80, y=245)
        self.password_entry = customtkinter.CTkEntry(master=self.configframence, show="*",
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.codepasswordConfigNCE)
        self.password_entry.place(x=95, y=275)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.configframence, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=425, y=272)
        # --------------------------------------IPAddressNode-------------------------------------------------------------
        apname2 = CTkLabel(self.configframence, width=120, text="IP Address*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=85, y=330)
        CTkLabel(self.configframence,text="can refer to screenshot output in Automate > Checklist Output > NCE-IP folder", 
                 fg_color="transparent",font=("Tahoma", 9)).place(x=100, y=350)
        apname3 = customtkinter.CTkEntry(master=self.configframence,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codeIPAddNodeConfigNCE)
        apname3.place(x=95, y=375)
        # --------------------------------------Port Number-------------------------------------------------------------
        apname2 = CTkLabel(self.configframence, width=120, text="xxx*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=500, y=155)
        CTkLabel(self.configframence,text="e.g. Gi0/2/1", fg_color="transparent",font=("Tahoma", 9)).place(x=500, y=175)
        apname3 = customtkinter.CTkEntry(master=self.configframence,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codePortNumberConfigNCE)
        apname3.place(x=495, y=195)
        # --------------------------------------PORTNO-VLANID-------------------------------------------------------------
        PortNo11 = CTkLabel(self.configframence, width=120, text="Port Number & VLAN ID", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        PortNo11.place(x=500, y=235)
        CTkLabel(self.configframence,text="e.g. xxxx", fg_color="transparent",font=("Tahoma", 9)).place(x=500, y=255)
        PortNo11 = customtkinter.CTkEntry(master=self.configframence,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codePortNoVLANIDConfigNCE)
        PortNo11.place(x=495, y=275)
        # Create a Submit button
        self.submitButton =CTkButton(self.configframence, text="Submit", width=100, height=30, border_width=0, command=self.POC321)
        self.submitButton.place(x=460, y=500)
        CTkButton(master=self.configframence, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=350, y=500)


    def POC321(self):
        # for HUAWEI ATN device only!!!!!!
        self.ID = self.codeusernameConfigNCE.get()
        self.Password = self.codepasswordConfigNCE.get()
        self.IPAddNode = self.codeIPAddNodeConfigNCE.get()
        self.PortNumber = self.codePortNumberConfigNCE.get()
        self.PortNumberVLANID = self.codePortNoVLANIDConfigNCE.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the CONFIG DESC program is finished running.')

        conn = paramiko.SSHClient()
        conn.set_missing_host_key_policy(paramiko.AutoAddPolicy())

        conn.connect(f"{self.IPAddNode}", 00, f"{self.ID}", f"{self.Password}")

        commands = conn.invoke_shell()

        # Display int info
        commands.send(f"dis int desc {self.PortNumber}\n")
        time.sleep(5)
        output1 = commands.recv(65535)
        output1 = output1.decode("utf-8")
        print(output1)
        time.sleep(5)

        # Display int info
        commands.send(f"dis current-configuration int GigabitEthernet{self.PortNumberVLANID}\n")
        time.sleep(5)
        output2 = commands.recv(65535)
        output2 = output2.decode("utf-8")
        print(output2)
        time.sleep(5)

        commands.send(b"q\n")
        time.sleep(3)
        output3 = commands.recv(65535)
        output3 = output3.decode("utf-8")
        print(output3)

        # send output to files
        lines = [output1, output2, output3]

        with open(f'C:\\Automate\\Checklist Output\\CONFIG\\POC3POC2POC1\\ConfigOutputNCEIP_{self.IPAddNode}.txt', 'w') as f:
            for line in lines:
                f.write(line)
            print(f)
            f.close()

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>CONFIG>POC3POC2POC1 folders.')
        # ___________________
        
    def Configs(self):
        self.configframe = customtkinter.CTkFrame(self.content_frame, fg_color="transparent", height=600, width=930,
                                                  border_color="Gray", border_width=20, corner_radius=20)
        self.configframe.place(x=40, y=40)

        my_image = customtkinter.CTkImage(dark_image=Image.open("images\\vector.png"), size=(400, 400))
        CTkLabel(self.configframe, image=my_image).place(x=60, y=120)

        self.codeusername = StringVar()
        self.codepassword = StringVar()
        self.codeIPAddNode = StringVar()
        self.codePortNumber = StringVar()

        # ---------------------------------------label-----------------------------------------------------------
        head = CTkLabel(self.configframe, width=120, text="PLEASE ENTER THE DETAILS BELOW (CONFIGURATION)*",
                        font=("Tahoma", 15, "bold"))
        head.place(x=305, y=30)
        # ---------------------------------------USERNAME-----------------------------------------------------------
        user2 = CTkLabel(self.configframe, width=120, text="USERNAME*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        user2.place(x=490, y=135)
        user1 = customtkinter.CTkEntry(master=self.configframe,
                                       width=320,
                                       height=25,
                                       border_width=2,
                                       corner_radius=10, textvariable=self.codeusername)
        user1.place(x=500, y=175)
        # --------------------------------------PASSWORD-------------------------------------------------------------
        kata1 = CTkLabel(self.configframe, width=120, text="PASSWORD*", fg_color="transparent",
                         font=("Tahoma", 15, "bold"))
        kata1.place(x=490, y=215)
        self.password_entry = customtkinter.CTkEntry(master=self.configframe, show="*",
                                      width=320,
                                      height=25,
                                      border_width=2,
                                      corner_radius=10, textvariable=self.codepassword)
        self.password_entry.place(x=500, y=255)
        self.show_password = tkinter.BooleanVar(value=False)
        self.load_eye_images()
        self.password_toggle_button = CTkButton(self.configframe, text="", width=5,
                                    image=self.eye_closed_image, command=self.toggle_password)
        self.password_toggle_button.place(x=830, y=252)
        # --------------------------------------IPAddressNode-------------------------------------------------------------
        apname2 = CTkLabel(self.configframe, width=120, text="IPAddressNode*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=505, y=305)
        apname3 = customtkinter.CTkEntry(master=self.configframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codeIPAddNode)
        apname3.place(x=500, y=345)
        # --------------------------------------Port Number-------------------------------------------------------------
        apname2 = CTkLabel(self.configframe, width=120, text="xxxx*", fg_color="transparent",
                           font=("Tahoma", 15, "bold"))
        apname2.place(x=500, y=385)
        apname3 = customtkinter.CTkEntry(master=self.configframe,
                                         width=320,
                                         height=25,
                                         border_width=2,
                                         corner_radius=10, textvariable=self.codePortNumber)
        apname3.place(x=500, y=425)

        CTkButton(master=self.configframe, text="Submit", width=100, height=30, border_width=0,
                  command=self.HUAWEI).place(x=750, y=500)
        CTkButton(master=self.configframe, text="Back", width=100, height=30, border_width=0,
                  command=self.NMS_Mainpage).place(x=600, y=500)

    def HUAWEI(self):
        # for HUAWEI ATN device only!!!!!!
        self.ID = self.codeusername.get()
        self.Password = self.codepassword.get()
        self.IPAddNode = self.codeIPAddNode.get()
        self.PortNumber = self.codePortNumber.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the CONFIG DESC program is finished running.')

        conn = paramiko.SSHClient()
        conn.set_missing_host_key_policy(paramiko.AutoAddPolicy())

        conn.connect("xxxxx", 22, "xxxx", "xxxx")

        commands = conn.invoke_shell()

        # Please change Device IP #stel -a 121.123.28.28 172.22.61.219 -vpn-ins ce_oam
        commands.send(b"system-v\n")
        time.sleep(.5)
        output1 = commands.recv(65535)
        output1 = output1.decode("utf-8")
        print(output1)

        # Please change Device IP
        commands.send("stel -a xxxx " + self.IPAddNode + " -vpn-ins ce_oam\n")
        time.sleep(.5)
        output2 = commands.recv(65535)
        output2 = output2.decode("utf-8")
        print(output2)

        commands.send(f"{self.ID}\n")
        time.sleep(.5)
        output3 = commands.recv(65535)
        output3 = output3.decode("utf-8")
        print(output3)
        time.sleep(5)

        commands.send(b"y\n")
        time.sleep(.5)
        output4 = commands.recv(65535)
        output4 = output4.decode("utf-8")
        print(output4)

        commands.send(f"{self.Password}\n")
        time.sleep(.5)
        output5 = commands.recv(65535)
        output5 = output5.decode("utf-8")
        print(output5)

        commands.send(b"screen-length 0 temp\n")
        time.sleep(3)
        output6 = commands.recv(65535)
        output6 = output6.decode("utf-8")
        print(output6)

        commands.send(f"dis int des {self.PortNumber}\n")
        time.sleep(3)
        output7 = commands.recv(65535)
        output7 = output7.decode("utf-8")
        print(output7)

        commands.send(b"dis int des\n")
        time.sleep(3)
        output8 = commands.recv(65535)
        output8 = output8.decode("utf-8")
        print(output8)

        commands.send(b"dis ip routing\n")
        time.sleep(3)
        output9 = commands.recv(65535)
        output9 = output9.decode("utf-8")
        print(output9)

        commands.send(b"dis bgp peer\n")
        time.sleep(3)
        output11 = commands.recv(65535)
        output11 = output11.decode("utf-8")
        print(output11)

        commands.send(b"dis cpu-usage history 72hour\n")
        time.sleep(3)
        output14 = commands.recv(65535)
        output14 = output14.decode("utf-8")
        print(output14)

        commands.send(b"dis memory-usage\n")
        time.sleep(3)
        output15 = commands.recv(65535)
        output15 = output15.decode("utf-8")
        print(output15)

        commands.send(b"dis alarm active\n")
        time.sleep(3)
        output16 = commands.recv(65535)
        output16 = output16.decode("utf-8")
        print(output16)

        commands.send(b"dis lldp nei\n")
        time.sleep(3)
        output17 = commands.recv(65535)
        output17 = output17.decode("utf-8")
        print(output17)

        commands.send(b"dis int brief\n")
        time.sleep(3)
        output18 = commands.recv(65535)
        output18 = output18.decode("utf-8")
        print(output18)

        commands.send(b"dis interface | inc CRC\n")
        time.sleep(3)
        output19 = commands.recv(65535)
        output19 = output19.decode("utf-8")
        print(output19)

        commands.send(b"dis interface | inc Half-duplex\n")
        time.sleep(3)
        output20 = commands.recv(65535)
        output20 = output20.decode("utf-8")
        print(output20)

        commands.send(b"dis log | inc /1/\n")
        time.sleep(3)
        output20 = commands.recv(65535)
        output20 = output20.decode("utf-8")
        print(output20)

        commands.send(b"dis lldp nei\n")
        time.sleep(3)
        output21 = commands.recv(65535)
        output21 = output21.decode("utf-8")
        print(output21)

        commands.send(b"dis cur | inc CRC\n")
        time.sleep(3)
        output22 = commands.recv(65535)
        output22 = output22.decode("utf-8")
        print(output22)

        # Display int info
        commands.send("display esn\n")
        time.sleep(3)
        output24 = commands.recv(65535)
        output24 = output24.decode("utf-8")
        print(output24)

        # Display int info
        commands.send(f"dis int {self.PortNumber}\n")
        time.sleep(5)
        output26 = commands.recv(65535)
        output26 = output26.decode("utf-8")
        print(output26)
        time.sleep(40)

        commands.send(b"quit\n")
        time.sleep(3)
        output27 = commands.recv(65535)
        output27 = output27.decode("utf-8")
        print(output27)

        # send output to files
        lines = [output1, output2, output3, output4, output5, output6, output7, output8, output9, output11,
                 output14, output15, output16, output17, output18, output19, output20, output21, output22,
                 output24, output26, output27]

        with open(f'C:\\Automate\\Checklist Output\\CONFIG\\ConfigOutputHuawei_{self.IPAddNode}.txt', 'w') as f:
            for line in lines:
                f.write(line)
            print(f)
            f.close()

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>CONFIG folders.')
        # ___________________
    
    #only for Cisco ISR & ASR only!!!
    def CISCO(self):
    
        self.ID = self.codeusername.get()
        self.Password = self.codepassword.get()
        self.IPAddNode = self.codeIPAddNode.get()
        self.PortNumber = self.codePortNumber.get()
        self.SlotNo = self.codeSlotNo.get()
        self.PortNo = self.codePortNo.get()

        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the CONFIG DESC program is finished running.')

        conn = paramiko.SSHClient()
        conn.set_missing_host_key_policy(paramiko.AutoAddPolicy())

        conn.connect("xxxxx", 22, "xxxx", "xxxx")

        commands = conn.invoke_shell()

        # Please change Device IP #stel -a 121.123.28.28 172.22.61.219 -vpn-ins ce_oam
        commands.send("stel -a xxxxx " + self.IPAddNode + " -vpn-ins ce_oam\n")
        time.sleep(.5)
        output1 = commands.recv(65535)
        output1 = output1.decode("utf-8")
        print(output1)

        commands.send(f"{self.ID}\n")
        time.sleep(.5)
        output2 = commands.recv(65535)
        output2 = output2.decode("utf-8")
        print(output2)

        commands.send(b"y\n")
        time.sleep(.5)
        output3 = commands.recv(65535)
        output3 = output3.decode("utf-8")
        print(output3)

        commands.send(f"{self.Password}\n")
        time.sleep(.5)
        output4 = commands.recv(65535)
        output4 = output4.decode("utf-8")
        print(output4)

        commands.send("en\n")
        time.sleep(3)
        output44 = commands.recv(65535)
        output44 = output44.decode("utf-8")
        print(output44)

        commands.send(f"Show interface {self.PortNumber}\n")
        time.sleep(3)
        output5 = commands.recv(65535)
        output5 = output5.decode("utf-8")
        print(output5)

        commands.send(f"Show hw-module subslot {self.SlotNo} transceiver {self.PortNo} status\n")
        time.sleep(3)
        output6 = commands.recv(65535)
        output6 = output6.decode("utf-8")
        print(output6)

        # Display int info
        commands.send("Show facility alarm-status\n")
        time.sleep(3)
        output7 = commands.recv(65535)
        output7 = output7.decode("utf-8")
        print(output7)

        # Display int info
        commands.send(f"Show logging\n")
        time.sleep(5)
        output8 = commands.recv(65535)
        output8 = output8.decode("utf-8")
        print(output8)
        time.sleep(3)

        commands.send(b"Show ip route\n")
        time.sleep(3)
        output9 = commands.recv(65535)
        output9 = output9.decode("utf-8")
        print(output9)

        commands.send(b"Show ip interface brief\n")
        time.sleep(3)
        output10 = commands.recv(65535)
        output10 = output10.decode("utf-8")
        print(output10)
        
        commands.send(f"Show run interface {self.PortNumber}\n")
        time.sleep(3)
        output11 = commands.recv(65535)
        output11 = output11.decode("utf-8")
        print(output11)

        commands.send(b"End\n")
        time.sleep(3)
        output12 = commands.recv(65535)
        output12 = output12.decode("utf-8")
        print(output12)

        # send output to files
        lines = [output1, output2, output3, output4, output5, output6, 
                 output7, output8, output9, output10, output11, output12]

        with open(f"C:\\Automate\\Checklist Output\\CONFIG\\ConfigOutputCisco-{self.IPAddNode}.txt", "w") as f:
            for line in lines:
                f.write(line)
            print(f)
            f.close()

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Checklist Output>CONFIG folders.')
        # ___________________


    def nmsbreezeair(self):
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is currently in progress.Please wait until the BREEZE AIR program is finished running.')
        ID = "xxxxx"
        PASSWORD = "xxxx"

        # set the download directory and circuit ID
        download_dir_NLOS = "C:\\Automate\\Checklist Output\\BREEZE AIR"

        self.options = Options()

        # this parameter tells Chrome that
        # it should be run without UI (Headless)
        self.options.add_argument("--disable-infobars")
        self.options.add_argument("--start-maximized")
        self.options.add_argument("--disable-extensions")
        self.options.add_argument('--window-size=1920,1080')
        self.options.add_argument("--ignore-certificate-errors")
        self.options.add_argument('--ignore-ssl-errors=yes')
        prefs = {"download.default_directory": download_dir_NLOS}
        self.options.add_experimental_option("prefs", prefs)
        self.options.add_argument("--headless=new")

        self.bot = webdriver.Chrome(service=self.s, options=self.options)
        self.bot.maximize_window()
        urlnmsbreezeair= "http://"

        
        try:
            self.bot.get(urlnmsbreezeair)
        except:
            messagebox.showerror("Error", "Unable to reach the site")
            self.bot.close()

        admin = WebDriverWait(self.bot, 90).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "input[type=\"text\" i] ")))
        admin.clear()

        # login
        user = WebDriverWait(self.bot, 90).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "input[type=\"text\" i] ")))
        user.send_keys(ID)
        self.bot.find_element(By.CSS_SELECTOR, "input[type=\"password\" i] ").send_keys(PASSWORD)
        self.bot.find_element(By.CSS_SELECTOR, "input[type=\"submit\" i] ").click()
        self.bot.implicitly_wait(5)
        time.sleep(40)

        # go to ANURULAG
        a = WebDriverWait(self.bot, 90).until(EC.element_to_be_clickable((By.XPATH, "//span[@id='userId']")))
        ActionChains(self.bot).move_to_element(a).click(a).perform()
        time.sleep(3)

        # go to Database
        b = WebDriverWait(self.bot, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//span[normalize-space()='Database']")))
        ActionChains(self.bot).move_to_element(b).click(b).perform()
        time.sleep(3)

        # go to Export SIte - download report to PC
        c = WebDriverWait(self.bot, 90).until(
            EC.element_to_be_clickable((By.XPATH, "//div[contains(text(),'Export Site...')]")))
        ActionChains(self.bot).move_to_element(c).click(c).perform()
        time.sleep(5)

        # message box appear
        messagebox.showinfo('information',
                            'Hi! Your BIS Automation process is done.Please refer to your Local Disk D>Checklist Output>BREEZE AIR>radios.csv file.')
        self.bot.close()
        time.sleep(5)


if __name__ == "__main__":
    app = MainFrame()
    app.iconbitmap('images\icon.ico')
    app.mainloop()
